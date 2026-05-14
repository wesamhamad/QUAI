#!/usr/bin/env python3
"""
Generate professional QMentor SRS DOCX — Arabic & English versions.
Features: branded header/footer, page numbers, professional tables, cover page.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os

# Brand colors
BLUE_DARK = RGBColor(31, 78, 121)    # #1F4E79
BLUE_MED  = RGBColor(68, 114, 196)   # #4472C4
BLUE_LIGHT= RGBColor(217, 226, 243)  # #D9E2F3
GREEN     = RGBColor(39, 124, 72)    # #277C48
GRAY_DARK = RGBColor(64, 64, 64)
GRAY_MED  = RGBColor(120, 120, 120)
RED       = RGBColor(192, 0, 0)
WHITE     = RGBColor(255, 255, 255)
ROW_ALT   = "E8F0FE"
HDR_BG    = "1F4E79"
HDR_BG2   = "2E75B6"
ACCENT_BG = "D9E2F3"

# ─── Language helper ───
class Lang:
    def __init__(self, is_arabic):
        self.ar = is_arabic
        self.align = WD_ALIGN_PARAGRAPH.RIGHT if is_arabic else WD_ALIGN_PARAGRAPH.LEFT
        self.center = WD_ALIGN_PARAGRAPH.CENTER
        self.font = 'Arial' if is_arabic else 'Calibri'

# ─── Core helpers ───
def setup_doc(lang):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = lang.font
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run('QMentor \u2014 AI-Powered Smart Academic Advisor  |  Qassim University')
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_MED
        run.font.name = lang.font
        # Footer with page number
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('SRS v2.1  |  ')
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_MED
        # Page number field
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="787878"/></w:rPr>'
            f'<w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._p.append(parse_xml(fld_xml))

    return doc


def set_rtl(paragraph):
    if paragraph is None:
        return
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>'))


def add_para(doc, text, lang, bold=False, size=11, color=None, align=None, spacing_after=4):
    p = doc.add_paragraph()
    p.alignment = align or lang.align
    p.paragraph_format.space_after = Pt(spacing_after)
    if lang.ar:
        set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = lang.font
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if lang.ar:
        rPr = run._r.get_or_add_rPr()
        rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} val="1"/>'))
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>'))
        rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{size * 2}"/>'))
    return p


def add_heading(doc, text, lang, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = lang.align
    if lang.ar:
        set_rtl(h)
    for run in h.runs:
        run.font.name = lang.font
        if lang.ar:
            rPr = run._r.get_or_add_rPr()
            rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} val="1"/>'))
            rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>'))
    return h


def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, lang, header_bg=HDR_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if lang.ar:
        tblPr = table._tbl.tblPr
        tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")} val="1"/>'))

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lang.ar:
            set_rtl(p)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = lang.font
        run.font.color.rgb = WHITE
        shade_cell(cell, header_bg)

    # Rows
    for ri, row_data in enumerate(rows):
        for ci, txt in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if not lang.ar else WD_ALIGN_PARAGRAPH.RIGHT
            if lang.ar:
                set_rtl(p)
            run = p.add_run(str(txt))
            run.font.size = Pt(8.5)
            run.font.name = lang.font
            if lang.ar:
                rPr = run._r.get_or_add_rPr()
                rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} val="1"/>'))
            if ri % 2 == 0:
                shade_cell(cell, ROW_ALT)

    doc.add_paragraph('')  # spacing
    return table


def add_blue_box(doc, text, lang):
    """Add a highlighted blue box paragraph."""
    p = add_para(doc, text, lang, bold=True, size=10, color=BLUE_DARK, spacing_after=8)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT_BG}" w:val="clear"/>')
    pPr.append(shd)
    # Border
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="4" w:color="1F4E79"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p


def add_section_divider(doc, lang):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = BLUE_LIGHT


def bullet(doc, text, lang, size=10):
    add_para(doc, f'\u2022  {text}', lang, size=size, spacing_after=2)


def check(doc, text, lang, size=10):
    add_para(doc, f'\u2713  {text}', lang, size=size, color=GREEN, spacing_after=2)


# ═══════════════════════════════════════════════════════════════
#                     ENGLISH VERSION
# ═══════════════════════════════════════════════════════════════
def build_english():
    L = Lang(is_arabic=False)
    doc = setup_doc(L)

    # ── COVER ──
    for _ in range(3):
        doc.add_paragraph('')

    add_para(doc, 'QASSIM UNIVERSITY', L, bold=True, size=22, color=BLUE_DARK, align=L.center, spacing_after=2)
    add_para(doc, 'Deanship of E-Learning & IT \u2014 Digital Transformation Team', L, size=12, color=GRAY_MED, align=L.center, spacing_after=12)
    add_section_divider(doc, L)

    add_para(doc, 'QMentor', L, bold=True, size=44, color=BLUE_DARK, align=L.center, spacing_after=2)
    add_para(doc, 'AI-Powered Smart Academic Advisor', L, bold=True, size=18, color=BLUE_MED, align=L.center, spacing_after=4)
    add_para(doc, '\u0627\u0644\u0645\u0631\u0634\u062f \u0627\u0644\u0623\u0643\u0627\u062f\u064a\u0645\u064a \u0627\u0644\u0630\u0643\u064a', L, size=16, color=GRAY_MED, align=L.center, spacing_after=16)

    add_section_divider(doc, L)
    add_para(doc, 'System Requirements Specification (SRS)', L, bold=True, size=16, align=L.center, spacing_after=4)
    add_para(doc, 'Version 2.1  |  IEEE 830-1998 (modified)  |  April 2026', L, size=11, color=GRAY_MED, align=L.center, spacing_after=16)

    for tag in [
        'Compliant with SDAIA AI Ethics Principles',
        'Innovative Solution \u2014 Qiyas Digital Transformation 2026',
        'Cloud-Native on AWS (Saudi Region) with GPU Servers',
    ]:
        add_para(doc, f'\u25c6  {tag}', L, size=10, color=GREEN, align=L.center, spacing_after=2)

    doc.add_page_break()

    # ── DOCUMENT INFO ──
    add_heading(doc, '1. Document Information', L, level=1)
    add_table(doc,
        ['Field', 'Value'], [
            ['Project Name', 'QMentor \u2014 AI-Powered Smart Academic Advisor'],
            ['Organization', 'Deanship of E-Learning & IT \u2014 Digital Transformation Team'],
            ['Core Technology', 'Agentic AI \u2014 Predictive Analytics \u2014 Student Digital Twin'],
            ['Target Users', 'All students (~70,000), Academic Advisors, Faculty, University Leadership'],
            ['Hosting', 'Cloud-Native on AWS (Saudi Region) \u2014 GPU Servers (G5/SageMaker)'],
            ['Compliance', 'SDAIA AI Ethics \u2014 Saudi PDPL \u2014 NCA ECC'],
            ['Version / Date', '2.1 / 2026-04-07'],
            ['Classification', 'Confidential'],
        ], L)

    doc.add_page_break()

    # ── TOC ──
    add_heading(doc, 'Table of Contents', L, level=1)
    toc = [
        ('1', 'Document Information'),
        ('2', 'Introduction & Scope'),
        ('3', 'Current Advisor Workflow vs. QMentor Automation'),
        ('4', 'Core Modules (6 modules)'),
        ('5', 'Extended Modules (9 modules)'),
        ('  5.1', 'AI Study Plan Builder & Graduation Optimizer (28 requirements)'),
        ('  5.A', 'Current Available API Endpoints (Data Sources)'),
        ('  5.B', 'Comprehensive Risk Criteria \u2014 66 Indicators in 9 Categories'),
        ('  5.C', 'Risk Scoring Methodology (Hybrid Weighted + Critical Override)'),
        ('6', 'Complete Task Autonomy Matrix (60 tasks: Agent vs. Human)'),
        ('7', 'SDAIA AI Ethics Compliance (6 principles, 37 requirements)'),
        ('8', 'Agentic AI Guardrails (5 autonomy levels, 15 guardrails)'),
        ('9', 'Security & Data Classification (4 tiers, 46 requirements)'),
        ('11', 'AWS Infrastructure & Monthly Cost Estimate'),
        ('12', 'Project Phases & Timeline (5 phases, 18 months)'),
    ]
    for num, title in toc:
        bold = not num.startswith(' ')
        add_para(doc, f'{num}{"." if bold else ""}\t{title}', L, bold=bold, size=11, spacing_after=3)

    doc.add_page_break()

    # ── INTRO ──
    add_heading(doc, '2. Introduction & Scope', L, level=1)

    add_heading(doc, '2.1 Purpose', L, level=2)
    add_para(doc,
        'This SRS defines the complete functional and non-functional requirements for QMentor, '
        'an Agentic AI-powered Smart Academic Advisor for Qassim University. It serves as the '
        'single source of truth for the development team, stakeholders, and QA throughout all 5 project phases.', L)

    add_heading(doc, '2.2 Scope', L, level=2)
    add_para(doc, 'QMentor is a cloud-native platform hosted on AWS GPU servers providing:', L)

    scope = [
        'Proactive academic advising via an autonomous AI agent monitoring every student 24/7',
        'Student Digital Twin \u2014 real-time holistic profile from Blackboard LMS, Oracle SIS, and MyQU',
        'Predictive analytics with 66 risk indicators across 9 categories',
        'Multi-channel smart alerts (in-app, email, SMS, WhatsApp) with auto-escalation',
        'Advising chatbot (Arabic + English) answering academic policy questions with RAG',
        'Microsoft Teams meeting scheduling directly from chatbot',
        'AI Study Plan Builder \u2014 builds personalized semester schedules with elective/free course selection based on full historical analysis',
        'Graduation path optimizer with what-if simulations',
        'Mobile app (iOS & Android), WhatsApp/Telegram bots',
        'Student-to-student academic help (voluntary peer matching), faculty dashboards, institutional benchmarking',
        'Emergency academic recovery program with week-by-week rescue plans',
    ]
    for s in scope:
        bullet(doc, s, L)

    add_heading(doc, '2.3 Target Users', L, level=2)
    add_table(doc, ['Role', 'Description'], [
        ['Student', 'All enrolled students (~70,000) across all colleges'],
        ['Academic Advisor', 'Faculty assigned as advisors'],
        ['Course Instructor', 'Faculty teaching specific sections'],
        ['Department Head', 'Heads of academic departments'],
        ['Dean / Vice-Dean', 'College leadership'],
        ['University Leadership', 'Provost, VP Academic Affairs'],
        ['System Admin', 'IT administrators'],
    ], L)

    doc.add_page_break()

    # ── CURRENT vs QMENTOR ──
    add_heading(doc, '3. Current Advisor Workflow vs. QMentor', L, level=1)

    add_blue_box(doc,
        'Based on the "Communication Guide for Academic Advising" (Deanship of Admission & Registration, 2020), '
        'the current system requires 7 manual steps through Oracle SIS with no student data visible. '
        'QMentor automates all steps while preserving the advisor\'s authority.', L)

    add_heading(doc, 'Advisor \u2192 Student: 7 Manual Steps \u2192 Automated', L, level=2)
    add_table(doc,
        ['Step', 'Current (Manual)', 'QMentor (Automated)'], [
            ['1\u20133', 'Login \u2192 Academic menu \u2192 Advisor service', 'Eliminated \u2014 continuous monitoring'],
            ['4', 'View student list (names only, NO data)', 'Digital Twin: GPA, attendance, risk, recommendations'],
            ['5', 'Click "Communicate" per student', 'Auto-triggered when risk detected'],
            ['6', 'Choose channel (portal/SMS 65-char/email), write content', 'AI-authored message, optimal channel auto-selected'],
            ['7', 'Click "Send"', 'Auto-sent for Low/Med; advisor approves High/Critical'],
        ], L)

    add_heading(doc, 'Pain Points Resolved', L, level=2)
    add_table(doc,
        ['Pain Point', 'Impact', 'QMentor Solution'], [
            ['SMS limited to 65 characters', 'Cannot convey meaningful advice', 'Multi-channel, unlimited content'],
            ['Portal messages = banner ads', 'Students ignore them', 'Push notifications + chatbot follow-up'],
            ['No performance data visible', 'Generic messages', 'Full Digital Twin in every interaction'],
            ['Email only (one-way, slow)', 'Response takes days', 'Chatbot: instant AI + human escalation'],
            ['Instructor sees names only', 'Can\'t identify at-risk students', 'Section risk heatmap + auto-alerts'],
            ['All communication is reactive', 'At-risk students missed', '24/7 proactive monitoring'],
        ], L)

    doc.add_page_break()

    # ── CORE MODULES ──
    add_heading(doc, '4. Core Modules', L, level=1)

    # MON
    add_heading(doc, '4.1 Continuous Monitoring Agent (FR-MON)', L, level=2)
    add_blue_box(doc, 'The core Agentic AI loop running autonomously 24/7 \u2014 pulls data, detects anomalies, updates Digital Twins.', L)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['MON-01', 'Pull attendance from Blackboard every 2 hours', 'Must'],
        ['MON-02', 'Pull grades from Oracle SIS every 6 hours', 'Must'],
        ['MON-03', 'Pull assignment submissions every 4 hours', 'Must'],
        ['MON-04', 'Pull registration & academic standing daily', 'Must'],
        ['MON-05', 'Pull student requests from MyQU daily', 'Must'],
        ['MON-06', 'Track LMS login frequency and content interaction', 'Should'],
        ['MON-07', 'Detect data anomalies (sudden drops, patterns)', 'Must'],
        ['MON-08', 'Log all pulls with timestamps + success/failure', 'Must'],
        ['MON-09', 'Retry failed pulls (3x, exponential backoff)', 'Must'],
    ], L)

    # PAE
    add_heading(doc, '4.2 Predictive Analytics Engine (FR-PAE)', L, level=2)
    add_blue_box(doc, '66 indicators across 9 categories. Hybrid scoring: weighted (0\u2013100) + critical override + auto-escalation.', L)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['PAE-01', 'Classify each student: Low / Medium / High / Critical', 'Must'],
        ['PAE-02', 'Predict course failure probability per student per course', 'Must'],
        ['PAE-03', 'Predict probability of reaching absence limit', 'Must'],
        ['PAE-04', 'Detect declining trends over 2+ weeks', 'Must'],
        ['PAE-06', 'Generate risk score 0\u2013100 with top 3 contributing factors', 'Should'],
        ['PAE-07', 'Retrain models each semester with outcome data', 'Must'],
        ['PAE-09', 'Achieve \u226585% prediction accuracy', 'Must'],
        ['PAE-10', 'Provide explainability for every flag', 'Must'],
    ], L)

    # SDT
    add_heading(doc, '4.3 Student Digital Twin (FR-SDT)', L, level=2)
    add_table(doc, ['Section', 'Contents'], [
        ['Academic Profile', 'GPA, credits earned/remaining, warnings, expected graduation grade'],
        ['Study Plan', 'Completed/remaining courses, prerequisites, optimal path'],
        ['Behavior Metrics', 'Attendance rate, submission rate, LMS hours per course'],
        ['Risk Points', 'At-risk courses, active alerts, intervention history + outcomes'],
        ['AI Recommendations', 'Learning content, courses to drop/add, advisor meeting suggestion'],
        ['Timeline', 'Full academic event history from admission to present'],
    ], L)

    # ALT
    add_heading(doc, '4.4 Smart Alert System (FR-ALT)', L, level=2)
    add_table(doc, ['Channel', 'When', 'Example'], [
        ['In-app push', 'Daily nudges', '"You missed CS101 lecture today"'],
        ['Email', 'Weekly summaries', '"3 courses need attention this week"'],
        ['SMS', 'Urgent only', '"Absence at 20% \u2014 limit is 25%"'],
        ['WhatsApp', 'Urgent + interactive', '"Student Ahmed is High risk \u2014 see Digital Twin"'],
        ['Advisor alert', 'Human intervention needed', 'Full Digital Twin + AI intervention plan'],
    ], L)

    # BOT
    add_heading(doc, '4.5 Advising Chatbot (FR-BOT)', L, level=2)
    add_blue_box(doc, '24/7 Arabic + English chatbot with RAG on university bylaws, personalized answers from Digital Twin, and Teams meeting scheduling.', L)

    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['BOT-01', 'Conversational chatbot via web and mobile', 'Must'],
        ['BOT-02', 'Arabic (MSA + Saudi dialect) and English', 'Must'],
        ['BOT-03', 'Answer academic policy questions (bylaws, registration, deadlines)', 'Must'],
        ['BOT-04', 'Personalized advice from Digital Twin data', 'Must'],
        ['BOT-05', 'Escalate to human advisor for complex/emotional issues', 'Must'],
        ['BOT-06', 'RAG with university regulations knowledge base', 'Must'],
        ['BOT-12', 'Citation/source for every policy answer', 'Must'],
        ['BOT-13', 'Answer regulation questions from official bylaws', 'Must'],
        ['BOT-14', 'Personalized answers combining policy + student data', 'Must'],
        ['BOT-15', 'Schedule Microsoft Teams meetings with advisor', 'Should'],
    ], L)

    add_heading(doc, 'Academic Policy Knowledge Base (RAG)', L, level=3)
    add_table(doc, ['Official Document', 'Example Questions'], [
        ['Study & Exams Bylaw', '"What are First Honors requirements?" \u2014 "How many retakes allowed?" \u2014 "When does exam barring happen?"'],
        ['Admission & Registration', '"Requirements to transfer colleges?" \u2014 "How to defer a semester?"'],
        ['Warnings & Dismissal', '"How many warnings before dismissal?" \u2014 "Can I appeal?"'],
        ['Academic Calendar', '"Last day for add/drop?" \u2014 "When do finals start?"'],
        ['Self-Service Regs', '"How to request a transcript?" \u2014 "How to change major?"'],
    ], L)

    add_heading(doc, 'Personalized Policy + Data Answers', L, level=3)
    add_table(doc, ['Question', 'Policy-Only', 'QMentor Personalized'], [
        ['"First Honors requirements?"', 'GPA \u22654.75, no failures', 'Your GPA is 4.82, no failures \u2014 you qualify \u2705'],
        ['"Can I drop MATH201?"', 'Available until week 10', 'Yes, but it\'s prerequisite for 2 courses \u2014 delays graduation 1 semester \u26a0\ufe0f'],
        ['"Absences left in CS101?"', 'Max 25%', '4 of 8 used (50%). 4 remaining. Miss 5 = danger zone \u26a0\ufe0f'],
        ['"Need advisor meeting"', 'Contact via portal', 'Dr. Al-Harbi, Tue 10AM available. Book Teams meeting? \u2705'],
    ], L)

    # ADV
    add_heading(doc, '4.6 Advisor Dashboard (FR-ADV)', L, level=2)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['ADV-01', 'All students sorted by risk level (critical first)', 'Must'],
        ['ADV-02', 'Aggregate stats: total, at-risk per level, trends', 'Must'],
        ['ADV-03', 'One-click Digital Twin access', 'Must'],
        ['ADV-04', 'AI-generated intervention plan per at-risk student', 'Must'],
        ['ADV-05', 'Log intervention notes and outcomes', 'Must'],
        ['ADV-06', 'Alert history + student response tracking', 'Must'],
        ['ADV-08', 'Calendar view of upcoming meetings', 'Should'],
    ], L)

    doc.add_page_break()

    # ── EXTENDED MODULES ──
    add_heading(doc, '5. Extended Modules', L, level=1)

    # ── GPO (the big new feature) ──
    add_heading(doc, '5.1 AI Study Plan Builder & Graduation Optimizer (FR-GPO)', L, level=2)
    add_blue_box(doc,
        'NEW: Analyzes the student\'s complete historical record to build a personalized study plan \u2014 '
        'selects elective and free courses, generates optimized semester-by-semester schedules, and simulates what-if scenarios.', L)

    add_heading(doc, '5.1.1 Core Graduation Path', L, level=3)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['GPO-01', 'Calculate remaining semesters based on current progress', 'Must'],
        ['GPO-02', 'Generate semester-by-semester plan respecting all prerequisites', 'Must'],
        ['GPO-03', 'Offer multiple paths: fastest, balanced workload, GPA-safe', 'Should'],
        ['GPO-04', 'Warn when choices delay graduation', 'Must'],
        ['GPO-05', 'Consider course availability (Fall-only, Spring-only, alternate years)', 'Should'],
        ['GPO-06', 'Auto-recalculate when student drops or fails a course', 'Must'],
        ['GPO-07', 'What-if simulations ("If I drop this \u2192 new grad date + GPA impact")', 'Should'],
        ['GPO-08', 'Suggest summer courses to accelerate graduation', 'Could'],
    ], L)

    add_heading(doc, '5.1.2 Intelligent Elective & Free Course Selection', L, level=3)
    add_para(doc, 'Builds personalized elective recommendations by analyzing the student\'s complete academic history, performance patterns, career interests, and peer outcomes.', L, size=10)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['GPO-09', 'Analyze full transcript to identify strengths, weaknesses, subject-area patterns', 'Must'],
        ['GPO-10', 'Classify remaining requirements: mandatory core, dept elective, college elective, free elective', 'Must'],
        ['GPO-11', 'Recommend free electives based on: GPA trend in similar areas, historical grade distributions, career alignment, schedule fit', 'Must'],
        ['GPO-12', 'Recommend dept/college electives by: peer GPA correlation, career relevance, prerequisite chain impact', 'Must'],
        ['GPO-13', 'Flag "GPA-booster" electives \u2014 courses where same-major students historically score \u22651 grade higher', 'Should'],
        ['GPO-14', 'Flag "GPA-risk" electives \u2014 high failure rate for student\'s major \u2014 warn before selection', 'Should'],
        ['GPO-15', 'Factor instructor-specific pass rates when available', 'Could'],
        ['GPO-16', 'Student sets preferences: interest areas, preferred days/times, workload level', 'Should'],
        ['GPO-17', 'NEVER auto-register \u2014 all recommendations need student/advisor review', 'Must'],
    ], L)

    add_heading(doc, '5.1.3 Semester Schedule Builder', L, level=3)
    add_para(doc, 'For each upcoming semester, generates a complete optimized timetable.', L, size=10)
    add_table(doc, ['ID', 'Requirement', 'Pri'], [
        ['GPO-18', 'Generate complete draft schedule: core + electives + free = recommended credit hours', 'Must'],
        ['GPO-19', 'Respect optimal load (12\u201318 hrs) based on student\'s historical GPA under different loads', 'Must'],
        ['GPO-20', 'No time conflicts: lectures, labs, or back-to-back final exams', 'Must'],
        ['GPO-21', 'Balance daily workload \u2014 no 8-hour continuous days or 2-day packing', 'Should'],
        ['GPO-22', 'Consider time-of-day performance (if lower grades at 8AM, prefer later sections)', 'Could'],
        ['GPO-23', 'Generate 2\u20133 ranked options: graduation impact, GPA-safety, schedule comfort', 'Should'],
        ['GPO-24', 'Display each option as visual weekly timetable (Sun\u2013Thu) with color-coded blocks', 'Must'],
        ['GPO-25', 'Lock/unlock courses or time slots, regenerate around constraints', 'Should'],
        ['GPO-26', 'Export chosen schedule as PDF, one-click share with advisor', 'Should'],
        ['GPO-27', 'Notify when seat opens in preferred section that was full', 'Should'],
        ['GPO-28', 'Alert if prerequisite status changes (fail prereq in current semester)', 'Must'],
    ], L)

    add_heading(doc, 'Example: AI-Generated Study Plan', L, level=3)
    example_text = (
        'Student: Ahmed M. | Major: CS | GPA: 3.45 | Semester 7 of 8\n\n'
        'ANALYSIS OF FULL ACADEMIC HISTORY:\n'
        '  Strengths: Programming (avg A-), Databases (avg B+)\n'
        '  Weaknesses: Math (avg C+), Theory (avg C)\n'
        '  Optimal load: 15 credit hours (GPA drops at 18+)\n'
        '  Best performance: 10 AM\u20132 PM sections\n\n'
        'REMAINING REQUIREMENTS:\n'
        '  Core: CS401, CS420, CS490 (Capstone)\n'
        '  Dept Elective: 2 courses needed\n'
        '  Free Elective: 1 course needed\n'
        '  Total: 6 courses across 2 semesters\n\n'
        'RECOMMENDED SCHEDULE (Option A \u2014 GPA-Safe):\n'
        '  Sun/Tue/Thu 10:00  CS401 (Core)\n'
        '  Sun/Tue     11:00  CS420 (Core)\n'
        '  Mon/Wed     12:00  MGT201 (Free Elective \u2014 GPA booster)\n'
        '  Mon/Wed      1:00  CS455-AI (Dept Elective \u2014 career aligned)\n\n'
        '  \u2605 Why MGT201? CS students historically score avg 4.2/5.0 (A-)\n'
        '  \u2605 Why CS455? Aligns with career interest in AI/ML\n'
        '  Credit hours: 15 | No exam conflicts | Graduation: on track\n\n'
        'Option B (Fastest): Add CS490 \u2192 graduate this semester\n'
        '  \u26a0 Warning: 18 hours \u2014 GPA typically drops 0.3 at this load'
    )
    p = add_para(doc, example_text, L, size=9, color=GRAY_DARK, spacing_after=8)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'))

    # Other extended modules (condensed) — removed: Career AI, Gamification, Parent Portal
    ext_modules = [
        ('5.2 Student-to-Student Academic Help', [
            ('PTM-01', 'Match high-performing students with struggling peers (fellow students only — not private tutors or university staff) by course, schedule, language', 'Should'),
            ('PTM-03', 'Both helper and helped student must opt-in; either can opt-out at any time', 'Must'),
            ('PTM-04', 'Track help sessions and impact on grades', 'Should'),
            ('PTM-05', 'Sessions are proposed as suggestions — students decide whether to accept, reschedule, or decline', 'Must'),
            ('PTM-06', 'Respect gender segregation policies', 'Must'),
        ]),
        ('5.3 Sentiment & Well-being Analysis', [
            ('SWB-01', 'Detect distress signals in chatbot conversations', 'Should'),
            ('SWB-02', 'Detect burnout patterns (sudden disengagement across all courses)', 'Should'),
            ('SWB-04', 'Provide self-help resources and counseling center info', 'Must'),
            ('SWB-05', 'NEVER diagnose or provide clinical mental health advice', 'Must'),
        ]),
        ('5.4 Mobile Application (Flutter)', [
            ('MOB-01', 'iOS 15+ / Android 12+ via Flutter', 'Must'),
            ('MOB-02', 'Push notifications for all alert types', 'Must'),
            ('MOB-03', 'Chatbot interface with voice input', 'Must'),
            ('MOB-05', 'Biometric authentication (Face ID / fingerprint)', 'Should'),
            ('MOB-07', 'Nafath national digital identity login', 'Should'),
            ('MOB-08', 'Dark mode and accessibility (VoiceOver, TalkBack)', 'Must'),
        ]),
        ('5.5 WhatsApp & Telegram Integration', [
            ('MSG-01', 'Send alerts via WhatsApp Business API', 'Should'),
            ('MSG-02', 'Basic chatbot Q&A via WhatsApp', 'Should'),
            ('MSG-03', 'Telegram bot as alternative channel', 'Could'),
            ('MSG-04', 'Student chooses preferred messaging platform', 'Should'),
        ]),
        ('5.6 Faculty Analytics Dashboard', [
            ('FAC-01', 'Section risk heatmap for instructor\u2019s students', 'Should'),
            ('FAC-02', 'Auto-alert instructor about at-risk students in their course', 'Must'),
            ('FAC-03', 'Department-level analytics (pass rates, risk distribution, trends)', 'Should'),
            ('FAC-04', 'College-level executive dashboard for deans', 'Should'),
        ]),
        ('5.7 Emergency Academic Recovery Program', [
            ('EAR-01', 'Auto-generate personalized recovery plan at Critical risk', 'Must'),
            ('EAR-02', 'Include priority courses, study schedule, recommended content, and suggested fellow students who can help', 'Must'),
            ('EAR-03', 'Week-by-week checklist with measurable milestones', 'Should'),
            ('EAR-06', 'What-if GPA simulator ("If I score X on remaining exams, GPA = Y")', 'Should'),
            ('EAR-07', 'Recovery plan shared with student + advisor', 'Must'),
        ]),
        ('5.8 Institutional Benchmarking & Reporting', [
            ('IBR-01', 'University-wide risk distribution dashboards (by college, department)', 'Should'),
            ('IBR-02', 'Semester-over-semester retention, GPA trends, intervention effectiveness', 'Should'),
            ('IBR-03', 'Accreditation-ready reports (NCAAA, ABET)', 'Could'),
            ('IBR-05', 'Export reports: PDF, Excel, JSON for ministry submissions', 'Should'),
            ('IBR-07', 'Track QMentor\u2019s own impact metrics (time saved, interventions, outcomes)', 'Must'),
        ]),
        ('5.9 Smart Scheduling Assistant', [
            ('SSA-01', 'Suggest optimal timetable combinations based on study plan', 'Should'),
            ('SSA-02', 'Detect time conflicts between courses, labs, and final exams', 'Must'),
            ('SSA-03', 'Warn when schedule creates 3+ back-to-back exams', 'Should'),
            ('SSA-06', 'Integrate with Graduation Optimizer to prioritize bottleneck courses', 'Should'),
            ('SSA-07', 'Notify when seat opens in previously full section', 'Should'),
        ]),
    ]

    for title, reqs in ext_modules:
        add_heading(doc, title, L, level=2)
        add_table(doc, ['ID', 'Requirement', 'Pri'],
                  [[r[0], r[1], r[2]] for r in reqs], L)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  SECTION 5.A: AVAILABLE API ENDPOINTS (DATA SOURCES)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, '5.A Current Available API Endpoints (Data Sources)', L, level=1)
    add_blue_box(doc,
        'These are the actual REST API endpoints currently available in the codebase for pulling student data. '
        'Each endpoint maps to specific risk indicators. Gaps are flagged.', L)

    add_heading(doc, 'UniversityApiClient (Oracle SIS + MyQU)', L, level=2)
    add_para(doc, 'Base URL: configured in quai.blackboard.api_base_url | Auth: Service Token (Bearer/X-Service-Token) | Caching: 300s default', L, size=9, color=GRAY_MED)
    add_table(doc, ['Method', 'Endpoint', 'Returns', 'Used By Indicators'], [
        ['getStudentProfile()', 'GET /api/v3/me', 'Name, ID, GPA, major, college, contact info', 'G-04, G-05, AC-01, AC-04, AC-05'],
        ['getAbsences()', 'GET /absences-with-details', 'Detailed absence records with dates, reasons, courses', 'A-01 to A-08 (all attendance)'],
        ['getAcademicTransactions()', 'GET /student-academic-transactions', 'Drop/add, withdrawals, warnings history', 'R-01, R-02, R-03, R-06, AC-02'],
        ['getCurrentCourses()', 'GET /student/courses', 'Enrolled courses with bb_course_id, code, name', 'AC-03, AC-06, AC-07, R-04'],
        ['getStudentPlan()', 'GET /student-plan', 'Academic plan, completed/remaining courses', 'P-01 to P-05 (graduation path)'],
        ['getAcademicPlanByMajor()', 'GET /api/v1/academic_plan/{majorNo}', 'Full degree requirements for a major', 'GPO module (elective classification)'],
        ['getTimeTable()', 'GET /time-table', 'Weekly class schedule', 'T-01, T-04 (schedule conflicts)'],
        ['getFinalExams()', 'GET /final-exams', 'Final exam dates, times, rooms', 'T-01, T-03 (exam conflicts)'],
        ['getAcademicCalendar()', 'GET /academic-calendar', 'Semester dates, deadlines (cached 3600s)', 'AGT-09 (temporal awareness)'],
        ['getAdvisorInfo()', 'GET /academic-advisor', 'Advisor name, contact, office', 'BOT-15 (Teams scheduling)'],
        ['getRewards()', 'GET /rewards', 'Financial rewards, scholarships', 'Institutional reporting'],
        ['getStudentSkills()', 'GET /api/v1/skills/student/{id}', 'Student skills and competencies', 'Student-to-student matching, future career module'],
        ['getDepartmentsAndColleges()', 'GET /api/v1/departments_and_collages', 'All departments & colleges (cached 3600s)', 'College-specific thresholds'],
        ['fetchMultiple()', 'GET (concurrent pool)', 'Parallel fetch of multiple endpoints', 'Digital Twin builder (batch)'],
    ], L)

    add_heading(doc, 'BlackboardApiClient (LMS)', L, level=2)
    add_para(doc, 'Base URL: configured in quai.blackboard.api_base_url | Auth: JWT Bearer/Service Token | Timeout: 30s', L, size=9, color=GRAY_MED)
    add_table(doc, ['Method', 'Endpoint', 'Returns', 'Used By Indicators'], [
        ['getCourseGrades()', 'GET /blackboard/courses/{id}/grades', 'Grade records: quizzes, midterms, assignments, final', 'G-01 to G-03, G-06 to G-11'],
        ['getAllCourseGrades()', 'GET (iterates all courses)', 'Grades across ALL enrolled courses with metadata', 'Cross-course risk analysis'],
        ['getCourseContent()', 'GET /blackboard/courses/{id}/contents', 'Content tree structure + access data', 'E-03 (content access rate)'],
        ['getAllCourseContents()', 'GET (iterates all courses)', 'All content across all courses', 'Cross-course engagement'],
        ['getCourseAttachments()', 'GET /blackboard/courses/{id}/contents/{cid}/attachments', 'Assignment attachments + submission data', 'S-01 to S-06 (submissions)'],
        ['getAnnouncements()', 'GET /blackboard/announcements', 'All announcements across courses', 'Communication tracking'],
        ['getCourseAnnouncements()', 'GET /blackboard/courses/{id}/announcements', 'Per-course announcements', 'Per-course engagement'],
        ['downloadAttachment()', 'GET .../attachments/{id}/download', 'Binary file download (PDF, etc.)', 'File extraction pipeline'],
    ], L)

    add_heading(doc, 'AI & Processing Services', L, level=2)
    add_table(doc, ['Service', 'Endpoint', 'Purpose'], [
        ['OllamaService', 'POST /v1/chat/completions', 'LLM inference (ALLaM:7b Arabic, Llama 3.1 English)'],
        ['AiCompletionService', 'POST {provider}/v1/chat/completions', 'Multi-provider AI (OpenAI-compatible)'],
        ['EmbeddingService', 'POST /api/embed', 'Vector embeddings for semantic search'],
        ['KnowledgeRetrievalService', 'POST /search', 'RAG semantic search for policy knowledge base'],
        ['WebSearchService', 'GET /search?q=...&format=json', 'SearXNG meta-search (supplementary context)'],
        ['FileExtractionService', 'POST /extract', 'PDF/document text extraction with OCR'],
    ], L)

    add_heading(doc, '\u26a0\ufe0f API Gaps (Missing Endpoints)', L, level=2)
    add_blue_box(doc,
        'The following risk indicators require data NOT currently available from existing APIs. '
        'These must be resolved during Phase 1 integration.', L)
    add_table(doc, ['Indicator', 'Data Needed', 'Gap', 'Proposed Resolution'], [
        ['E-01: LMS login frequency', 'Per-student login count per week', 'Blackboard REST API does not expose login frequency', 'Blackboard Data (Analytics) integration or Building Block'],
        ['E-02: Time on LMS', 'Weekly hours spent per student', 'Not available in standard Blackboard API', 'Blackboard Data warehouse or custom activity log'],
        ['E-04: Engagement drop', 'Activity comparison vs. 3-week average', 'Requires historical activity data', 'Build custom activity tracker from available content access logs'],
        ['E-05: Forum participation', 'Discussion post count', 'Blackboard discussions API not integrated', 'Add /blackboard/courses/{id}/discussions endpoint'],
        ['E-06: Days since last login', 'Last login timestamp', 'Not in current API scope', 'Blackboard session API or Data warehouse'],
        ['S-03: Last-minute submissions', 'Submission timestamp vs. deadline', 'Timestamps available in attachments but deadline data may be incomplete', 'Enrich assignment metadata from Blackboard'],
    ], L)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  SECTION: COMPLETE 66 RISK CRITERIA
    # ══════════════════════════════════════════════════════════
    add_heading(doc, '5.B Comprehensive Risk Criteria \u2014 66 Indicators in 9 Categories', L, level=1)

    add_blue_box(doc,
        'The Predictive Analytics Engine evaluates students using 66 measurable indicators. '
        'Each maps to a specific API endpoint. Thresholds define Low/Medium/High/Critical levels.', L)

    # Summary table
    add_table(doc, ['Category', 'Count', 'Key Examples'], [
        ['A: Attendance', '8', 'Absence rate, consecutive absences, distance from barring, absence after warning'],
        ['G: Grades', '11', 'Midterm, quizzes, cumulative GPA, grade trend, gap vs. section average'],
        ['S: Assignments', '6', 'Missing submissions, late submissions, last-minute, zero-score'],
        ['E: LMS Engagement', '6', 'Login frequency, time on platform, content access, sudden drop'],
        ['AC: Academic Standing', '7', 'Warnings, cumulative failures, credit hours behind plan'],
        ['R: Registration Behavior', '6', 'Course drops, withdrawals, major changes, overload/underload'],
        ['T: Schedule & Exams', '5', 'Exam conflicts, missed midterms/finals, packed schedule'],
        ['C: Compound Indicators', '12', 'Silent withdrawal, absence+declining grades, collapse after success'],
        ['P: Graduation Path', '5', 'Behind plan, core courses not passed, near max enrollment'],
    ], L)

    # A: Attendance
    add_heading(doc, 'Category A: Attendance (8 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getAbsences() \u2192 GET /absences-with-details', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['A-01', 'Absence % per course', '5\u201310%', '11\u201315%', '16\u201320%', '>20% (barred)'],
        ['A-02', 'Consecutive absences', '1', '2', '3', '\u22654'],
        ['A-03', 'Distance from barring limit', '>10% remaining', '5\u201310%', '2\u20135%', '\u22642%'],
        ['A-04', 'Absence after academic warning', '\u2014', 'Any absence', '2+ absences', 'Pattern continues'],
        ['A-05', 'Absence pattern (specific days)', 'Random', 'Recurring day', 'Recurring pattern', 'All morning/week'],
        ['A-06', 'Medical vs. unexcused ratio', '>80% medical', '50\u201380%', '20\u201350%', '<20% medical'],
        ['A-07', 'Absence trend (3-week window)', 'Improving', 'Stable', 'Worsening', 'Accelerating'],
        ['A-08', 'Cross-course absence correlation', '1 course', '2 courses', '3+ courses', 'All courses'],
    ], L)

    # G: Grades
    add_heading(doc, 'Category G: Grades (11 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getCourseGrades() \u2192 GET /blackboard/courses/{id}/grades + getStudentProfile()', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['G-01', 'Midterm exam score', '\u226570%', '50\u201369%', '35\u201349%', '<35%'],
        ['G-02', 'Quiz average', '\u226570%', '50\u201369%', '35\u201349%', '<35%'],
        ['G-03', 'Assignment average', '\u226570%', '50\u201369%', '35\u201349%', '<35%'],
        ['G-04', 'Cumulative GPA', '\u22653.0', '2.0\u20132.99', '1.5\u20131.99', '<1.5'],
        ['G-05', 'Semester GPA trend', 'Improving', 'Stable (\u00b10.2)', 'Dropped 0.2\u20130.5', 'Dropped >0.5'],
        ['G-06', 'Grade gap vs. section avg', 'Within 1\u03c3', '1\u20131.5\u03c3 below', '1.5\u20132\u03c3 below', '>2\u03c3 below'],
        ['G-07', 'Performance trajectory', 'Rising', 'Flat', 'Declining', 'Steep decline'],
        ['G-08', 'Courses at risk of failure', '0', '1', '2', '\u22653'],
        ['G-09', 'Failed courses (current sem)', '0', '1', '2', '\u22653'],
        ['G-10', 'Historical failure count', '0', '1\u20132', '3\u20134', '\u22655'],
        ['G-11', 'Predicted final grade (ML)', '\u2265C', 'D+ to C-', 'D to D-', 'F'],
    ], L)

    # S: Assignments
    add_heading(doc, 'Category S: Assignments (6 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getCourseAttachments() \u2192 GET /blackboard/courses/{id}/contents/{cid}/attachments', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['S-01', 'Missing assignment %', '0\u201310%', '11\u201325%', '26\u201350%', '>50%'],
        ['S-02', 'Late submission count', '0\u20131', '2\u20133', '4\u20135', '\u22656'],
        ['S-03', 'Last-minute submissions (<1hr)', '<20%', '20\u201340%', '41\u201360%', '>60%'],
        ['S-04', 'Assignment score trend', 'Improving', 'Stable', 'Declining', 'Steep decline'],
        ['S-05', 'Zero-score submissions', '0', '1', '2', '\u22653'],
        ['S-06', 'Missing high-weight assignments (>10%)', '0', '1', '2', '\u22653'],
    ], L)

    # E: LMS Engagement
    add_heading(doc, 'Category E: LMS Engagement (6 Indicators)', L, level=2)
    add_para(doc, '\u26a0\ufe0f Data Source: PARTIALLY AVAILABLE \u2014 some require additional Blackboard API endpoints (see API Gaps)', L, size=9, color=RED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical', 'API Status'], [
        ['E-01', 'Login frequency/week', '\u22655', '3\u20134', '1\u20132', '0', '\u26a0\ufe0f GAP'],
        ['E-02', 'Time on LMS (weekly hrs)', '\u22655', '2\u20135', '0.5\u20132', '<0.5', '\u26a0\ufe0f GAP'],
        ['E-03', 'Content access rate (%)', '\u226570%', '40\u201369%', '15\u201339%', '<15%', '\u2705 Available'],
        ['E-04', 'Sudden engagement drop', '<20%', '20\u201340%', '41\u201370%', '>70%', '\u26a0\ufe0f GAP'],
        ['E-05', 'Forum participation', 'Active', 'Occasional', 'Rare', 'None', '\u26a0\ufe0f GAP'],
        ['E-06', 'Days since last login', '0\u20132', '3\u20135', '6\u201310', '>10', '\u26a0\ufe0f GAP'],
    ], L)

    # AC: Academic Standing
    add_heading(doc, 'Category AC: Academic Standing (7 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getStudentProfile() + getAcademicTransactions()', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['AC-01', 'Active academic warnings', '0', '1st warning', '2nd warning', '3rd+'],
        ['AC-02', 'Cumulative failed courses', '0', '1\u20132', '3\u20135', '\u22656'],
        ['AC-03', 'Credit hours behind plan', 'On track', '3\u20136 hrs', '7\u201312 hrs', '>12 hrs'],
        ['AC-04', 'Probation status', 'None', 'Watch list', 'Probation', 'Final probation'],
        ['AC-05', 'GPA distance from min (2.0)', '>1.0 above', '0.5\u20131.0', '0.1\u20130.5', 'Below 2.0'],
        ['AC-06', 'Failed prereqs blocking progress', '0', '1', '2', '\u22653'],
        ['AC-07', 'Repeated course count', '0', '1', '2\u20133', '\u22654'],
    ], L)

    # R: Registration Behavior
    add_heading(doc, 'Category R: Registration Behavior (6 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getAcademicTransactions() + getCurrentCourses()', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['R-01', 'Courses dropped this semester', '0', '1', '2', '\u22653'],
        ['R-02', 'Withdrawal requests history', '0', '1', '2', '\u22653'],
        ['R-03', 'Major change requests', '0', '1 (recent)', '2+ attempts', 'Under processing'],
        ['R-04', 'Course load vs recommended', 'Normal (15\u201318)', 'Light (12\u201314)', 'Overload (>18) or <12', '<9 hrs'],
        ['R-05', 'Late registration pattern', 'On-time', '1\u20132 days late', '3\u20135 days', '>5 days'],
        ['R-06', 'Add/drop frequency (first 2 wks)', '0\u20131', '2\u20133', '4\u20135', '\u22656'],
    ], L)

    # T: Schedule & Exams
    add_heading(doc, 'Category T: Schedule & Exams (5 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getTimeTable() + getFinalExams() + getCourseGrades()', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['T-01', 'Final exam same-day conflicts', '0\u20131', '2', '3', '3+ back-to-back'],
        ['T-02', 'Missed midterm exams', '0', '1', '2', '\u22653'],
        ['T-03', 'Missed final (prev semester)', '0', '1', '2', '\u22653'],
        ['T-04', 'Schedule gap issues', 'Balanced', 'Minor', '5+ hrs consecutive', '7+ hrs'],
        ['T-05', 'Exam week absence', 'Present', '1 absence', '2+', 'All review missed'],
    ], L)

    # C: Compound Indicators
    add_heading(doc, 'Category C: Compound Indicators (12 Indicators)', L, level=2)
    add_para(doc, 'Combine multiple simple indicators to detect complex behavioral patterns.', L, size=10)
    add_table(doc, ['ID', 'Indicator', 'Combination Logic', 'Triggers'], [
        ['C-01', 'Silent withdrawal', 'E-06>7d + S-01>50% + A-01>15%', 'Critical'],
        ['C-02', 'Absence + declining grades', 'A-01>10% AND G-07="Declining"', 'High'],
        ['C-03', 'High effort, low results', 'E-02\u22655hrs AND G-01<50%', 'Medium (needs help)'],
        ['C-04', 'Collapse after success', 'G-05 dropped>0.5 AND prev GPA\u22653.5', 'High (well-being)'],
        ['C-05', 'Multi-course disengagement', 'E-04>50% drop in \u22653 courses', 'Critical'],
        ['C-06', 'Drop cascade', 'R-01\u22652 AND G-08\u22651 AND A-01\u2191', 'High'],
        ['C-07', 'Grade-attendance disconnect', 'A-01<5% BUT G-01<40%', 'Medium (learning)'],
        ['C-08', 'Procrastination spiral', 'S-03>60% AND S-01\u2191 over 3 weeks', 'Medium'],
        ['C-09', 'Social isolation signal', 'E-05=None AND no advisor meeting 4wk', 'Medium (well-being)'],
        ['C-10', 'Registration confusion', 'R-06\u22654 AND R-04\u2260Normal AND AC-06\u22651', 'Medium'],
        ['C-11', 'Graduation path drift', 'AC-03>6hrs AND R-01\u22651 AND G-04\u2193', 'High'],
        ['C-12', 'Recovery failure', 'Previous intervention + no improvement 2wk', 'Auto-escalate +1'],
    ], L)

    # P: Graduation Path
    add_heading(doc, 'Category P: Graduation Path (5 Indicators)', L, level=2)
    add_para(doc, 'Data Source: getStudentPlan() + getAcademicPlanByMajor()', L, size=9, color=GRAY_MED)
    add_table(doc, ['ID', 'Indicator', 'Low', 'Medium', 'High', 'Critical'], [
        ['P-01', 'Semesters behind plan', '0', '1', '2', '\u22653'],
        ['P-02', 'Core courses not yet passed', 'On track', '1 delayed', '2\u20133', '\u22654'],
        ['P-03', 'Approaching max enrollment', '>4 sem remaining', '2\u20134', '1', 'At limit'],
        ['P-04', 'Required electives unfulfilled', 'On track', '1\u20132 missing', '3\u20134', '\u22655'],
        ['P-05', 'Summer dependency', 'None', '1 course', '2', '\u22653'],
    ], L)

    doc.add_page_break()

    # ── SCORING METHODOLOGY ──
    add_heading(doc, '5.C Risk Scoring Methodology', L, level=1)

    add_heading(doc, 'Hybrid Scoring Approach', L, level=2)
    add_para(doc, 'QMentor uses a hybrid model combining weighted scoring with critical-indicator override:', L, bold=True, size=11)

    add_heading(doc, 'Layer 1 \u2014 Weighted Score (0\u2013100)', L, level=3)
    add_table(doc, ['Category', 'Weight', 'Rationale'], [
        ['G: Grades', '30%', 'Strongest predictor of academic outcomes'],
        ['A: Attendance', '25%', 'Direct correlation with barring and failure'],
        ['S: Assignments', '15%', 'Leading indicator before grades finalize'],
        ['AC: Academic Standing', '15%', 'Cumulative risk from institutional records'],
        ['E: LMS Engagement', '10%', 'Behavioral predictor of disengagement'],
        ['R: Registration', '5%', 'Supplementary signal'],
    ], L)

    add_heading(doc, 'Layer 2 \u2014 Critical Override', L, level=3)
    add_para(doc, 'Any single indicator at Critical in these categories auto-escalates the student:', L)
    add_table(doc, ['Indicator', 'Condition', 'Auto-Escalates To'], [
        ['A-03', 'Approaching barring limit (\u22642%)', 'Minimum High'],
        ['AC-01', '2nd+ academic warning', 'Minimum High'],
        ['G-09', '3+ courses failing', 'Minimum Critical'],
        ['C-01', 'Silent withdrawal detected', 'Minimum Critical'],
        ['C-05', 'Multi-course disengagement', 'Minimum Critical'],
    ], L)

    add_heading(doc, 'Layer 3 \u2014 Auto-Escalation (No Improvement)', L, level=3)
    add_table(doc, ['Current Level', 'If No Improvement After', 'Escalates To'], [
        ['Medium', '2 weeks', 'High'],
        ['High', '1 week', 'Critical'],
        ['Critical', 'Immediate', 'Dept Head notified; Dean if no response in 48hrs'],
    ], L)

    add_heading(doc, 'Risk Level Thresholds', L, level=3)
    add_table(doc, ['Level', 'Score', 'Color', 'Arabic'], [
        ['Low', '0\u201325', 'Green', '\u0623\u062e\u0636\u0631'],
        ['Medium', '26\u201350', 'Yellow', '\u0623\u0635\u0641\u0631'],
        ['High', '51\u201375', 'Orange', '\u0628\u0631\u062a\u0642\u0627\u0644\u064a'],
        ['Critical', '76\u2013100', 'Red', '\u0623\u062d\u0645\u0631'],
    ], L)

    add_heading(doc, 'College-Specific Threshold Overrides', L, level=3)
    add_table(doc, ['College', 'Customization', 'Rationale'], [
        ['Medicine', 'A-01 thresholds reduced by 5% (Medium at 6% not 11%)', 'Clinical patient safety'],
        ['Engineering / CS', 'E-01\u2013E-06 stricter; lab deadlines zero tolerance', 'Lab-based, LMS-heavy curriculum'],
        ['Sharia / Arabic', 'E-01\u2013E-06 relaxed; S-01 less critical', 'Less LMS dependency; oral exams'],
        ['Pharmacy / Dentistry', 'Same as Medicine for attendance + co-op weight', 'Clinical training component'],
        ['Business', 'Standard thresholds', 'Balanced curriculum'],
    ], L)

    doc.add_page_break()

    # ── AUTONOMY MATRIX ──
    add_heading(doc, '6. Task Autonomy Matrix \u2014 Agent vs. Human', L, level=1)

    add_blue_box(doc,
        'This matrix defines every system action and whether it is performed autonomously by QMentor '
        'or requires human involvement. This is the authoritative reference for the development team.', L)

    add_para(doc, '\U0001f916 = Agent (Fully Autonomous)     \U0001f916\U0001f464 = Agent + Notify Human     \U0001f464\U0001f916 = Human Approves     \U0001f464 = Human Only', L, bold=True, size=10)

    matrix_sections = [
        ('Data Collection & Monitoring', [
            ['1', 'Pull attendance from Blackboard (2hr)', '\U0001f916 Agent', 'QMentor'],
            ['2', 'Pull grades from SIS (6hr)', '\U0001f916 Agent', 'QMentor'],
            ['3', 'Pull assignments from Blackboard (4hr)', '\U0001f916 Agent', 'QMentor'],
            ['4', 'Pull registration/standing daily', '\U0001f916 Agent', 'QMentor'],
            ['5', 'Track LMS engagement', '\U0001f916 Agent', 'QMentor'],
            ['6', 'Detect anomalies', '\U0001f916 Agent', 'QMentor'],
            ['7', 'Build/update Digital Twin', '\U0001f916 Agent', 'QMentor'],
        ]),
        ('Risk Assessment', [
            ['8', 'Calculate risk score (66 criteria)', '\U0001f916 Agent', 'QMentor'],
            ['9', 'Classify risk level', '\U0001f916 Agent', 'QMentor'],
            ['10', 'Predict course failure', '\U0001f916 Agent', 'QMentor'],
            ['11', 'Predict absence limit', '\U0001f916 Agent', 'QMentor'],
            ['12', 'Detect declining trends', '\U0001f916 Agent', 'QMentor'],
            ['13', 'Generate risk explanation', '\U0001f916 Agent', 'QMentor'],
            ['14', 'Manually adjust risk level', '\U0001f464 Human', 'Advisor'],
        ]),
        ('Alerts \u2014 Low/Medium Risk', [
            ['15', 'In-app nudge (missed lecture)', '\U0001f916 Agent', 'QMentor'],
            ['16', 'Weekly summary email', '\U0001f916 Agent', 'QMentor'],
            ['17', 'Suggest learning content', '\U0001f916 Agent', 'QMentor'],
            ['18', 'Attendance warning (near limit)', '\U0001f916\U0001f464 Agent+Notify', 'QMentor (advisor notified)'],
            ['19', 'Suggest fellow student help matches (optional)', '\U0001f916 Agent', 'QMentor (suggestion only)'],
        ]),
        ('Alerts \u2014 High Risk', [
            ['20', 'SMS urgent alert to student', '\U0001f916\U0001f464 Agent+Notify', 'QMentor (advisor notified)'],
            ['21', 'Digital Twin report to advisor', '\U0001f916 Agent', 'QMentor'],
            ['22', 'Generate AI intervention plan', '\U0001f916 Agent', 'QMentor (recommendation)'],
            ['23', 'Advisor reviews/approves plan', '\U0001f464\U0001f916 Human Approves', 'Advisor'],
            ['24', 'Send intervention message', '\U0001f464\U0001f916 Human Approves', 'Agent after approval'],
        ]),
        ('Alerts \u2014 Critical Risk', [
            ['25', 'Generate emergency recovery plan', '\U0001f916 Agent', 'QMentor (recommendation)'],
            ['26', 'Advisor reviews recovery plan', '\U0001f464\U0001f916 Human Approves', 'Advisor'],
            ['27', 'Escalate to Department Head', '\U0001f464\U0001f916 Human Approves', 'Advisor approves'],
            ['28', 'Escalate to Dean', '\U0001f464 Human', 'Dept Head initiates'],
            ['29', 'Recommend probation/dismissal', '\U0001f464 Human', 'University committee'],
            ['30', 'Issue formal academic warning', '\U0001f464 Human', 'Registrar only'],
        ]),
        ('Chatbot Interactions', [
            ['31', 'Answer policy questions (bylaws)', '\U0001f916 Agent', 'Chatbot'],
            ['32', 'Personalized answers (policy + Twin)', '\U0001f916 Agent', 'Chatbot'],
            ['33', 'Schedule Teams meeting', '\U0001f916 Agent', 'Via Graph API'],
            ['34', 'Send Teams calendar invitations', '\U0001f916 Agent', 'Via Graph API'],
            ['35', 'Suggest drop + graduation impact', '\U0001f916 Agent', 'Recommendation only'],
            ['36', 'Actually drop a course in SIS', '\U0001f464 Human', '\u274c NEVER \u2014 student does it'],
            ['37', 'Detect emotional distress', '\U0001f916 Agent', 'Auto-flagged'],
            ['38', 'Mental health counseling', '\U0001f464 Human', '\u274c NEVER \u2014 routes to center'],
            ['39', 'Escalate complex question', '\U0001f916 Agent', 'Auto-escalated'],
        ]),
        ('Study Plan & Graduation', [
            ['40', 'Calculate remaining semesters', '\U0001f916 Agent', 'QMentor'],
            ['41', 'Generate semester course plan', '\U0001f916 Agent', 'Recommendation'],
            ['42', 'Build elective/free course recommendations', '\U0001f916 Agent', 'Based on full history analysis'],
            ['43', 'Generate visual semester timetable', '\U0001f916 Agent', '2\u20133 options ranked'],
            ['44', 'What-if simulations', '\U0001f916 Agent', 'QMentor'],
            ['45', 'Warn if choices delay graduation', '\U0001f916 Agent', 'QMentor'],
            ['46', 'Approve actual course registration', '\U0001f464 Human', 'Advisor in SIS'],
            ['47', 'Approve major change', '\U0001f464 Human', 'Dept + Dean committee'],
        ]),
        ('Administration', [
            ['51', 'Generate analytics reports', '\U0001f916 Agent', 'QMentor'],
            ['54', 'Configure college thresholds', '\U0001f464 Human', 'Dept Head + Admin'],
            ['55', 'Kill switch (disable automation)', '\U0001f464 Human', 'System Admin'],
            ['56', 'Retrain ML models', '\U0001f464\U0001f916 Human Approves', 'AI Team approves'],
            ['57', 'Approve college-specific threshold changes', '\U0001f464 Human', 'Dept Head + governance'],
        ]),
    ]

    for title, rows in matrix_sections:
        add_heading(doc, title, L, level=3)
        add_table(doc, ['#', 'Task', 'Mode', 'Who Acts'], rows, L)

    add_heading(doc, 'Summary', L, level=3)
    add_table(doc, ['Mode', 'Count', '%'], [
        ['\U0001f916 Agent (Fully Autonomous)', '32', '53%'],
        ['\U0001f916\U0001f464 Agent + Notify', '4', '7%'],
        ['\U0001f464\U0001f916 Human Approves', '8', '13%'],
        ['\U0001f464 Human Only', '16', '27%'],
        ['Total', '60', '100%'],
    ], L)

    add_blue_box(doc,
        'KEY PRINCIPLE: QMentor automates 60% of tasks but NEVER performs irreversible academic actions '
        '(course drops, grade changes, sanctions, dismissals). These remain exclusively human decisions. '
        'The agent ensures the right information reaches the right person at the right time.', L)

    doc.add_page_break()

    # ── SDAIA ETHICS ──
    add_heading(doc, '7. SDAIA AI Ethics Compliance', L, level=1)
    add_blue_box(doc, '6 principles, 37 verifiable requirements. Compliant with SDAIA AI Ethics Framework (2023).', L)

    principles = [
        ('7.1 Fairness & Non-Discrimination', [
            'Risk models audited for bias (gender, college, nationality) before deployment + after each retraining',
            'Nationality/gender/socioeconomic status NEVER used as direct model features',
            'Quarterly Fairness Reports; >5% disparity triggers retraining within 2 weeks',
            'Alert messages culturally neutral \u2014 no assumptions about family/finances',
        ]),
        ('7.2 Transparency & Explainability', [
            'Every risk score includes top 3 contributing factors with weights',
            '"My Data" page: student sees all collected data',
            'All AI content labeled "Generated by AI"',
            'Chatbot discloses AI nature + offers human escalation',
            'Advisor sees confidence level per prediction (e.g., "87%")',
        ]),
        ('7.3 Accountability & Governance', [
            'AI Governance Committee (deans + IT + legal + student rep) reviews quarterly',
            'Every automated action logged: timestamp, reasoning, data, outcome',
            'Model Registry: versions, training data, metrics, approval status',
            'Student grievance mechanism: "Challenge Assessment" button',
            'Annual external audit by SDAIA-recognized auditor',
        ]),
        ('7.4 Privacy & Data Protection', [
            'Purpose limitation \u2014 only academic advising data collected',
            'Granular consent: separate opt-ins for monitoring, chatbot, WhatsApp',
            'Consent revocable anytime \u2014 graceful degradation to manual advising',
            'All data in Saudi Arabia \u2014 no cross-border transfer',
            'Privacy Impact Assessment (DPIA) before each module launch',
        ]),
        ('7.5 Safety & Reliability', [
            'Kill switch: admin disables all automation in <5 minutes',
            'Mental health: AI routes only, NEVER diagnoses',
            'NEVER performs irreversible actions (no drops, no grade changes, no sanctions)',
            'Critical false positive rate <10%',
            'Canary deployments: shadow mode 2 weeks before production',
        ]),
        ('7.6 Human Oversight & Control', [
            'Advisor overrides any AI recommendation (with documented reason)',
            'Manual risk level adjustment (up/down) with justification',
            'AI uses "Suggested" / "Recommended" \u2014 never "Must" / "Required"',
            'Tiered: Low=auto, Medium=auto+notify, High=advisor approves, Critical=advisor+dept head',
            'Override rate >30% triggers automatic model review',
        ]),
    ]

    for title, items in principles:
        add_heading(doc, title, L, level=2)
        for item in items:
            check(doc, item, L)

    doc.add_page_break()

    # ── AGENTIC AI ──
    add_heading(doc, '8. Agentic AI Guardrails', L, level=1)

    add_heading(doc, '5 Autonomy Levels', L, level=2)
    add_table(doc, ['Level', 'Description', 'Application', 'Approval?'], [
        ['L0: Informational', 'Collect & display', 'Digital Twin, dashboards', 'No'],
        ['L1: Advisory', 'Recommend, human decides', 'Recovery plans, grad path', 'No (rec. only)'],
        ['L2: Supervised Auto', 'Act + notify', 'Low/Med alerts', 'No, advisor notified'],
        ['L3: Conditional Auto', 'Act after pre-approval', 'High risk escalation', 'Yes'],
        ['L4: Human-Only', 'Data only, human decides', 'Critical, mental health', 'Mandatory'],
    ], L)

    add_heading(doc, '15 Safety Guardrails', L, level=2)
    guardrails = [
        ('AGT-05: No Irreversible Actions', 'NEVER drops courses, changes grades, or issues sanctions. Recommends only.'),
        ('AGT-06: Rate Limiting', 'Max 1 alert/student/channel/day (except Critical).'),
        ('AGT-07: No Hallucination', 'References only data in Digital Twin. No fabricated grades/attendance.'),
        ('AGT-08: Scope Containment', 'Academic advising ONLY. No legal, medical, financial, or personal advice.'),
        ('AGT-09: Temporal Awareness', 'Considers semester phase (add/drop, midterm, finals) in thresholds.'),
        ('AGT-10: No Conflict of Interest', 'Risk ignores advisor workload. Based solely on student indicators.'),
        ('AGT-11: Feedback Loop Integrity', 'Own alerts NOT used as model input (prevents self-reinforcing bias).'),
        ('AGT-12: Graceful Degradation', 'Missing data source \u2192 continues + flags reduced confidence.'),
        ('AGT-13: Student Agency', 'Options, not directives. Student can dismiss/snooze non-critical alerts.'),
        ('AGT-14: Cultural Sensitivity', 'Arabic reviewed by native speakers. No patronizing or shaming language.'),
        ('AGT-15: Adversarial Robustness', 'Tested against students gaming LMS to manipulate risk scores.'),
    ]
    for title, desc in guardrails:
        add_para(doc, title, L, bold=True, size=10, spacing_after=1)
        add_para(doc, desc, L, size=9, color=GRAY_MED, spacing_after=6)

    doc.add_page_break()

    # ── SECURITY ──
    add_heading(doc, '9. Security & Data Classification', L, level=1)

    add_heading(doc, '4-Tier Data Classification', L, level=2)
    add_table(doc, ['Tier', 'Examples', 'Handling'], [
        ['T1: Public', 'University name, course catalog', 'No extra encryption; CDN cacheable'],
        ['T2: Internal', 'Aggregate stats, anonymized analytics', 'Encrypted at rest; authenticated access'],
        ['T3: Confidential', 'GPA, grades, attendance, risk scores', 'Field-level KMS encryption; RBAC; audit logged'],
        ['T4: Restricted', 'National ID, phone, medical, well-being flags', 'Column-level encryption (separate KMS); MFA + justification; alert on access'],
    ], L)

    add_heading(doc, 'Sensitive Data Scenarios', L, level=2)
    add_table(doc, ['Scenario', 'Tier', 'Protocol'], [
        ['Chatbot conversation', 'T3', 'Encrypted; 2yr retention; student can delete'],
        ['Mental distress detected', 'T4', 'Separate KMS key; counseling only; not visible to advisor w/o consent'],
        ['Advisor views Digital Twin', 'T3', 'Access logged (who, when, which student)'],
        ['Dept Head views college report', 'T3', 'Role-based access; logged; aggregated data only'],
        ['ML model training', 'T2', 'k-anonymized (k\u22655); differential privacy (\u03b5\u22641.0); isolated bucket'],
        ['Nafath login (National ID)', 'T4', 'NEVER stored in QMentor; passthrough only'],
    ], L)

    doc.add_page_break()

    # ── AWS ──
    add_heading(doc, '11. AWS Infrastructure', L, level=1)
    add_table(doc, ['Layer', 'Service', 'Purpose'], [
        ['Compute', 'ECS Fargate', 'Laravel backend'],
        ['GPU', 'EC2 G5.xlarge', 'LLM inference (NVIDIA A10G)'],
        ['GPU', 'SageMaker', 'ML model serving'],
        ['Serverless', 'Lambda', 'Event processing'],
        ['LLM', 'Bedrock', 'Claude/Llama API'],
        ['NLP', 'Comprehend', 'Arabic sentiment'],
        ['Database', 'RDS PostgreSQL', 'Primary data (Multi-AZ)'],
        ['Cache', 'ElastiCache Redis', 'Cache + rate limiting'],
        ['Vector', 'ChromaDB on ECS', 'Semantic search'],
        ['Queue', 'SQS + EventBridge', 'Async + scheduling'],
        ['Auth', 'Cognito + SAML', 'SSO via MyQU'],
        ['Security', 'WAF + Shield + KMS', 'Protection + encryption'],
    ], L)

    add_heading(doc, 'Monthly Cost Estimate', L, level=2)
    add_table(doc, ['Service', 'Monthly (USD)'], [
        ['ECS Fargate', '$400\u2013600'],
        ['EC2 G5 GPU (1 base + autoscale)', '$800\u20132,000'],
        ['SageMaker', '$500\u20131,000'],
        ['Bedrock API', '$300\u2013800'],
        ['RDS PostgreSQL', '$400\u2013600'],
        ['ElastiCache + S3 + CDN', '$350\u2013600'],
        ['SES + SNS + Other', '$400\u2013800'],
        ['TOTAL', '$3,150\u20136,400/mo'],
    ], L)

    doc.add_page_break()

    # ── PHASES ──
    add_heading(doc, '12. Project Phases & Timeline', L, level=1)
    phases = [
        ('Phase 1: Foundation (Months 1\u20133) \u2014 Q2 2026', [
            'AWS infrastructure (VPC, RDS, ECS, Cognito, WAF, KMS)',
            'Data pipeline: Blackboard + Oracle SIS integration',
            'Digital Twin v1 (GPA, courses, attendance)',
            'Alert system v1 (email + in-app for high/critical)',
            'Advisor dashboard v1 + MyQU SSO',
        ]),
        ('Phase 2: Intelligence (Months 4\u20136) \u2014 Q3 2026', [
            'Predictive analytics v1 (risk classification model)',
            'Chatbot v1: Arabic + English + policy RAG + Teams scheduling',
            'AI Study Plan Builder v1 (core path + elective recommendations)',
            'GPU infrastructure (G5 + SageMaker)',
            'Mobile app beta',
        ]),
        ('Phase 3: Expansion (Months 7\u20139) \u2014 Q4 2026', [
            'Graduation optimizer (multi-path + what-if)',
            'Semester Schedule Builder (visual timetable + options)',
            'Student-to-student academic help (voluntary) + WhatsApp + Faculty dashboard',
            'Sentiment analysis + well-being flags',
        ]),
        ('Phase 4: Optimization (Months 10\u201312) \u2014 Q1 2027', [
            'Advanced study plan builder v2 + elective optimizer',
            'Predictive analytics v2 (refined models)',
            'Leadership dashboard + Security audit (NCA)',
        ]),
        ('Phase 5: Maturity (Months 13\u201318) \u2014 Q2\u2013Q3 2027', [
            'Full university rollout (all colleges)',
            'NCAAA/ABET reports + code.gov.sa deposit',
            'Award submissions (DGA, SDAIA Summit, WSIS)',
        ]),
    ]
    for title, items in phases:
        add_heading(doc, title, L, level=2)
        for item in items:
            bullet(doc, item, L)

    # ── FOOTER ──
    doc.add_paragraph('')
    add_section_divider(doc, L)
    add_para(doc, 'Prepared by: Digital Transformation Team \u2014 Deanship of E-Learning & IT, Qassim University',
             L, size=9, color=GRAY_MED, align=L.center)
    add_para(doc, 'Next revision: Upon Phase 1 completion (end Q2 2026)',
             L, size=9, color=GRAY_MED, align=L.center)

    return doc


# ═══════════════════════════════════════════════════════════════
#                         MAIN
# ═══════════════════════════════════════════════════════════════
def main():
    base = os.path.dirname(__file__)

    # English
    doc_en = build_english()
    en_path = os.path.join(base, 'QMentor_SRS_EN.docx')
    doc_en.save(en_path)
    print(f'English: {en_path}')

if __name__ == '__main__':
    main()

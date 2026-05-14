#!/usr/bin/env python3
"""
Generate 66 Risk Indicators DOCX — maps each indicator to:
  1. How to get it (which API/system/endpoint)
  2. Whether it exists in the Blackboard/Oracle official API docs
  3. Whether it's implemented in THIS project's codebase
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BLUE = RGBColor(31, 78, 121)
GREEN = RGBColor(39, 124, 72)
GRAY = RGBColor(100, 100, 100)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(200, 120, 0)
WHITE = RGBColor(255, 255, 255)


def shade(cell, h):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}" w:val="clear"/>'))


def add_p(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def add_indicator_table(doc, headers, rows):
    """7-column table with color-coded status cells."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Wider table
    t.autofit = True

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.name = 'Calibri'
        r.font.color.rgb = WHITE
        shade(c, "1F4E79")

    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(txt))
            r.font.size = Pt(7.5)
            r.font.name = 'Calibri'

            # Color-code status columns (5 and 6)
            if ci in (5, 6):
                txt_lower = str(txt).lower()
                if '\u2705' in txt or 'yes' in txt_lower:
                    shade(c, "D5F5E3")  # green
                elif '\u26a0' in txt or 'partial' in txt_lower:
                    shade(c, "FEF9E7")  # yellow
                elif '\u274c' in txt or 'no' in txt_lower or 'missing' in txt_lower:
                    shade(c, "FADBD8")  # red
            elif ri % 2 == 0:
                shade(c, "EBF5FB")

    doc.add_paragraph('')
    return t


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.page_width = Cm(29.7)  # Landscape
        sec.page_height = Cm(21)
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(1.5)
        hdr = sec.header.paragraphs[0]
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr.add_run('QMentor V2 \u2014 66 Risk Indicators: API Mapping & Implementation Status')
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        ftr = sec.footer.paragraphs[0]
        ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ftr.add_run('66 Indicators Matrix  |  ')
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        ftr._p.append(parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="787878"/></w:rPr>'
            f'<w:t>1</w:t></w:r></w:fldSimple>'))

    # ═══ COVER ═══
    for _ in range(2):
        doc.add_paragraph('')
    add_p(doc, 'QMentor V2', bold=True, size=36, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_p(doc, '66 Risk Indicators \u2014 API Mapping & Implementation Status', bold=True, size=18, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_p(doc, '9 Categories | Per-Indicator Data Source | Official API Availability | Project Implementation Status',
          size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    add_p(doc, 'Legend:', bold=True, size=12, after=4)
    add_p(doc, '\u2705 = Available and implemented  |  \u26a0\ufe0f = Partially available (needs processing/inference)  |  \u274c = Not available (needs new endpoint)', size=10, after=2)
    add_p(doc, 'Column "API Exists?" = Does the endpoint exist in the official Blackboard/Oracle API documentation (regardless of our project)', size=10, after=2)
    add_p(doc, 'Column "In Project?" = Is it currently called/used in our Smart Advisor V1 codebase', size=10, after=8)

    # Summary table
    doc.add_heading('Summary: 66 Indicators by Category', level=2)
    sum_headers = ['Category', 'Count', '\u2705 Available', '\u26a0\ufe0f Partial', '\u274c Missing']
    sum_rows = [
        ['A: Attendance', '8', '4', '3', '1'],
        ['G: Grades', '11', '9', '1', '1'],
        ['S: Assignments', '6', '1', '3', '2'],
        ['E: LMS Engagement', '6', '0', '1', '5'],
        ['AC: Academic Standing', '7', '4', '1', '2'],
        ['R: Registration Behavior', '6', '1', '1', '4'],
        ['T: Schedule & Exams', '5', '3', '2', '0'],
        ['C: Compound Indicators', '12', '7', '5', '0'],
        ['P: Graduation Path', '5', '3', '2', '0'],
        ['TOTAL', '66', '32 (48%)', '19 (29%)', '15 (23%)'],
    ]
    t = doc.add_table(rows=11, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(sum_headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        shade(c, "1F4E79")
    for ri, row in enumerate(sum_rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.bold = (ri == len(sum_rows) - 1)
            if ci == 2: shade(c, "D5F5E3")
            elif ci == 3: shade(c, "FEF9E7")
            elif ci == 4: shade(c, "FADBD8")
            elif ri % 2 == 0: shade(c, "EBF5FB")

    doc.add_page_break()

    # ═══ Column definitions ═══
    H = ['ID', 'Indicator', 'How to Get (Data Source & Endpoint)', 'Exact Field(s) / Logic', 'API Exists?\n(Official Docs)', 'In Project?\n(Codebase)', 'Notes']

    # ═══ A: ATTENDANCE ═══
    doc.add_heading('Category A: Attendance \u2014 8 Indicators', level=1)
    add_p(doc, 'Primary Source: UniversityApiClient::getAbsences() \u2192 GET /absences-with-details', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['A-01', 'Absence % per course', 'Oracle SIS: GET /absences-with-details', 'absence_all_percent (per course object)', '\u2705 Yes \u2014 Oracle SIS API returns this directly', '\u2705 Yes \u2014 getAbsences() in UniversityApiClient, formatted in SmartAdvisorService::formatAbsences()', ''],
        ['A-02', 'Consecutive absences', 'Oracle SIS: GET /absences-with-details', 'Parse absences[] array, sort by absence_date, count consecutive dates', '\u26a0\ufe0f Partial \u2014 Raw dates available, need sequential counting logic', '\u26a0\ufe0f Partial \u2014 Raw absence dates pulled but no consecutive counting implemented', 'Need: sort absences by date, detect runs of consecutive class days'],
        ['A-03', 'Distance from barring limit', 'Oracle SIS: GET /absences-with-details', 'Calculate: 25% - absence_all_percent (barring = 25% in most colleges)', '\u2705 Yes \u2014 Can derive from existing data', '\u2705 Yes \u2014 formatAbsences() already calculates remaining allowed absences', 'Threshold varies by college (configurable)'],
        ['A-04', 'Absence after academic warning', 'Oracle SIS: GET /absences-with-details + NEW: /student/academic-warnings', 'Cross-reference: warning_date < absence_date', '\u26a0\ufe0f Partial \u2014 Absences yes, but academic warnings endpoint does NOT exist in current SIS API', '\u274c No \u2014 No warnings endpoint implemented', 'Need: new Oracle SIS endpoint for academic warning history'],
        ['A-05', 'Absence pattern (specific days)', 'Oracle SIS: GET /absences-with-details', 'Group absences by day-of-week from absence_date, detect recurring patterns', '\u2705 Yes \u2014 Date data available for analysis', '\u26a0\ufe0f Partial \u2014 Dates available but no day-of-week pattern analysis built', 'Need: groupBy dayOfWeek logic on absence_date array'],
        ['A-06', 'Medical vs unexcused ratio', 'Oracle SIS: GET /absences-with-details', 'absence_excused field ("0" or "1") per absence record, compare absence_excused_percent vs absence_all_percent', '\u2705 Yes \u2014 Both fields returned by SIS', '\u2705 Yes \u2014 Both fields pulled in getAbsences()', ''],
        ['A-07', 'Absence trend (3-week window)', 'Oracle SIS: GET /absences-with-details', 'Group absence_date by week, compare week-over-week counts', '\u2705 Yes \u2014 Date data available for windowed analysis', '\u26a0\ufe0f Partial \u2014 Dates pulled but no time-window trend analysis built', 'Need: sliding 3-week window with week-over-week comparison'],
        ['A-08', 'Cross-course absence correlation', 'Oracle SIS: GET /absences-with-details', 'Count courses where absence_all_percent > threshold simultaneously', '\u2705 Yes \u2014 Multi-course data returned in single call', '\u2705 Yes \u2014 getAbsences() returns all courses; formatAbsences() processes all', 'Just need cross-course threshold check logic'],
    ])

    doc.add_page_break()

    # ═══ G: GRADES ═══
    doc.add_heading('Category G: Grades \u2014 11 Indicators', level=1)
    add_p(doc, 'Primary Sources: BlackboardApiClient::getCourseGrades() + UniversityApiClient::getStudentProfile() + getAcademicTransactions()', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['G-01', 'Midterm exam score', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Filter grade items by name containing "midterm"/"\u0646\u0635\u0641\u064a": score, possible', '\u2705 Yes \u2014 Blackboard REST API returns all grade columns', '\u2705 Yes \u2014 getCourseGrades() + getAllCourseGrades() implemented', 'Grade column name matching needed (midterm/quiz/etc.)'],
        ['G-02', 'Quiz average', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Filter by name containing "quiz"/"\u0643\u0648\u064a\u0632": avg(score/possible)', '\u2705 Yes \u2014 Individual quiz columns available', '\u2705 Yes \u2014 Same endpoint, need name-based filtering', ''],
        ['G-03', 'Assignment average', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Filter by name containing "assignment"/"\u0648\u0627\u062c\u0628": avg(score/possible)', '\u2705 Yes \u2014 Individual assignment columns available', '\u2705 Yes \u2014 Same endpoint, need name-based filtering', ''],
        ['G-04', 'Cumulative GPA', 'Oracle SIS: GET /api/v3/me', 'profile.academic.cumulative_gpa (fallback: last_recorded_gpa)', '\u2705 Yes \u2014 SIS returns GPA directly', '\u2705 Yes \u2014 getStudentProfile() pulls this, used in formatProfile()', ''],
        ['G-05', 'Semester GPA trend', 'Oracle SIS: GET /student-academic-transactions', 'Compare semester_gpa across semesters array (chronological)', '\u2705 Yes \u2014 Per-semester GPA in transactions', '\u2705 Yes \u2014 getAcademicTransactions() returns semester-level GPAs', 'Need: compare last 2\u20133 semesters'],
        ['G-06', 'Gap vs section average', 'Blackboard: GET /blackboard/courses/{id}/grades/statistics (NOT IMPLEMENTED)', 'student_score - section_mean / section_stddev', '\u26a0\ufe0f Partial \u2014 Blackboard Learn API has grade statistics endpoint but NOT in our proxy API', '\u274c No \u2014 No section average endpoint implemented', 'Need: new proxy endpoint for Blackboard grade statistics'],
        ['G-07', 'Performance trajectory', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Sort grade items by date/sequence, compute score/possible trend (rising/declining)', '\u2705 Yes \u2014 Multiple grade items per course available', '\u2705 Yes \u2014 Grade data pulled, need trend calculation', ''],
        ['G-08', 'Courses at risk of failure', 'Blackboard: GET /blackboard/courses/{id}/grades (all courses)', 'Count courses where weighted_score < 60% (passing threshold)', '\u2705 Yes \u2014 All course grades available', '\u2705 Yes \u2014 getAllCourseGrades() iterates all courses', 'Need: passing threshold config per college'],
        ['G-09', 'Failed courses (current sem)', 'Oracle SIS: GET /student-academic-transactions', 'Count courses in latest semester where letter_grade in ("F", "DN", "NP")', '\u2705 Yes \u2014 Transaction data includes letter_grade', '\u2705 Yes \u2014 getAcademicTransactions() implemented', ''],
        ['G-10', 'Historical failure count', 'Oracle SIS: GET /student-academic-transactions', 'Count all courses across all semesters where letter_grade = "F"', '\u2705 Yes \u2014 Full history available', '\u2705 Yes \u2014 All semesters returned', 'Just need count(F) across all semesters'],
        ['G-11', 'Predicted final grade (ML)', 'SageMaker ML model (QMentor V2)', 'Input: G-01 to G-10 features \u2192 ML model \u2192 predicted letter grade', '\u274c N/A \u2014 Requires custom ML model', '\u274c No \u2014 Not implemented; V2 feature', 'Phase 2: Train model on historical outcomes'],
    ])

    doc.add_page_break()

    # ═══ S: ASSIGNMENTS ═══
    doc.add_heading('Category S: Assignments & Submissions \u2014 6 Indicators', level=1)
    add_p(doc, 'Primary Source: BlackboardApiClient::getCourseAttachments() + getCourseGrades()', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['S-01', 'Missing assignment %', 'Blackboard: grades + assignments', 'Count grade items with score=0 or null / total assignment-type items', '\u26a0\ufe0f Partial \u2014 Grades show scores but no explicit "submitted" flag', '\u26a0\ufe0f Partial \u2014 Can infer from zero scores in getCourseGrades()', 'Blackboard Learn API has /attempts endpoint but not proxied'],
        ['S-02', 'Late submission count', 'Blackboard: GET /courses/{id}/gradebook/columns/{id}/attempts', 'attempt.created vs. column.dueDate', '\u2705 Yes \u2014 Blackboard Learn REST API has attempts endpoint with timestamps', '\u274c No \u2014 Attempts endpoint not implemented in project', 'Need: proxy /attempts endpoint; compare created vs dueDate'],
        ['S-03', 'Last-minute submissions (<1hr)', 'Blackboard: GET /courses/{id}/gradebook/columns/{id}/attempts', 'attempt.created within 1hr of column.dueDate', '\u2705 Yes \u2014 Same attempts endpoint', '\u274c No \u2014 Not implemented', 'Same endpoint as S-02'],
        ['S-04', 'Assignment score trend', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Sort assignment-type grades chronologically, compute trend', '\u2705 Yes \u2014 Score data available', '\u26a0\ufe0f Partial \u2014 Scores pulled, no trend logic', 'Need: chronological sort + slope calculation'],
        ['S-05', 'Zero-score submissions', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Count grade items where score = 0 AND has submission', '\u2705 Yes \u2014 Score=0 visible in grades', '\u26a0\ufe0f Partial \u2014 Can detect score=0 from grades data', 'Ambiguous: 0 could mean "not submitted" or "submitted bad work"'],
        ['S-06', 'Missing high-weight assignments', 'Blackboard: GET /blackboard/courses/{id}/grades', 'Filter by pointsPossible > 10% of total, check score=0/null', '\u2705 Yes \u2014 pointsPossible field available', '\u26a0\ufe0f Partial \u2014 pointsPossible pulled, need weight calculation', 'Need: calculate weight as pointsPossible / sum(all pointsPossible)'],
    ])

    doc.add_page_break()

    # ═══ E: LMS ENGAGEMENT ═══
    doc.add_heading('Category E: LMS Engagement \u2014 6 Indicators', level=1)
    add_p(doc, '\u26a0\ufe0f CRITICAL GAP: Most LMS engagement data requires Blackboard Data (Analytics) or Building Blocks \u2014 NOT available via standard REST API', size=10, bold=True, color=RED)
    add_indicator_table(doc, H, [
        ['E-01', 'LMS login frequency/week', 'Blackboard Data / Activity Logs', 'Count distinct login sessions per week per student', '\u26a0\ufe0f Partial \u2014 Blackboard Learn REST API does NOT expose login logs. Blackboard Data (BbData/Snowflake) does.', '\u274c No \u2014 No activity log endpoint', 'Option 1: Blackboard Data warehouse\nOption 2: Custom Building Block\nOption 3: Track via LTI launch events'],
        ['E-02', 'Time on LMS (weekly hrs)', 'Blackboard Data / Session Analytics', 'Sum session durations per week', '\u274c No \u2014 Not in REST API. Available only in Blackboard Data.', '\u274c No', 'Requires Blackboard Data (Snowflake) integration'],
        ['E-03', 'Content access rate (%)', 'Blackboard: GET /courses/{id}/contents + /courses/{id}/contents/{id}/children (review status)', 'Count accessed content items / total content items', '\u26a0\ufe0f Partial \u2014 Blackboard REST has content listing but individual access tracking requires reviewStatus or Blackboard Data', '\u26a0\ufe0f Partial \u2014 getCourseContent() gets content list but not per-student access status', 'Need: Blackboard Data content_access table or REST reviewStatus if available'],
        ['E-04', 'Sudden engagement drop', 'Derived from E-01/E-02/E-03', 'Compare current week metrics vs. 3-week rolling average', '\u274c No \u2014 Depends on E-01/E-02 availability', '\u274c No', 'Blocked by E-01 and E-02 gaps'],
        ['E-05', 'Forum/discussion participation', 'Blackboard: GET /courses/{id}/discussions/threads', 'Count posts by student in discussion forums', '\u2705 Yes \u2014 Blackboard REST API has discussion thread endpoints', '\u274c No \u2014 Discussion endpoint not implemented in project', 'Need: proxy /discussions endpoint; count student posts'],
        ['E-06', 'Days since last LMS login', 'Blackboard Data / Session Analytics', 'TODAY - MAX(last_login_date)', '\u274c No \u2014 Not in REST API', '\u274c No', 'Same dependency as E-01'],
    ])

    doc.add_page_break()

    # ═══ AC: ACADEMIC STANDING ═══
    doc.add_heading('Category AC: Academic Standing \u2014 7 Indicators', level=1)
    add_p(doc, 'Primary Source: UniversityApiClient::getStudentProfile() + getAcademicTransactions()', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['AC-01', 'Active academic warnings', 'Oracle SIS: NEW endpoint needed: /student/academic-warnings', 'warning_type, warning_date, status', '\u274c No \u2014 Not in current SIS API. May exist in Oracle Student module but not exposed.', '\u274c No \u2014 No warnings endpoint', 'Need: Oracle DBA to expose academic warning data via API'],
        ['AC-02', 'Cumulative failed courses', 'Oracle SIS: GET /student-academic-transactions', 'Count all courses where letter_grade in ("F", "DN") across all semesters', '\u2705 Yes', '\u2705 Yes \u2014 getAcademicTransactions() returns full history', 'Simple count filter on existing data'],
        ['AC-03', 'Credit hours behind plan', 'Oracle SIS: GET /api/v3/me + /student-plan', 'total_plan_hours - total_earned_hours, compared to expected by current level', '\u2705 Yes \u2014 Both fields available', '\u2705 Yes \u2014 Both endpoints implemented', 'Need: "expected hours by semester" lookup table per major'],
        ['AC-04', 'Probation status', 'Oracle SIS: GET /api/v3/me', 'profile.academic.academic_status field', '\u26a0\ufe0f Partial \u2014 academic_status exists but may not distinguish probation types', '\u26a0\ufe0f Partial \u2014 Field pulled but not parsed for probation logic', 'Need: confirm possible values of academic_status field'],
        ['AC-05', 'GPA distance from min (2.0)', 'Oracle SIS: GET /api/v3/me', 'cumulative_gpa - 2.0', '\u2705 Yes', '\u2705 Yes', 'Simple calculation on existing data'],
        ['AC-06', 'Failed prereqs blocking progress', 'Oracle SIS: GET /student-plan + /student-academic-transactions', 'Find courses in plan with status="not_registered" where prereq has letter_grade="F"', '\u2705 Yes \u2014 Plan + grades both available', '\u26a0\ufe0f Partial \u2014 Both data sources pulled, need cross-reference logic', 'Need: prerequisite chain mapping + failure cross-check'],
        ['AC-07', 'Repeated course count', 'Oracle SIS: GET /student-academic-transactions', 'Count distinct course_codes appearing 2+ times across semesters', '\u2705 Yes', '\u2705 Yes', 'Simple groupBy + count > 1 on transactions'],
    ])

    doc.add_page_break()

    # ═══ R: REGISTRATION ═══
    doc.add_heading('Category R: Registration Behavior \u2014 6 Indicators', level=1)
    add_p(doc, 'Primary Source: Oracle SIS (mostly NEW endpoints needed)', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['R-01', 'Courses dropped this semester', 'Oracle SIS: NEW endpoint needed: /student/course-changes', 'Count drop transactions for current semester', '\u274c No \u2014 Not in current SIS API. Oracle may have registration_history table.', '\u274c No', 'Need: Oracle DBA to expose add/drop history'],
        ['R-02', 'Withdrawal (semester deferral) history', 'Oracle SIS: NEW endpoint needed: /student/deferrals', 'Count deferral records', '\u274c No \u2014 Not exposed', '\u274c No', 'Need: new SIS endpoint'],
        ['R-03', 'Major change requests', 'Oracle SIS: NEW endpoint needed: /student/major-changes', 'Count major change records', '\u274c No \u2014 Not exposed', '\u274c No', 'Need: new SIS endpoint'],
        ['R-04', 'Course load vs recommended', 'Oracle SIS: GET /api/v3/me', 'current_registered_hours compared to 15\u201318 normal range', '\u2705 Yes \u2014 current_registered_hours available', '\u2705 Yes \u2014 Pulled in getStudentProfile()', 'Simple range check on existing data'],
        ['R-05', 'Late registration pattern', 'Oracle SIS: NEW endpoint needed: /student/registration-dates', 'Compare registration_date to semester_start_date', '\u274c No \u2014 Registration timestamp not exposed', '\u274c No', 'Need: new SIS endpoint with registration timestamps'],
        ['R-06', 'Add/drop frequency (first 2 wks)', 'Oracle SIS: NEW /student/course-changes', 'Count add+drop transactions within first 2 weeks of semester', '\u26a0\ufe0f Partial \u2014 If course-changes endpoint existed, filtering by date would work', '\u274c No', 'Depends on R-01 endpoint'],
    ])

    doc.add_page_break()

    # ═══ T: SCHEDULE & EXAMS ═══
    doc.add_heading('Category T: Schedule & Exams \u2014 5 Indicators', level=1)
    add_p(doc, 'Primary Source: UniversityApiClient::getTimeTable() + getFinalExams()', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['T-01', 'Final exam same-day conflicts', 'Oracle SIS: GET /final-exams', 'Group exams by exam_date, count per day', '\u2705 Yes \u2014 exam_date returned', '\u2705 Yes \u2014 getFinalExams() implemented', 'Simple groupBy date + count'],
        ['T-02', 'Missed midterm exams', 'Blackboard grades + Oracle absences', 'Cross-ref: midterm grade=0 AND absence on midterm date', '\u26a0\ufe0f Partial \u2014 Both data sources exist, need correlation logic', '\u26a0\ufe0f Partial \u2014 Both pulled separately, no cross-reference', 'Need: date matching between absences and midterm grade items'],
        ['T-03', 'Missed final (prev semester)', 'Oracle SIS: GET /student-academic-transactions', 'letter_grade = "IC" (incomplete) or "AB" (absent) in previous semester', '\u2705 Yes \u2014 Letter grades include IC/AB codes', '\u2705 Yes', 'Check for IC/AB grade codes'],
        ['T-04', 'Schedule gap issues', 'Oracle SIS: GET /time-table', 'Analyze times[] per day: max(end) - min(start) vs. actual class hours', '\u2705 Yes \u2014 start/end times available', '\u2705 Yes \u2014 getTimeTable() returns full schedule', 'Need: per-day gap analysis (idle hours between classes)'],
        ['T-05', 'Exam week absence', 'Oracle absences + calendar', 'Count absences where absence_date falls within exam period from academic_calendar', '\u26a0\ufe0f Partial \u2014 Both data available, need date range matching', '\u26a0\ufe0f Partial \u2014 Both pulled, no cross-ref logic', 'Need: match absence_dates against calendar exam period dates'],
    ])

    doc.add_page_break()

    # ═══ C: COMPOUND ═══
    doc.add_heading('Category C: Compound Indicators \u2014 12 Indicators', level=1)
    add_p(doc, 'These combine multiple simple indicators. "Available" means all source indicators are available.', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['C-01', 'Silent withdrawal', 'E-06 + S-01 + A-01', 'E-06>7d AND S-01>50% AND A-01>15%', '\u26a0\ufe0f Partial \u2014 E-06 is missing (LMS login)', '\u26a0\ufe0f Partial \u2014 A-01 and S-01 available; E-06 not', 'Blocked by E-06 gap; can use A-01+S-01 as proxy'],
        ['C-02', 'Absence + declining grades', 'A-01 + G-07', 'A-01>10% AND G-07="Declining"', '\u2705 Yes', '\u2705 Yes \u2014 Both data sources available', ''],
        ['C-03', 'High effort, low results', 'E-02 + G-01', 'E-02\u22655hrs AND G-01<50%', '\u26a0\ufe0f Partial \u2014 E-02 (time on LMS) missing', '\u26a0\ufe0f Partial', 'Can proxy with content access instead of time'],
        ['C-04', 'Collapse after success', 'G-05 (GPA trend)', 'G-05 dropped>0.5 AND previous GPA\u22653.5', '\u2705 Yes', '\u2705 Yes \u2014 GPA history in transactions', ''],
        ['C-05', 'Multi-course disengagement', 'E-04 (engagement drop)', 'E-04>50% in \u22653 courses', '\u26a0\ufe0f Partial \u2014 E-04 depends on E-01', '\u26a0\ufe0f Partial', 'Can proxy with multi-course absence spike'],
        ['C-06', 'Drop cascade', 'R-01 + G-08 + A-01', 'R-01\u22652 AND G-08\u22651 AND A-01 increasing', '\u26a0\ufe0f Partial \u2014 R-01 (drops) missing', '\u26a0\ufe0f Partial', 'Blocked by R-01; can use grades+absence'],
        ['C-07', 'Grade-attendance disconnect', 'A-01 + G-01', 'A-01<5% (good) BUT G-01<40%', '\u2705 Yes', '\u2705 Yes', 'Both available'],
        ['C-08', 'Procrastination spiral', 'S-03 + S-01', 'S-03>60% AND S-01 increasing', '\u26a0\ufe0f Partial \u2014 S-03 needs attempts endpoint', '\u26a0\ufe0f Partial', 'Blocked by S-03'],
        ['C-09', 'Social isolation signal', 'E-05 + advisor meeting data', 'E-05=None AND no advisor meeting 4 weeks', '\u274c Missing \u2014 E-05 not available, no meeting tracker', '\u274c No', 'Need both forum data and meeting tracking'],
        ['C-10', 'Registration confusion', 'R-06 + R-04 + AC-06', 'R-06\u22654 AND R-04\u2260Normal AND AC-06\u22651', '\u26a0\ufe0f Partial \u2014 R-06 needs course-changes', '\u26a0\ufe0f Partial', 'R-04 and AC-06 available; R-06 blocked'],
        ['C-11', 'Graduation path drift', 'AC-03 + R-01 + G-04', 'AC-03>6hrs AND R-01\u22651 AND G-04 declining', '\u26a0\ufe0f Partial \u2014 R-01 missing', '\u26a0\ufe0f Partial', 'AC-03 + G-04 available'],
        ['C-12', 'Recovery failure', 'Previous intervention + no change', 'QMentor intervention log + re-evaluation', '\u2705 Yes (internal) \u2014 Once QMentor tracking exists', '\u274c No \u2014 V2 feature', 'Need: intervention tracking in QMentor DB'],
    ])

    doc.add_page_break()

    # ═══ P: GRADUATION ═══
    doc.add_heading('Category P: Graduation Path \u2014 5 Indicators', level=1)
    add_p(doc, 'Primary Source: UniversityApiClient::getStudentPlan() + getStudentProfile() + getAcademicPlanByMajor()', size=9, color=GRAY)
    add_indicator_table(doc, H, [
        ['P-01', 'Semesters behind plan', 'Oracle SIS: GET /student-plan + /api/v3/me', 'Compare academic_level / earned_hours vs. plan level expectations', '\u2705 Yes \u2014 Both plan and profile available', '\u2705 Yes \u2014 Both implemented', 'Need: expected-hours-per-level lookup table'],
        ['P-02', 'Core courses not yet passed', 'Oracle SIS: GET /student-plan', 'Count courses in plan with status="not_registered" in upper levels', '\u2705 Yes \u2014 Plan has status field', '\u2705 Yes \u2014 getStudentPlan() implemented', ''],
        ['P-03', 'Approaching max enrollment', 'Oracle SIS: GET /api/v3/me + policy data', 'Compare current semester count vs. max allowed (usually 7 years = 14 semesters)', '\u26a0\ufe0f Partial \u2014 No explicit "semester count" or "admission date" in current API', '\u26a0\ufe0f Partial \u2014 Can estimate from academic_level but not exact', 'Need: admission_date or semester_count from SIS'],
        ['P-04', 'Required electives unfulfilled', 'Oracle SIS: GET /student-plan', 'Count elective-type courses with status="not_registered"', '\u2705 Yes \u2014 Plan shows all required courses + status', '\u2705 Yes', 'Need: classify courses as elective vs. core from plan structure'],
        ['P-05', 'Summer dependency', 'Oracle SIS: GET /student-plan + course availability data', 'Check if remaining courses are only offered in certain semesters', '\u26a0\ufe0f Partial \u2014 Plan exists but course availability schedule not in API', '\u26a0\ufe0f Partial', 'Need: course offering schedule data (which semester each course is offered)'],
    ])

    doc.add_page_break()

    # ═══ ACTION PLAN ═══
    doc.add_heading('Implementation Priority & Action Plan', level=1)

    add_p(doc, 'Based on the analysis above, here is the phased implementation plan:', bold=True, size=11)

    phase_h = ['Phase', 'Indicators', 'Count', 'Action Required', 'Effort']
    phase_rows = [
        ['Phase 1\nImmediate', 'A-01, A-03, A-06, A-08, G-01\u2013G-05, G-07\u2013G-10, AC-02, AC-03, AC-05, AC-07, R-04, T-01, T-03, T-04, C-02, C-04, C-07, P-01, P-02, P-04', '26', 'Use existing API data as-is. Build calculation/threshold logic only.', 'Low \u2014 2\u20133 weeks'],
        ['Phase 2\nProcessing', 'A-02, A-05, A-07, A-08, G-06 (partial), S-01, S-04\u2013S-06, AC-04, AC-06, T-02, T-05, C-01*, C-03*, C-05*, C-06*, C-08*, C-10*, C-11*, P-03, P-05', '19', 'Build inference/correlation logic on existing data. Date parsing, cross-referencing, trend analysis.', 'Medium \u2014 3\u20134 weeks'],
        ['Phase 3a\nNew BB Endpoints', 'S-02, S-03 (attempts), E-05 (forums), G-06 (stats)', '4', 'Proxy new Blackboard REST API endpoints: /attempts, /discussions, /grades/statistics', 'Medium \u2014 2 weeks'],
        ['Phase 3b\nNew SIS Endpoints', 'A-04, AC-01 (warnings), R-01 (drops), R-02, R-03, R-05, R-06', '7', 'Request Oracle DBA to expose: academic_warnings, course_changes, deferrals, major_changes', 'High \u2014 depends on Oracle DBA availability'],
        ['Phase 3c\nBlackboard Data', 'E-01, E-02, E-03, E-04, E-06', '5', 'Integrate Blackboard Data (Snowflake/BbData) for activity logs, login frequency, time-on-platform', 'High \u2014 requires Blackboard Data license + Snowflake setup'],
        ['Phase 4\nML Models', 'G-11 (predicted grade), C-12 (recovery failure)', '2', 'Train ML models on historical data via SageMaker', 'High \u2014 Phase 2 of QMentor project'],
    ]

    t = doc.add_table(rows=1 + len(phase_rows), cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(phase_h):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        shade(c, "1F4E79")
    for ri, row in enumerate(phase_rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(7.5)
            if ri % 2 == 0: shade(c, "EBF5FB")

    # FOOTER
    doc.add_paragraph('')
    add_p(doc, '\u2501' * 60, size=8, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Prepared by: Digital Transformation Team \u2014 Qassim University', size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    path = os.path.join(os.path.dirname(__file__), 'QMentor_66_Indicators_API_Map.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Generate QMentor SRS English DOCX.
Uses python-docx.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def add_para(doc, text, bold=False, size=12, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cells[col_idx].text = ''
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>')
                cells[col_idx]._tc.get_or_add_tcPr().append(shading)

    return table


def generate_docx():
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # ==================== COVER PAGE ====================
    for _ in range(4):
        doc.add_paragraph('')

    add_para(doc, 'Qassim University', bold=True, size=28,
             color=RGBColor(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Deanship of E-Learning & IT \u2014 Digital Transformation Team',
             size=14, color=RGBColor(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 40)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_para(doc, 'QMentor', bold=True, size=40,
             color=RGBColor(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'AI-Powered Smart Academic Advisor', bold=True, size=20,
             color=RGBColor(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph('')
    add_para(doc, 'System Requirements Specification (SRS)', bold=True, size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Version 2.1 \u2014 1447 H / 2026 CE',
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Per DGA Digital Transformation Standards V5.0 \u2014 Standard 5.24',
             size=11, color=RGBColor(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph('')
    tags = [
        'Compliant with SDAIA AI Ethics Principles',
        'Innovative Solution \u2014 Qiyas Digital Transformation 2026',
        'Year of Artificial Intelligence 2026',
    ]
    for tag in tags:
        add_para(doc, f'\u25c6 {tag}', size=10,
                 color=RGBColor(39, 124, 72), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ==================== DOCUMENT INFO ====================
    add_heading(doc, '1. Document Information', level=1)

    add_table(doc,
        ['Field', 'Value'],
        [
            ['Project Name', 'QMentor \u2014 AI-Powered Smart Academic Advisor'],
            ['Organization', 'Deanship of E-Learning & IT \u2014 Digital Transformation Team'],
            ['Core Technology', 'Agentic AI \u2014 Predictive Analytics \u2014 Digital Twin'],
            ['Target Users', 'All students (~70,000), Academic Advisors, University Leadership'],
            ['Hosting', 'Cloud-Native on AWS (Saudi Region)'],
            ['Compliance', 'SDAIA AI Ethics \u2014 PDPL \u2014 NCA ECC'],
            ['Document Version', '2.1'],
            ['Date', '2026-04-07'],
            ['Classification', 'Confidential'],
            ['Standard', 'IEEE 830-1998 (modified)'],
        ])

    doc.add_page_break()

    # ==================== TOC ====================
    add_heading(doc, 'Table of Contents', level=1)

    toc_items = [
        ('1', 'Document Information'),
        ('2', 'Introduction & Scope'),
        ('3', 'Current Advisor Workflow vs. QMentor'),
        ('4', 'Functional Requirements \u2014 Core Modules (6)'),
        ('5', 'Functional Requirements \u2014 Extended Modules (12)'),
        ('6', 'Task Autonomy Matrix \u2014 Agent vs. Human-in-the-Loop (58 tasks)'),
        ('7', 'SDAIA AI Ethics Compliance (6 Principles, 37 Requirements)'),
        ('8', 'Agentic AI Principles & Guardrails (15 Guardrails)'),
        ('9', 'Security & Sensitive Data Handling (4-Tier Classification)'),
        ('10', 'Non-Functional Requirements'),
        ('11', 'AWS Infrastructure & Cost Estimate'),
        ('12', 'Project Phases & Timeline'),
        ('A', 'Appendix: Risk Criteria Matrix (66 Indicators)'),
    ]
    for num, title in toc_items:
        add_para(doc, f'{num}.  {title}', size=12)

    doc.add_page_break()

    # ==================== SECTION 2: INTRO ====================
    add_heading(doc, '2. Introduction & Scope', level=1)

    add_heading(doc, '2.1 Purpose', level=2)
    add_para(doc,
        'This SRS defines the complete functional and non-functional requirements for QMentor, '
        'an Agentic AI-powered Smart Academic Advisor system for Qassim University. It serves as '
        'the single source of truth for the development team, stakeholders, and QA throughout the project lifecycle.')

    add_heading(doc, '2.2 Scope', level=2)
    add_para(doc,
        'QMentor is a cloud-native platform hosted on AWS GPU-based servers that provides:')

    scope_items = [
        'Proactive academic advising via an autonomous AI agent that monitors every student 24/7',
        'Student Digital Twin \u2014 a real-time holistic profile combining Blackboard LMS, Oracle SIS, and MyQU data',
        'Predictive analytics that detect academic risk before it escalates (66 indicators across 9 categories)',
        'Multi-channel smart alerts (in-app, email, SMS, WhatsApp) with automatic escalation',
        'Natural-language advising chatbot (Arabic + English) with human escalation',
        'Chatbot answers academic regulation questions (e.g., "What are the requirements for First Honors?")',
        'Microsoft Teams meeting scheduling directly from chatbot',
        'Advisor dashboard with AI-generated intervention plans',
        'Learning content recommendation engine',
        'Mobile application (iOS & Android)',
        'WhatsApp / Telegram bot integration',
        'Graduation path optimizer',
        'Career guidance AI module',
        'Student-to-Student Academic Help — matching high-performing students with struggling peers (fully voluntary, not a private tutor or university instructor)',
        'Student sentiment & well-being analysis',
        'Faculty & leadership analytics dashboards',
        'Gamification & engagement engine',
        'Parent/guardian portal (with student consent)',
        'Smart scheduling assistant with conflict detection',
        'Emergency academic recovery program',
        'Institutional benchmarking & accreditation reporting',
    ]
    for item in scope_items:
        add_para(doc, f'\u2022 {item}', size=11)

    add_heading(doc, '2.3 Target Users', level=2)
    add_table(doc,
        ['Role', 'Description'],
        [
            ['Student', 'All enrolled students (~70,000) across all colleges'],
            ['Academic Advisor', 'Faculty members assigned as academic advisors'],
            ['Course Instructor', 'Faculty teaching specific sections'],
            ['Department Head', 'Heads of academic departments'],
            ['Dean / Vice-Dean', 'College leadership'],
            ['University Leadership', 'Provost, VP Academic Affairs'],
            ['System Admin', 'IT administrators managing the platform'],
            ['Parent / Guardian', '(Optional) with student consent'],
        ])

    doc.add_page_break()

    # ==================== SECTION 3: CURRENT vs QMENTOR ====================
    add_heading(doc, '3. Current Advisor Workflow vs. QMentor', level=1)

    add_para(doc,
        'Based on the "Communication Guide for Academic Advising" (Deanship of Admission & Registration, 2020), '
        'the current advising system relies on a fully manual, multi-step process through the Oracle SIS portal. '
        'QMentor automates each step while preserving the advisor\'s authority.', size=11)

    add_heading(doc, 'Advisor \u2192 Student Communication (Current: 7 Manual Steps)', level=2)
    add_table(doc,
        ['Step', 'Current Manual Process', 'QMentor Automation'],
        [
            ['1', 'Advisor logs into Oracle SIS portal', 'Eliminated \u2014 system monitors continuously'],
            ['2', 'Navigates to "Academic" menu', 'Eliminated'],
            ['3', 'Selects "Academic Advisor" service', 'Eliminated \u2014 advisor sees pre-sorted dashboard'],
            ['4', 'Views student list (names only, no data)', 'Digital Twin: full profile with GPA, attendance, risk, recommendations'],
            ['5', 'Clicks "Communicate" next to student', 'Auto-triggered: alerts sent proactively when risk detected'],
            ['6', 'Chooses channel (portal/SMS-65 chars/email), writes content', 'AI-authored personalized messages; system selects optimal channel'],
            ['7', 'Clicks "Send"', 'Auto-sent for Low/Medium; advisor approves for High/Critical'],
        ])

    add_heading(doc, 'Student \u2192 Advisor Communication (Current: 6 Manual Steps)', level=2)
    add_table(doc,
        ['Step', 'Current Manual Process', 'QMentor Automation'],
        [
            ['1', 'Student logs into SIS \u2192 "Personal" menu', 'Opens QMentor app or web portal'],
            ['2', 'Selects "Send Email"', 'Types/speaks natural language in chatbot'],
            ['3', 'Sees list of instructors + advisor', 'Not needed \u2014 chatbot routes automatically'],
            ['4', 'Selects advisor checkbox', 'Not needed'],
            ['5', 'Writes subject + content', 'Speaks or types in Arabic/English naturally'],
            ['6', 'Clicks "Send", waits days for response', 'Instant AI response in seconds; escalates to human only for complex cases'],
        ])

    add_heading(doc, 'Key Pain Points Resolved', level=2)
    add_table(doc,
        ['Pain Point', 'Impact', 'QMentor Solution'],
        [
            ['SMS limited to 65 characters', 'Cannot convey meaningful advice', 'Multi-channel with unlimited content (app, email, WhatsApp)'],
            ['Portal messages appear as banner ads', 'Students ignore them', 'Push notifications + personalized chatbot follow-up'],
            ['No student performance data visible', 'Advisor sends generic messages', 'Full Digital Twin visible during every interaction'],
            ['Email-only for student (one-way, slow)', 'Response takes days', 'Chatbot instant response; human escalation when needed'],
            ['Instructor sees only student names', 'Cannot identify struggling students', 'Section risk heatmap with per-student metrics'],
            ['All communication is reactive', 'At-risk students missed entirely', '24/7 proactive monitoring with auto-escalation'],
        ])

    doc.add_page_break()

    # ==================== SECTION 4: CORE MODULES ====================
    add_heading(doc, '4. Functional Requirements \u2014 Core Modules', level=1)

    # 4.1 Monitoring Agent
    add_heading(doc, '4.1 Continuous Monitoring Agent (FR-MON)', level=2)
    add_para(doc, 'The core Agentic AI loop that runs autonomously 24/7.', size=11)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-MON-01', 'Pull attendance data from Blackboard LMS every 2 hours via REST API', 'Must'],
            ['FR-MON-02', 'Pull grade data from Oracle SIS every 6 hours via scheduled job', 'Must'],
            ['FR-MON-03', 'Pull assignment submission status from Blackboard every 4 hours', 'Must'],
            ['FR-MON-04', 'Pull registration and academic standing from Oracle SIS daily', 'Must'],
            ['FR-MON-05', 'Pull student requests (drop/add, withdrawal) from MyQU daily', 'Must'],
            ['FR-MON-06', 'Track LMS login frequency and content interaction per student', 'Should'],
            ['FR-MON-07', 'Detect data anomalies (sudden grade drops, attendance patterns)', 'Must'],
            ['FR-MON-08', 'Log all data pulls with timestamps and success/failure status', 'Must'],
            ['FR-MON-09', 'Retry failed data pulls up to 3 times with exponential backoff', 'Must'],
            ['FR-MON-10', 'Support configurable polling intervals per data source', 'Should'],
        ])

    # 4.2 Predictive Analytics
    add_heading(doc, '4.2 Predictive Analytics Engine (FR-PAE)', level=2)
    add_para(doc, 'ML models that analyze student data and predict academic risk.', size=11)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-PAE-01', 'Classify each student into 4 risk levels: Low, Medium, High, Critical', 'Must'],
            ['FR-PAE-02', 'Predict probability of course failure per student per course', 'Must'],
            ['FR-PAE-03', 'Predict probability of reaching absence limit per course', 'Must'],
            ['FR-PAE-04', 'Detect declining performance trends over 2+ weeks', 'Must'],
            ['FR-PAE-05', 'Factor in historical data (previous semesters, warnings)', 'Must'],
            ['FR-PAE-06', 'Generate risk scores on 0\u2013100 scale with explanation (top 3 factors)', 'Should'],
            ['FR-PAE-07', 'Retrain models at end of each semester with new outcome data', 'Must'],
            ['FR-PAE-09', 'Achieve minimum 85% accuracy in risk predictions', 'Must'],
            ['FR-PAE-10', 'Provide explainability \u2014 why student was flagged (contributing factors)', 'Must'],
        ])

    # 4.3 Digital Twin
    add_heading(doc, '4.3 Student Digital Twin (FR-SDT)', level=2)
    add_para(doc, 'A real-time, comprehensive digital profile for every student.', size=11)
    add_table(doc,
        ['Section', 'Contents'],
        [
            ['Academic Profile', 'Cumulative GPA, credits earned/remaining, warnings count, expected graduation grade'],
            ['Study Plan', 'Completed/remaining courses, prerequisites, suggested optimal path for remaining semesters'],
            ['Behavior Metrics', 'Attendance rate per course, assignment submission rate, LMS interaction hours'],
            ['Risk Points', 'At-risk courses, active alerts, history of past interventions and outcomes'],
            ['Recommendations', 'Suggested learning content, courses to consider dropping, suggested advisor meeting'],
            ['Timeline', 'Full academic event timeline from admission to present'],
        ])

    # 4.4 Smart Alerts
    add_heading(doc, '4.4 Smart Alert System (FR-ALT)', level=2)
    add_table(doc,
        ['Channel', 'When Used', 'Example'],
        [
            ['In-app notification', 'Daily light nudges', '"You missed Programming Fundamentals lecture today"'],
            ['Email', 'Weekly summaries & recommendations', '"Weekly performance summary: 3 courses need attention"'],
            ['SMS', 'Urgent alerts only', '"Warning: Your absence in Course X reached 20% \u2014 limit is 25%"'],
            ['WhatsApp', 'Urgent + interactive', '"Your student Ahmed reached High risk \u2014 see Digital Twin"'],
            ['Advisor notification', 'Cases requiring human intervention', 'Full Digital Twin + AI-generated intervention plan'],
        ])

    # 4.5 Chatbot
    add_heading(doc, '4.5 Advising Chatbot (FR-BOT)', level=2)
    add_para(doc, 'Natural-language AI assistant available 24/7 with university regulations knowledge base.', size=11)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-BOT-01', 'Provide conversational chatbot accessible via web and mobile app', 'Must'],
            ['FR-BOT-02', 'Understand and respond in Arabic (MSA + Saudi dialect) and English', 'Must'],
            ['FR-BOT-03', 'Answer questions about academic policies, registration, deadlines, procedures', 'Must'],
            ['FR-BOT-04', 'Provide personalized advice based on the student\'s Digital Twin data', 'Must'],
            ['FR-BOT-05', 'Escalate to human advisor for complex emotional or academic issues', 'Must'],
            ['FR-BOT-06', 'Use RAG with university regulations knowledge base', 'Must'],
            ['FR-BOT-07', 'Support voice input via speech-to-text (Arabic + English)', 'Should'],
            ['FR-BOT-12', 'Provide citation/source for policy-related answers', 'Must'],
            ['FR-BOT-13', 'Answer academic regulation questions from official university bylaws', 'Must'],
            ['FR-BOT-14', 'Provide personalized answers combining policy + Digital Twin data', 'Must'],
            ['FR-BOT-15', 'Schedule Microsoft Teams meetings between student and advisor', 'Should'],
        ])

    add_heading(doc, 'Academic Policy Knowledge Base (RAG)', level=3)
    add_table(doc,
        ['Document', 'Example Questions Answered'],
        [
            ['Study & Exams Bylaw', '"What are the requirements for First Honors?" \u2014 "How many times can I retake a course?" \u2014 "When does exam barring happen?"'],
            ['Admission & Registration', '"What are the requirements to transfer between colleges?" \u2014 "How do I defer a semester?"'],
            ['Warnings & Dismissal', '"How many warnings before dismissal?" \u2014 "Can I appeal an academic warning?"'],
            ['Academic Calendar', '"When is the last day for add/drop?" \u2014 "When do final exams start?"'],
            ['Self-Service Regulations', '"How do I request an academic transcript?" \u2014 "How do I change my major?"'],
        ])

    add_heading(doc, 'Personalized Policy + Data Answers', level=3)
    add_table(doc,
        ['Student Question', 'Policy-Only Answer', 'QMentor Personalized Answer'],
        [
            ['"What are First Honors requirements?"', 'GPA >= 4.75 with no failures', 'Requirements: GPA >= 4.75 with no failures. Your current GPA is 4.82 with no failures \u2014 you currently qualify \u2705'],
            ['"Can I drop MATH201?"', 'Drop available until week 10', 'Yes, available until Nov 15. But WARNING: MATH201 is a prerequisite for 2 courses \u2014 dropping it delays graduation by 1 semester \u26a0\ufe0f'],
            ['"How many absences left in CS101?"', 'Maximum is 25% of lectures', 'You have 4 absences out of 8 allowed (50%). 4 lectures remaining. If you miss 5, you enter the danger zone \u26a0\ufe0f'],
            ['"I need a meeting with my advisor"', 'Contact your advisor via the portal', 'Your advisor is Dr. Faisal Al-Harbi. Next available: Tuesday 10:00 AM. Shall I book a Teams meeting? \u2705'],
        ])

    add_heading(doc, 'Microsoft Teams Meeting Scheduling', level=3)
    add_para(doc, 'The chatbot integrates with Microsoft Graph API to:', size=11)
    for item in [
        'Show available time slots from the advisor\'s Teams calendar',
        'Let the student pick a slot',
        'Create a Teams meeting and send calendar invitations to both parties',
        'Attach Digital Twin summary in meeting description (visible to advisor only)',
        'Send reminder 1 hour before the meeting',
    ]:
        add_para(doc, f'\u2022 {item}', size=11)

    # 4.6 Advisor Dashboard
    add_heading(doc, '4.6 Advisor Dashboard (FR-ADV)', level=2)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-ADV-01', 'Display all assigned students sorted by risk level (critical first)', 'Must'],
            ['FR-ADV-02', 'Show aggregate statistics: total students, at-risk count per level, trends', 'Must'],
            ['FR-ADV-03', 'Provide one-click access to any student\'s Digital Twin', 'Must'],
            ['FR-ADV-04', 'Show AI-generated intervention plan for each at-risk student', 'Must'],
            ['FR-ADV-05', 'Allow advisor to log intervention notes and outcomes', 'Must'],
            ['FR-ADV-06', 'Show alert history and student response tracking', 'Must'],
            ['FR-ADV-07', 'Allow advisor to manually trigger an alert to a student', 'Should'],
            ['FR-ADV-08', 'Provide calendar view of upcoming advisor-student meetings', 'Should'],
        ])

    doc.add_page_break()

    # ==================== SECTION 5: EXTENDED MODULES ====================
    add_heading(doc, '5. Functional Requirements \u2014 Extended Modules', level=1)

    modules = [
        ('5.1 Graduation Path Optimizer (FR-GPO)', [
            ('FR-GPO-01', 'Calculate remaining semesters to graduation', 'Must'),
            ('FR-GPO-02', 'Generate semester-by-semester course plan respecting prerequisites', 'Must'),
            ('FR-GPO-03', 'Offer multiple paths: fastest, balanced workload, GPA-safe', 'Should'),
            ('FR-GPO-07', 'Simulate "what-if" scenarios (e.g., "What if I drop this course?")', 'Should'),
        ]),
        ('5.2 Career Guidance AI (FR-CGD)', [
            ('FR-CGD-01', 'Map student\'s major, skills, and GPA to relevant career paths', 'Should'),
            ('FR-CGD-02', 'Provide Saudi job market insights relevant to student\'s field', 'Should'),
            ('FR-CGD-03', 'Suggest elective courses that strengthen specific career paths', 'Could'),
        ]),
        ('5.3 Student-to-Student Academic Help (FR-PTM)', [
            ('FR-PTM-01', 'Identify students who excel in courses where their peers struggle — match fellow students (NOT private tutors or university staff)', 'Should'),
            ('FR-PTM-02', 'Suggest peer help matches based on course, schedule, and language — presented as optional suggestions, not mandatory assignments', 'Should'),
            ('FR-PTM-03', 'Both helper and helped student must opt-in; either party can opt-out at any time', 'Must'),
            ('FR-PTM-05', 'Study sessions are proposed as suggestions — students decide whether to accept, reschedule, or decline', 'Must'),
            ('FR-PTM-06', 'Respect gender segregation policies when applicable', 'Must'),
        ]),
        ('5.4 Sentiment & Well-being Analysis (FR-SWB)', [
            ('FR-SWB-01', 'Analyze chatbot conversation tone for distress signals', 'Should'),
            ('FR-SWB-04', 'Provide self-help resources and counseling center info', 'Must'),
            ('FR-SWB-05', 'NEVER diagnose or provide clinical mental health advice', 'Must'),
        ]),
        ('5.5 Mobile Application (FR-MOB)', [
            ('FR-MOB-01', 'Available on iOS (15+) and Android (12+) via Flutter', 'Must'),
            ('FR-MOB-02', 'Support push notifications for all alert types', 'Must'),
            ('FR-MOB-03', 'Provide chatbot interface with voice input', 'Must'),
            ('FR-MOB-07', 'Support Nafath national digital identity for login', 'Should'),
        ]),
        ('5.6 WhatsApp & Telegram Integration (FR-MSG)', [
            ('FR-MSG-01', 'Send alert notifications via WhatsApp Business API', 'Should'),
            ('FR-MSG-02', 'Provide basic chatbot Q&A via WhatsApp', 'Should'),
        ]),
        ('5.7 Faculty Analytics Dashboard (FR-FAC)', [
            ('FR-FAC-01', 'Show instructor a risk heatmap of their section\'s students', 'Should'),
            ('FR-FAC-02', 'Alert instructor about at-risk students in their course', 'Must'),
        ]),
        ('5.8 Gamification Engine (FR-GAM)', [
            ('FR-GAM-01', 'Award points for positive behaviors (attendance streak, on-time submissions)', 'Should'),
            ('FR-GAM-03', 'Award badges for milestones ("Perfect Attendance Week", "GPA Improver")', 'Should'),
        ]),
        ('5.9 Parent/Guardian Portal (FR-PGP)', [
            ('FR-PGP-01', 'Allow students to grant/revoke parent access at any time', 'Must'),
            ('FR-PGP-02', 'Parent sees read-only summary: GPA trend, risk level, attendance', 'Must'),
            ('FR-PGP-04', 'Parent SHALL NOT see chatbot conversations or advisor notes', 'Must'),
        ]),
        ('5.10 Emergency Academic Recovery (FR-EAR)', [
            ('FR-EAR-01', 'Auto-generate personalized recovery plan at Critical risk', 'Must'),
            ('FR-EAR-02', 'Include priority courses, study schedule, recommended content, and suggested fellow students who can help', 'Must'),
            ('FR-EAR-06', 'Simulate what-if scenarios ("If I score X, my GPA = Y")', 'Should'),
        ]),
        ('5.11 Institutional Benchmarking (FR-IBR)', [
            ('FR-IBR-01', 'Generate university-wide risk distribution dashboards', 'Should'),
            ('FR-IBR-03', 'Produce accreditation-ready reports (NCAAA, ABET)', 'Could'),
        ]),
        ('5.12 Smart Scheduling Assistant (FR-SSA)', [
            ('FR-SSA-01', 'Suggest optimal timetable based on study plan and preferences', 'Should'),
            ('FR-SSA-02', 'Detect time conflicts between courses, labs, and final exams', 'Must'),
        ]),
    ]

    for title, reqs in modules:
        add_heading(doc, title, level=2)
        add_table(doc,
            ['ID', 'Requirement', 'Priority'],
            [[r[0], r[1], r[2]] for r in reqs])

    doc.add_page_break()

    # ==================== SECTION 6: AUTONOMY MATRIX ====================
    add_heading(doc, '6. Task Autonomy Matrix \u2014 Agent vs. Human-in-the-Loop', level=1)

    add_para(doc,
        'This matrix defines every system action and whether it is performed autonomously by the '
        'QMentor AI Agent or requires human involvement. This is the authoritative reference for '
        'the development team.', bold=True, size=11)

    doc.add_paragraph('')
    add_para(doc, 'Legend:', bold=True, size=12)
    add_para(doc, '\U0001f916 = Agent (Fully Autonomous) \u2014 no human approval needed', size=11)
    add_para(doc, '\U0001f916\U0001f464 = Agent + Notify \u2014 acts autonomously but notifies the human', size=11)
    add_para(doc, '\U0001f464\U0001f916 = Human Approves \u2014 agent recommends, human must approve', size=11)
    add_para(doc, '\U0001f464 = Human Only \u2014 agent provides data/context, human decides', size=11)
    doc.add_paragraph('')

    # Data Collection
    add_heading(doc, 'Data Collection & Monitoring', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['1', 'Pull attendance data from Blackboard every 2 hours', '\U0001f916', 'QMentor Agent'],
            ['2', 'Pull grades from Oracle SIS every 6 hours', '\U0001f916', 'QMentor Agent'],
            ['3', 'Pull assignment submissions from Blackboard every 4 hours', '\U0001f916', 'QMentor Agent'],
            ['4', 'Pull registration & academic standing from SIS daily', '\U0001f916', 'QMentor Agent'],
            ['5', 'Track LMS login frequency and content interaction', '\U0001f916', 'QMentor Agent'],
            ['6', 'Detect data anomalies (sudden drops, patterns)', '\U0001f916', 'QMentor Agent'],
            ['7', 'Build and update Digital Twin for every student', '\U0001f916', 'QMentor Agent'],
        ])

    add_heading(doc, 'Risk Assessment & Prediction', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['8', 'Calculate risk score (0\u2013100) using 66 criteria', '\U0001f916', 'QMentor Agent'],
            ['9', 'Classify student risk level (Low/Medium/High/Critical)', '\U0001f916', 'QMentor Agent'],
            ['10', 'Predict course failure probability', '\U0001f916', 'QMentor Agent'],
            ['11', 'Predict absence limit probability', '\U0001f916', 'QMentor Agent'],
            ['12', 'Detect declining performance trends', '\U0001f916', 'QMentor Agent'],
            ['13', 'Generate risk score explanation (top 3 factors)', '\U0001f916', 'QMentor Agent'],
            ['14', 'Manually adjust student risk level', '\U0001f464', 'Academic Advisor'],
        ])

    add_heading(doc, 'Alerts \u2014 Low/Medium Risk', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['15', 'Send in-app nudge for missed lecture', '\U0001f916', 'QMentor Agent'],
            ['16', 'Send weekly performance summary email', '\U0001f916', 'QMentor Agent'],
            ['17', 'Suggest learning content for weak courses', '\U0001f916', 'QMentor Agent'],
            ['18', 'Send attendance warning (approaching limit)', '\U0001f916\U0001f464', 'Agent (advisor notified)'],
            ['19', 'Award gamification points/badges', '\U0001f916', 'QMentor Agent'],
        ])

    add_heading(doc, 'Alerts \u2014 High Risk', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['20', 'Send SMS urgent alert to student', '\U0001f916\U0001f464', 'Agent (advisor notified)'],
            ['21', 'Send full Digital Twin report to advisor', '\U0001f916', 'QMentor Agent'],
            ['22', 'Generate AI intervention plan for advisor', '\U0001f916', 'QMentor (recommendation)'],
            ['23', 'Advisor reviews and approves intervention plan', '\U0001f464\U0001f916', 'Academic Advisor'],
            ['24', 'Send intervention message to student (after approval)', '\U0001f464\U0001f916', 'Agent after advisor approval'],
        ])

    add_heading(doc, 'Alerts \u2014 Critical Risk', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['25', 'Generate emergency recovery plan', '\U0001f916', 'QMentor (recommendation only)'],
            ['26', 'Advisor reviews recovery plan', '\U0001f464\U0001f916', 'Academic Advisor'],
            ['27', 'Escalate to Department Head', '\U0001f464\U0001f916', 'Advisor must approve'],
            ['28', 'Escalate to Dean (if no response in 48hrs)', '\U0001f464', 'Department Head initiates'],
            ['29', 'Recommend academic probation/dismissal', '\U0001f464', 'University committee only'],
            ['30', 'Issue formal academic warning', '\U0001f464', 'Registrar office only'],
        ])

    add_heading(doc, 'Chatbot Interactions', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['31', 'Answer academic policy questions (bylaws & regulations)', '\U0001f916', 'QMentor Chatbot'],
            ['32', 'Answer personalized questions (policy + Digital Twin)', '\U0001f916', 'QMentor Chatbot'],
            ['33', 'Schedule Teams meeting with advisor', '\U0001f916', 'QMentor via Graph API'],
            ['34', 'Send Teams calendar invitation to both parties', '\U0001f916', 'QMentor via Graph API'],
            ['35', 'Suggest course drop with graduation impact analysis', '\U0001f916', 'QMentor (recommendation)'],
            ['36', 'Actually drop a course in SIS', '\U0001f464', '\u274c NEVER \u2014 student does it themselves'],
            ['37', 'Detect emotional distress in conversation', '\U0001f916', 'QMentor flags automatically'],
            ['38', 'Provide mental health counseling', '\U0001f464', '\u274c NEVER \u2014 routes to counseling center'],
            ['39', 'Escalate complex question to human advisor', '\U0001f916', 'QMentor escalates automatically'],
        ])

    add_heading(doc, 'Academic Planning & Administration', level=3)
    add_table(doc,
        ['#', 'Task', 'Mode', 'Who Acts'],
        [
            ['40', 'Calculate remaining semesters to graduation', '\U0001f916', 'QMentor Agent'],
            ['41', 'Generate semester-by-semester course plan', '\U0001f916', 'QMentor (recommendation)'],
            ['42', 'Simulate "what-if" scenarios', '\U0001f916', 'QMentor Agent'],
            ['43', 'Warn when choices may delay graduation', '\U0001f916', 'QMentor Agent'],
            ['44', 'Approve student\'s actual course registration', '\U0001f464', 'Advisor approves in SIS'],
            ['45', 'Approve major change request', '\U0001f464', 'Department + Dean committee'],
            ['51', 'Generate institutional analytics reports', '\U0001f916', 'QMentor Agent'],
            ['54', 'Configure college-specific thresholds', '\U0001f464', 'Dept Head + Admin'],
            ['55', 'Disable/enable automated features (kill switch)', '\U0001f464', 'System Admin'],
            ['56', 'Retrain ML models with new semester data', '\U0001f464\U0001f916', 'AI Team approves, pipeline executes'],
        ])

    doc.add_paragraph('')
    add_heading(doc, 'Summary Count', level=3)
    add_table(doc,
        ['Mode', 'Count', 'Percentage'],
        [
            ['\U0001f916 Agent (Fully Autonomous)', '30', '52%'],
            ['\U0001f916\U0001f464 Agent + Notify', '5', '9%'],
            ['\U0001f464\U0001f916 Human Approves, Agent Executes', '8', '14%'],
            ['\U0001f464 Human Only', '15', '25%'],
            ['Total', '58', '100%'],
        ])

    doc.add_paragraph('')
    add_para(doc,
        'KEY PRINCIPLE: QMentor automates 61% of tasks (Agent + Agent+Notify) but NEVER performs '
        'irreversible academic actions (course drops, grade changes, sanctions, dismissals). '
        'These remain exclusively human decisions. The agent\'s role is to ensure the right information '
        'reaches the right person at the right time.',
        bold=True, size=11, color=RGBColor(31, 78, 121))

    doc.add_page_break()

    # ==================== SECTION 7: SDAIA ETHICS ====================
    add_heading(doc, '7. SDAIA AI Ethics Compliance', level=1)
    add_para(doc,
        'QMentor complies with all 6 principles of the SDAIA AI Ethics Framework (2023). '
        'Each principle maps to concrete, verifiable requirements.', size=11)

    principles = [
        ('Principle 1: Fairness & Non-Discrimination', [
            'Risk models audited for bias across gender, college, nationality, disability before deployment and after each retraining',
            'Nationality, gender, and socioeconomic status SHALL NOT be direct features in risk models',
            'Quarterly Fairness Reports comparing risk distribution across demographic groups',
            'If bias detected (>5% disparity), model retrained or decommissioned within 2 weeks',
            'Alert messages are culturally neutral \u2014 no assumptions about family, finances, or personal life',
        ]),
        ('Principle 2: Transparency & Explainability', [
            'Every risk score includes human-readable explanation with top 3 contributing factors',
            'Student can view all collected data via "My Data" page in the app',
            'All AI-generated content labeled as "Generated by AI"',
            'Chatbot discloses it is an AI at the start of every conversation',
            'Advisor sees confidence level for each prediction (e.g., "87% confidence")',
            'AI Decision Log accessible to governance committee',
        ]),
        ('Principle 3: Accountability & Governance', [
            'AI Governance Committee (deans + IT + legal + student rep) reviews quarterly',
            'Every automated action logged with timestamp, reasoning, data used, and outcome',
            'Model Registry documenting every version, training data, and approval status',
            'Student has formal grievance mechanism to challenge AI risk assessments',
            'Annual external AI audit by SDAIA-recognized auditor',
        ]),
        ('Principle 4: Privacy & Data Protection', [
            'Purpose limitation \u2014 only data necessary for academic advising is collected',
            'Explicit, informed, granular consent \u2014 separate opt-ins for monitoring, chatbot, WhatsApp, parent access',
            'Student can withdraw consent at any time \u2014 system degrades gracefully',
            'All data remains within Saudi Arabia \u2014 no cross-border transfer',
            'Privacy Impact Assessment (DPIA) before launching each new module',
        ]),
        ('Principle 5: Safety & Reliability', [
            'Kill switch \u2014 admin can disable all automated interventions within 5 minutes',
            'Mental health flags handled by human counselors ONLY \u2014 AI routes, never diagnoses',
            'System NEVER performs irreversible actions (no course drops, no grade changes, no sanctions)',
            'False positive rate for Critical alerts SHALL NOT exceed 10%',
            'Canary deployments \u2014 new models run in shadow mode 2 weeks before production',
        ]),
        ('Principle 6: Human Oversight & Control', [
            'Advisor can override any AI recommendation with documented reason',
            'Advisor can manually adjust student risk level (up or down) with justification',
            'AI presents recommendations as suggestions \u2014 never directives',
            'Tiered automation: Low=auto-send, Medium=auto+notify, High=advisor approves, Critical=advisor+dept head approve',
            'If override rate >30% for a model, automatic model review triggered',
        ]),
    ]

    for title, items in principles:
        add_heading(doc, title, level=2)
        for item in items:
            add_para(doc, f'\u2713 {item}', size=10)

    doc.add_page_break()

    # ==================== SECTION 8: AGENTIC AI ====================
    add_heading(doc, '8. Agentic AI Principles & Guardrails', level=1)

    add_para(doc,
        'QMentor is an Agentic AI system \u2014 an autonomous agent that observes, decides, and acts '
        'continuously. This requires specific guardrails beyond standard AI ethics.', size=11)

    add_heading(doc, 'Autonomy Levels', level=2)
    add_table(doc,
        ['Level', 'Description', 'QMentor Application', 'Human Approval?'],
        [
            ['Level 0: Informational', 'AI collects and presents data', 'Digital Twin, dashboards, reports', 'No'],
            ['Level 1: Advisory', 'AI recommends, human decides', 'Recovery plans, graduation suggestions', 'No (recommendation)'],
            ['Level 2: Supervised Autonomous', 'AI acts, human notified', 'Low/Medium risk alerts', 'No, but advisor notified'],
            ['Level 3: Conditional Autonomous', 'AI acts after human pre-approves', 'High risk escalation', 'Yes \u2014 advisor confirms'],
            ['Level 4: Human-Only', 'AI provides data, human decides', 'Critical interventions, mental health', 'Yes \u2014 mandatory'],
        ])

    add_heading(doc, '15 Safety Guardrails', level=2)

    guardrails = [
        ('AGT-05', 'No Irreversible Actions', 'Agent NEVER drops courses, changes grades, or issues sanctions. Recommends only.'),
        ('AGT-06', 'Rate Limiting', 'Max 1 alert per student per channel per day (except Critical).'),
        ('AGT-07', 'Hallucination Prevention', 'Agent references only data actually in Digital Twin. No fabricated grades/attendance.'),
        ('AGT-08', 'Scope Containment', 'Operates ONLY within academic advising. No legal, medical, financial, or personal advice.'),
        ('AGT-09', 'Temporal Awareness', 'Considers semester phase (add/drop, midterm, finals) when calibrating thresholds.'),
        ('AGT-10', 'No Conflict of Interest', 'Risk model ignores advisor workload. Risk based solely on student indicators.'),
        ('AGT-11', 'Feedback Loop Integrity', 'Own previous alerts NOT used as model input features (prevents self-reinforcing bias).'),
        ('AGT-12', 'Graceful Degradation', 'If data source unavailable, continues with available data + flags reduced confidence.'),
        ('AGT-13', 'Student Agency', 'Always presents options, not directives. Student can dismiss/snooze non-critical alerts.'),
        ('AGT-14', 'Cultural Sensitivity', 'Arabic output reviewed by native speakers. No patronizing, shaming, or anxiety-inducing language.'),
        ('AGT-15', 'Adversarial Robustness', 'Tested against adversarial inputs (students gaming LMS to manipulate risk scores).'),
    ]

    for gid, title, desc in guardrails:
        add_para(doc, f'{gid}: {title}', bold=True, size=11)
        add_para(doc, desc, size=10, color=RGBColor(89, 89, 89))

    doc.add_page_break()

    # ==================== SECTION 9: SECURITY ====================
    add_heading(doc, '9. Security & Sensitive Data Handling', level=1)

    add_heading(doc, '4-Tier Data Classification', level=2)
    add_table(doc,
        ['Tier', 'Classification', 'Examples', 'Handling'],
        [
            ['T1: Public', 'Non-sensitive', 'University name, course catalog', 'No extra encryption; cacheable via CDN'],
            ['T2: Internal', 'Low sensitivity', 'Aggregate statistics, anonymized analytics', 'Encrypted at rest; authenticated access'],
            ['T3: Confidential', 'High sensitivity', 'GPA, grades, attendance, risk scores, advisor notes', 'Field-level encryption (KMS); role-based access; audit logged'],
            ['T4: Restricted', 'Critical sensitivity', 'National ID, phone, medical records, well-being flags', 'Column-level encryption (separate KMS key); MFA + justification; alert on access'],
        ])

    add_heading(doc, 'Sensitive Data Handling Scenarios', level=2)
    add_table(doc,
        ['Scenario', 'Data', 'Tier', 'Protocol'],
        [
            ['Student asks chatbot about dropping a course', 'Conversation + study plan', 'T3', 'Encrypted; retained 2 years; student can request deletion'],
            ['AI detects potential mental distress', 'Well-being flag + conversation', 'T4', 'Dedicated KMS key; ONLY counseling services can access; not visible to advisor unless student consents'],
            ['Advisor views Digital Twin', 'GPA, grades, attendance, risk', 'T3', 'Logged (who, when, which student); session-based access'],
            ['Parent requests student data', 'Summary view (if consent)', 'T3', 'Nafath-verified identity; read-only; subset of Digital Twin'],
            ['ML model training', 'Historical grades/attendance', 'T2 (anonymized)', 'k-anonymized (k>=5); differential privacy; no PII; isolated S3 bucket'],
            ['National ID for Nafath login', 'National ID number', 'T4', 'NEVER stored in QMentor DB; passthrough to Nafath API only'],
        ])

    doc.add_page_break()

    # ==================== SECTION 11: AWS ====================
    add_heading(doc, '11. AWS Infrastructure & Cost Estimate', level=1)

    add_table(doc,
        ['Layer', 'AWS Service', 'Purpose'],
        [
            ['Compute', 'ECS Fargate', 'Laravel backend containers'],
            ['GPU Compute', 'EC2 G5.xlarge', 'Ollama LLM inference (NVIDIA A10G)'],
            ['GPU Compute', 'SageMaker Endpoints', 'ML model serving'],
            ['Serverless', 'Lambda', 'Event processing, data pipelines'],
            ['LLM', 'Bedrock', 'Claude/Llama API access'],
            ['NLP', 'Comprehend', 'Arabic sentiment analysis'],
            ['Database', 'RDS PostgreSQL', 'Primary relational data (Multi-AZ)'],
            ['Cache', 'ElastiCache Redis', 'Application cache, rate limiting'],
            ['Vector DB', 'ChromaDB on ECS', 'Semantic search embeddings'],
            ['Storage', 'S3', 'Files, data lake, backups'],
            ['Queue', 'SQS + EventBridge', 'Async tasks + scheduled jobs'],
            ['Auth', 'Cognito + SAML 2.0', 'SSO via MyQU'],
            ['Security', 'WAF + Shield + KMS', 'DDoS protection, encryption'],
            ['Monitoring', 'CloudWatch + X-Ray', 'Logs, metrics, tracing'],
        ])

    add_heading(doc, 'Monthly Cost Estimate', level=2)
    add_table(doc,
        ['Service', 'Monthly Cost (USD)'],
        [
            ['ECS Fargate (Laravel + microservices)', '$400\u2013600'],
            ['EC2 G5.xlarge (GPU \u2014 1 baseline + auto-scale)', '$800\u20132,000'],
            ['SageMaker endpoints', '$500\u20131,000'],
            ['Bedrock API calls', '$300\u2013800'],
            ['RDS PostgreSQL (Multi-AZ)', '$400\u2013600'],
            ['ElastiCache Redis', '$150\u2013200'],
            ['S3 + CloudFront', '$100\u2013200'],
            ['SES + SNS (email + SMS)', '$200\u2013500'],
            ['Other (CloudWatch, KMS, WAF...)', '$200\u2013300'],
            ['TOTAL ESTIMATED', '$3,050\u20136,200/month'],
        ])

    doc.add_page_break()

    # ==================== SECTION 12: PHASES ====================
    add_heading(doc, '12. Project Phases & Timeline', level=1)

    phases = [
        ('Phase 1: Foundation (Months 1\u20133) \u2014 Q2 2026', [
            'AWS infrastructure setup (VPC, RDS, ECS, S3, Cognito, WAF, KMS)',
            'Data pipeline MVP: Blackboard + Oracle SIS integration (read-only)',
            'Digital Twin v1: Basic academic profile (GPA, courses, attendance)',
            'Alert system v1: Email + in-app notifications for high/critical risk',
            'Advisor dashboard v1: Student list with risk levels + Digital Twin view',
            'MyQU SSO integration (SAML 2.0)',
        ]),
        ('Phase 2: Intelligence (Months 4\u20136) \u2014 Q3 2026', [
            'Predictive analytics v1: Risk classification model trained on historical data',
            'Chatbot v1: Arabic + English with university policy RAG',
            'Academic policy knowledge base (honors, drops, warnings, transfers...)',
            'Microsoft Teams meeting scheduling from chatbot',
            'Smart alert escalation with configurable thresholds per college',
            'GPU infrastructure: G5 instances + SageMaker',
            'Mobile app v1 (Beta): Student chatbot + alerts + Digital Twin',
        ]),
        ('Phase 3: Expansion (Months 7\u20139) \u2014 Q4 2026', [
            'Graduation path optimizer with prerequisite awareness',
            'Student-to-student academic help (voluntary peer matching) + session tracking',
            'WhatsApp integration: Alert delivery + basic Q&A',
            'Faculty dashboard: Section risk heatmap + alerts for instructors',
            'Gamification v1: Points, badges, streaks',
            'Sentiment analysis: Chatbot tone analysis + well-being flags',
        ]),
        ('Phase 4: Optimization (Months 10\u201312) \u2014 Q1 2027', [
            'Career guidance AI: Career path mapping + elective suggestions',
            'Predictive analytics v2: Refined models with first-semester data',
            'Parent/guardian portal (optional, with student consent)',
            'Leadership dashboard: University-wide analytics',
            'Performance optimization: Cost, caching, model fine-tuning',
            'Security audit: NCA compliance + penetration testing',
        ]),
        ('Phase 5: Maturity (Months 13\u201318) \u2014 Q2\u2013Q3 2027', [
            'Full university rollout: All colleges after pilot validation',
            'Telegram bot: Additional messaging channel',
            'Accreditation reports: Automated NCAAA/ABET data generation',
            'Open-source release on code.gov.sa',
            'Award submissions: DGA Award, Global AI Summit, WSIS Prizes',
        ]),
    ]

    for title, items in phases:
        add_heading(doc, title, level=2)
        for item in items:
            add_para(doc, f'\u25c6 {item}', size=10)

    # ==================== FOOTER ====================
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2501' * 50)
    run.font.color.rgb = RGBColor(31, 78, 121)

    add_para(doc, 'Prepared by: Digital Transformation Team \u2014 Deanship of E-Learning & IT, Qassim University',
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Next revision due: Upon completion of Phase 1 (end of Q2 2026)',
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'QMentor_SRS_EN.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_docx()

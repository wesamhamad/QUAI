#!/usr/bin/env python3
"""
Generate Smart Advisor V1 (المرشد الذكي) — Current Solution Technical Document.
Documents what EXISTS and WORKS today.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BLUE = RGBColor(31, 78, 121)
BLUE2 = RGBColor(68, 114, 196)
GREEN = RGBColor(39, 124, 72)
GRAY = RGBColor(100, 100, 100)
RED = RGBColor(192, 0, 0)
WHITE = RGBColor(255, 255, 255)


def add_p(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def add_h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))


def add_t(doc, headers, rows, hdr_bg="1F4E79"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Calibri'
        r.font.color.rgb = WHITE
        shade(c, hdr_bg)
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(8.5)
            r.font.name = 'Calibri'
            if ri % 2 == 0:
                shade(c, "E8F0FE")
    doc.add_paragraph('')
    return t


def blue_box(doc, text):
    p = add_p(doc, text, bold=True, size=10, color=BLUE, after=8)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'))
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="4" w:color="1F4E79"/></w:pBdr>'))
    return p


def bullet(doc, text, size=10):
    add_p(doc, f'\u2022  {text}', size=size, after=2)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2)
        hdr = sec.header.paragraphs[0]
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr.add_run('Smart Advisor V1  |  Qassim University  |  Current Implementation')
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        ftr = sec.footer.paragraphs[0]
        ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ftr.add_run('V1.0  |  ')
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        ftr._p.append(parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:sz w:val="16"/><w:color w:val="787878"/></w:rPr>'
            f'<w:t>1</w:t></w:r></w:fldSimple>'))

    # ═══════════ COVER ═══════════
    for _ in range(3):
        doc.add_paragraph('')
    add_p(doc, 'QASSIM UNIVERSITY', bold=True, size=22, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_p(doc, 'Deanship of E-Learning & IT \u2014 Digital Transformation Team', size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_p(doc, '\u2501' * 40, size=8, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '\u0627\u0644\u0645\u0631\u0634\u062f \u0627\u0644\u0630\u0643\u064a', bold=True, size=36, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_p(doc, 'Smart Advisor V1', bold=True, size=28, color=BLUE2, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_p(doc, 'AI-Powered Academic Advising Chatbot', size=16, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_p(doc, '\u2501' * 40, size=8, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Current Implementation Technical Document', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_p(doc, 'Version 1.0  |  April 2026', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    for tag in [
        'Production-Ready \u2014 Currently Serving Students',
        'Arabic-First (RTL) with English Support',
        'RAG Knowledge Base: University Bylaws & Regulations',
        'Real-Time Student Data from Blackboard LMS + Oracle SIS',
        'Human Escalation with Conversation Context',
    ]:
        add_p(doc, f'\u2705  {tag}', size=10, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    doc.add_page_break()

    # ═══════════ DOC INFO ═══════════
    add_h(doc, '1. Document Information', 1)
    add_t(doc, ['Field', 'Value'], [
        ['Product Name', 'Smart Advisor V1 (\u0627\u0644\u0645\u0631\u0634\u062f \u0627\u0644\u0630\u0643\u064a)'],
        ['Organization', 'Deanship of E-Learning & IT \u2014 Digital Transformation Team'],
        ['Platform', 'Laravel 11 (PHP 8.3) + Ollama (Self-hosted LLM)'],
        ['Status', 'Production-Ready \u2014 Currently deployed and serving students'],
        ['AI Models', 'command-r7b-arabic (fluency) + Qwen2.5:14b (regulation accuracy)'],
        ['Data Sources', 'Blackboard LMS API + Oracle SIS API + MyQU Portal'],
        ['Knowledge Base', 'RAG with BM25 ranking on university bylaws & regulations'],
        ['Authentication', 'SAML 2.0 SSO via MyQU (students only)'],
        ['Version', '1.0'],
        ['Date', '2026-04-07'],
    ])

    doc.add_page_break()

    # ═══════════ TOC ═══════════
    add_h(doc, 'Table of Contents', 1)
    for num, title in [
        ('1', 'Document Information'),
        ('2', 'System Overview & Architecture'),
        ('3', 'Current Features (What Works Today)'),
        ('4', 'Intent Detection & Data Routing'),
        ('5', 'Available API Endpoints & Data Sources'),
        ('6', 'Knowledge Base (RAG) \u2014 Academic Regulations'),
        ('7', 'AI Model Strategy (Hybrid Selection)'),
        ('8', 'Conversation Management & Storage'),
        ('9', 'Human Escalation Flow'),
        ('10', 'Response Quality Validation'),
        ('11', 'Security & Content Moderation'),
        ('12', 'Technical Configuration'),
        ('13', 'File Locations & Codebase Map'),
        ('14', 'Current Limitations & Known Gaps'),
        ('15', 'Roadmap to QMentor V2'),
    ]:
        add_p(doc, f'{num}.\t{title}', bold=True, size=11, after=3)

    doc.add_page_break()

    # ═══════════ 2. OVERVIEW ═══════════
    add_h(doc, '2. System Overview & Architecture', 1)

    add_p(doc,
        'The Smart Advisor is a production AI-powered academic advising chatbot that provides '
        'real-time, personalized answers to students by combining three data sources: '
        'the student\'s live academic data (Blackboard + SIS), a knowledge base of university '
        'regulations (RAG), and conversational AI (self-hosted LLM).', size=11)

    add_h(doc, '2.1 Data Flow', 2)
    flow = (
        'User Question\n'
        '    \u2193\n'
        'API Controller (SmartAdvisorController)\n'
        '    \u2193\n'
        'Content Moderation (SDAIA compliance check)\n'
        '    \u2193\n'
        'Intent Detection (AI-powered + keyword fallback)\n'
        '    \u2193\n'
        'Student Context Building (UniversityApiClient \u2192 Blackboard/SIS)\n'
        '    \u2193\n'
        'Knowledge Base Retrieval (BM25 search on regulations PDFs)\n'
        '    \u2193\n'
        'Message Payload Construction (system prompt + context + history)\n'
        '    \u2193\n'
        'AI Model Selection (Hybrid: command-r7b-arabic or Qwen2.5:14b)\n'
        '    \u2193\n'
        'Response Quality Validation (0\u2013100 score)\n'
        '    \u2193\n'
        'Conversation Storage (advisor_messages DB)\n'
        '    \u2193\n'
        'Return to User (Streaming SSE or JSON)'
    )
    p = add_p(doc, flow, size=9, color=RGBColor(64, 64, 64), after=8)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'))

    add_h(doc, '2.2 Tech Stack', 2)
    add_t(doc, ['Component', 'Technology', 'Details'], [
        ['Backend', 'Laravel 11 (PHP 8.3)', 'Main application framework'],
        ['LLM (Primary)', 'Ollama (self-hosted)', 'command-r7b-arabic + Qwen2.5:14b'],
        ['LLM (Fallback)', 'Groq Cloud API', 'llama-3.3-70b-versatile (optional, costs money)'],
        ['Data API', 'UniversityApiClient', 'REST calls to Blackboard LMS + Oracle SIS'],
        ['Knowledge Base', 'BM25 keyword ranking', 'Chunked PDFs from storage/app/advisor_knowledge/'],
        ['Semantic Search', 'Python microservice (optional)', 'http://localhost:8002 \u2014 disabled by default'],
        ['Search Engine', 'SearXNG on Docker', 'http://localhost:8888 \u2014 for web search context'],
        ['Authentication', 'SAML 2.0 SSO', 'MyQU portal integration'],
        ['Database', 'MySQL/PostgreSQL', 'advisor_conversations, advisor_messages, advisor_inquiries'],
        ['Streaming', 'Server-Sent Events (SSE)', 'Token-by-token streaming to UI'],
        ['Frontend', 'Blade + Alpine.js', 'RTL Arabic chat interface'],
    ])

    doc.add_page_break()

    # ═══════════ 3. CURRENT FEATURES ═══════════
    add_h(doc, '3. Current Features (What Works Today)', 1)

    blue_box(doc, 'Everything listed below is IMPLEMENTED, TESTED, and DEPLOYED. This is not a wishlist \u2014 it is the current state of the system.')

    add_t(doc, ['#', 'Feature', 'Status', 'Description'], [
        ['1', 'AI Chat Interface', '\u2705 Live', 'Web chat with streaming responses, RTL Arabic UI, conversation history sidebar'],
        ['2', 'Streaming Responses (SSE)', '\u2705 Live', 'Token-by-token streaming via Server-Sent Events for real-time typing effect'],
        ['3', 'Intent Detection', '\u2705 Live', 'AI-powered intent classification (13 intents) with keyword fallback'],
        ['4', 'Real-Time Student Data', '\u2705 Live', 'Live pull from Blackboard + SIS: courses, grades, absences, timetable, exams, plan'],
        ['5', 'Knowledge Base (RAG)', '\u2705 Live', 'BM25 search on university bylaws: study & exam regulations, student guide, calendar'],
        ['6', 'Hybrid Model Selection', '\u2705 Live', 'Auto-selects best model: command-r7b for fluency, Qwen2.5 for regulation accuracy'],
        ['7', 'Conversation History', '\u2705 Live', 'Full history stored in DB, resumable conversations, auto-generated titles'],
        ['8', 'Human Escalation', '\u2705 Live', '5 categories, attaches conversation summary, tracks status (pending \u2192 resolved)'],
        ['9', 'Reference Citations', '\u2705 Live', 'PDF download links with page numbers for every regulation-based answer'],
        ['10', 'Content Moderation', '\u2705 Live', 'SDAIA-compliant input filtering before AI processing'],
        ['11', 'Response Quality Scoring', '\u2705 Live', '0\u2013100 quality score with 7 checks (language, length, data usage, contradictions)'],
        ['12', 'Student-Only Access', '\u2705 Live', 'EnsureStudentAccess middleware blocks non-students (admins exempt for testing)'],
        ['13', 'Data Formatting', '\u2705 Live', '12 formatters convert raw API data to readable Arabic tables with emoji status indicators'],
        ['14', 'Conversation Archiving', '\u2705 Live', 'Soft-delete (archive) conversations, preserved for analytics'],
        ['15', 'Groq Cloud Fallback', '\u2705 Ready', 'Optional Groq API for larger model (70b) on regulation queries, disabled by default'],
    ])

    doc.add_page_break()

    # ═══════════ 4. INTENT DETECTION ═══════════
    add_h(doc, '4. Intent Detection & Data Routing', 1)

    add_p(doc, 'The system uses a two-tier intent detection to understand what the student is asking, then routes to the appropriate data sources.', size=11)

    add_h(doc, '4.1 Supported Intents (13)', 2)
    add_t(doc, ['Intent', 'Arabic Examples', 'APIs Called', 'Data Retrieved'], [
        ['courses', '\u0643\u0645 \u0645\u0642\u0631\u0631 \u0645\u0633\u062c\u0644 \u0647\u0627\u0644\u0641\u0635\u0644\u061f', '/student/courses + /student-plan', 'Current courses, study plan, completion %'],
        ['schedule', '\u0645\u062a\u0649 \u0645\u062d\u0627\u0636\u0631\u0629 \u0627\u0644\u0628\u0631\u0645\u062c\u0629\u061f', '/time-table + /student/courses', 'Weekly timetable, room numbers, instructors'],
        ['exams', '\u0645\u062a\u0649 \u0627\u062e\u062a\u0628\u0627\u0631 \u0627\u0644\u0631\u064a\u0627\u0636\u064a\u0627\u062a\u061f', '/final-exams', 'Exam dates, times, locations, conflicts'],
        ['absences', '\u0643\u0645 \u063a\u064a\u0627\u0628\u064a \u0641\u064a \u0643\u0644 \u0645\u0642\u0631\u0631\u061f', '/absences-with-details + courses', 'Per-course absence %, barring risk, allowed remaining'],
        ['academic_standing', '\u0648\u0634 \u0648\u0636\u0639\u064a \u0627\u0644\u0623\u0643\u0627\u062f\u064a\u0645\u064a\u061f', '/student-academic-transactions', 'GPA, warnings, probation status, academic history'],
        ['grades', '\u0648\u0634 \u062f\u0631\u062c\u0627\u062a\u064a\u061f', '/blackboard/courses/{id}/grades (all)', 'Per-course grades: quizzes, midterm, assignments, final'],
        ['financial', '\u0645\u062a\u0649 \u062a\u0646\u0632\u0644 \u0627\u0644\u0645\u0643\u0627\u0641\u0623\u0629\u061f', '/rewards', 'Scholarship status, allowance amounts'],
        ['advisor', '\u0645\u0646 \u0645\u0631\u0634\u062f\u064a \u0627\u0644\u0623\u0643\u0627\u062f\u064a\u0645\u064a\u061f', '/academic-advisor', 'Advisor name, email, phone, office'],
        ['calendar', '\u0645\u062a\u0649 \u0622\u062e\u0631 \u0645\u0648\u0639\u062f \u0644\u0644\u062d\u0630\u0641\u061f', '/academic-calendar', 'Semester dates, registration deadlines, holidays'],
        ['announcements', '\u0641\u064a\u0647 \u0625\u0639\u0644\u0627\u0646\u0627\u062a \u062c\u062f\u064a\u062f\u0629\u061f', '/blackboard/announcements (all)', 'Course announcements, grouped by course'],
        ['course_content', '\u0648\u064a\u0646 \u0645\u0644\u0641\u0627\u062a \u0627\u0644\u0645\u0627\u062f\u0629\u061f', '/blackboard/courses/{id}/contents (all)', 'Course materials, files, upload dates'],
        ['regulations', '\u0645\u0627 \u0634\u0631\u0648\u0637 \u0645\u0631\u062a\u0628\u0629 \u0627\u0644\u0634\u0631\u0641\u061f', 'Knowledge Base (RAG)', 'University bylaws, policies, regulations'],
        ['general', '\u0633\u0644\u0627\u0645 / \u0639\u0646\u062f\u064a \u0633\u0624\u0627\u0644', 'courses + absences + timetable', 'Default data set for general queries'],
    ])

    add_h(doc, '4.2 Detection Method', 2)
    add_p(doc, 'Tier 1 (Default): AI-powered \u2014 lightweight model analyzes the user\'s message and returns structured JSON with intents + confidence. Results cached 1 hour per unique message.', size=10)
    add_p(doc, 'Tier 2 (Fallback): Keyword matching \u2014 Arabic/English keyword dictionaries, used if AI detection fails or is disabled.', size=10)

    doc.add_page_break()

    # ═══════════ 5. API ENDPOINTS ═══════════
    add_h(doc, '5. Available API Endpoints & Data Sources', 1)

    add_h(doc, '5.1 UniversityApiClient (Oracle SIS + Blackboard)', 2)
    add_p(doc, 'Base: quai.blackboard.api_base_url | Auth: Service Token | Cache: 300s', size=9, color=GRAY)
    add_t(doc, ['Method', 'Endpoint', 'Returns', 'Risk Indicators'], [
        ['getStudentProfile()', 'GET /api/v3/me', 'Name, ID, GPA, major, college, contact', 'G-04, G-05, AC-01, AC-04, AC-05'],
        ['getAbsences()', 'GET /absences-with-details', 'Absence records: dates, reasons, per-course', 'A-01 to A-08'],
        ['getAcademicTransactions()', 'GET /student-academic-transactions', 'Drops, withdrawals, warnings history', 'R-01 to R-06, AC-02'],
        ['getCurrentCourses()', 'GET /student/courses', 'Enrolled courses + bb_course_id', 'AC-03, AC-06, R-04'],
        ['getStudentPlan()', 'GET /student-plan', 'Full academic plan, completed/remaining', 'P-01 to P-05'],
        ['getAcademicPlanByMajor()', 'GET /api/v1/academic_plan/{major}', 'Degree requirements per major', 'GPO elective classification'],
        ['getTimeTable()', 'GET /time-table', 'Weekly class schedule', 'T-01, T-04'],
        ['getFinalExams()', 'GET /final-exams', 'Exam dates, times, rooms', 'T-01, T-03'],
        ['getAcademicCalendar()', 'GET /academic-calendar', 'Semester dates, deadlines', 'AGT-09 temporal'],
        ['getAdvisorInfo()', 'GET /academic-advisor', 'Advisor name, email, office', 'Escalation routing'],
        ['getRewards()', 'GET /rewards', 'Scholarships, allowances', 'Financial queries'],
        ['getStudentSkills()', 'GET /api/v1/skills/student/{id}', 'Student competencies', 'Peer matching'],
        ['fetchMultiple()', 'GET (concurrent pool)', 'Parallel fetch of multiple endpoints', 'Digital Twin batch'],
    ])

    add_h(doc, '5.2 BlackboardApiClient (LMS)', 2)
    add_p(doc, 'Base: quai.blackboard.api_base_url | Auth: JWT Bearer | Timeout: 30s', size=9, color=GRAY)
    add_t(doc, ['Method', 'Endpoint', 'Returns'], [
        ['getCourseGrades()', 'GET /blackboard/courses/{id}/grades', 'Quizzes, midterms, assignments, final per course'],
        ['getAllCourseGrades()', 'Iterates all enrolled courses', 'Grades across ALL courses with metadata'],
        ['getCourseContent()', 'GET /blackboard/courses/{id}/contents', 'Content tree structure + access data'],
        ['getAllCourseContents()', 'Iterates all enrolled courses', 'All materials across all courses'],
        ['getCourseAttachments()', 'GET .../contents/{cid}/attachments', 'Assignment attachments + submission data'],
        ['getAnnouncements()', 'GET /blackboard/announcements', 'All announcements across courses'],
        ['downloadAttachment()', 'GET .../attachments/{id}/download', 'Binary file download (PDF, etc.)'],
    ])

    add_h(doc, '5.3 AI & Processing Services', 2)
    add_t(doc, ['Service', 'Endpoint', 'Purpose'], [
        ['OllamaService', 'POST /v1/chat/completions', 'Local LLM inference (ALLaM, command-r7b, Qwen)'],
        ['AiCompletionService', 'POST {provider}/v1/chat/completions', 'Multi-provider AI abstraction'],
        ['EmbeddingService', 'POST /api/embed', 'Vector embeddings for semantic search'],
        ['KnowledgeRetrievalService', 'POST /search', 'RAG semantic search (optional)'],
        ['WebSearchService', 'GET /search?q=...', 'SearXNG meta-search engine'],
        ['FileExtractionService', 'POST /extract', 'PDF/document text extraction + OCR'],
        ['ContentModerationService', 'Internal', 'SDAIA-compliant content filtering'],
    ])

    doc.add_page_break()

    # ═══════════ 6. KNOWLEDGE BASE ═══════════
    add_h(doc, '6. Knowledge Base (RAG) \u2014 Academic Regulations', 1)

    add_p(doc, 'The chatbot uses Retrieval-Augmented Generation to answer regulation questions from official university documents.', size=11)

    add_h(doc, '6.1 Documents Indexed', 2)
    add_t(doc, ['Document', 'Type', 'Example Questions'], [
        ['Study & Exams Bylaw (\u0644\u0627\u0626\u062d\u0629 \u0627\u0644\u062f\u0631\u0627\u0633\u0629 \u0648\u0627\u0644\u0627\u062e\u062a\u0628\u0627\u0631\u0627\u062a)', 'PDF', '"What are First Honors requirements?" \u2014 "When does exam barring happen?"'],
        ['Student Guide (\u062f\u0644\u064a\u0644 \u0627\u0644\u0637\u0627\u0644\u0628)', 'PDF', '"Where is the counseling office?" \u2014 "How to apply for merit award?"'],
        ['Academic Calendar (\u0627\u0644\u062a\u0642\u0648\u064a\u0645 \u0627\u0644\u0623\u0643\u0627\u062f\u064a\u0645\u064a)', 'PDF', '"Last day for add/drop?" \u2014 "When do finals start?"'],
    ])

    add_h(doc, '6.2 RAG Pipeline', 2)
    bullet(doc, 'Storage: storage/app/advisor_knowledge/')
    bullet(doc, 'PDFs extracted to .pdf.extracted.txt via FileExtractionService')
    bullet(doc, 'Chunking: 1000 chars per chunk, 200 char overlap')
    bullet(doc, 'Ranking: BM25 keyword relevance scoring')
    bullet(doc, 'Top 15 chunks retrieved (up to 40,000 chars for Ollama, 12,000 for Groq)')
    bullet(doc, 'Page markers (--- Page X ---) preserved for citation')
    bullet(doc, 'Reference links auto-appended with PDF download + page number')

    add_h(doc, '6.3 Citation Format', 2)
    add_p(doc, 'Every regulation answer includes downloadable source references:', size=10)
    code = (
        '\U0001f4da **References:**\n'
        '- [\U0001f4c4 Study & Exams Bylaw.pdf](/api/v1/smart-advisor/download-reference?file=...pdf) - Page 15'
    )
    p = add_p(doc, code, size=9, color=RGBColor(64, 64, 64), after=8)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'))

    doc.add_page_break()

    # ═══════════ 7. AI MODEL ═══════════
    add_h(doc, '7. AI Model Strategy (Hybrid Selection)', 1)

    add_t(doc, ['Scenario', 'Model Used', 'Temperature', 'Why'], [
        ['Regulation/policy questions', 'Qwen2.5:14b', '0.1', 'Higher accuracy for detailed rules; lower creativity'],
        ['General academic questions', 'command-r7b-arabic', '0.3', 'Better Arabic fluency for conversational answers'],
        ['Fallback (short response)', 'Alternate model', '0.2', 'If primary model gives <50 chars, retries with other'],
        ['Cloud fallback (optional)', 'Groq llama-3.3-70b', '0.2', 'Larger model for complex regulation queries'],
    ])

    add_h(doc, '7.1 Context Window Management', 2)
    add_t(doc, ['Provider', 'Max Context', 'History', 'Knowledge Base'], [
        ['Ollama (local)', '40,000 chars', 'Up to 20 recent messages + older if space', 'Full KB chunks (up to 30,000 chars)'],
        ['Groq (cloud)', '12,000 chars', 'Limited recent messages', 'KB limited to 1,000 chars'],
    ])

    doc.add_page_break()

    # ═══════════ 8. CONVERSATIONS ═══════════
    add_h(doc, '8. Conversation Management & Storage', 1)

    add_h(doc, '8.1 Database Schema', 2)
    add_t(doc, ['Table', 'Key Fields', 'Purpose'], [
        ['advisor_conversations', 'id (UUID), user_id, title, status, student_context (JSON)', 'Conversation sessions, auto-titled from first message'],
        ['advisor_messages', 'id (UUID), conversation_id, role, content, tokens_used, response_time_ms, metadata (JSON)', 'Individual messages with quality metrics'],
        ['advisor_inquiries', 'id (UUID), user_id, student_id, category, subject, description, status, advisor_response, conversation_summary (JSON)', 'Human escalation requests'],
    ])

    add_h(doc, '8.2 Message Metadata (per message)', 2)
    add_t(doc, ['Field', 'Type', 'Example'], [
        ['intent', 'Array', '["absences", "courses"]'],
        ['apis_called', 'Array', '["/absences-with-details", "/student/courses"]'],
        ['quality_score', 'Integer 0\u2013100', '85'],
        ['quality_issues', 'Array', '["too_short", "no_data_reference"]'],
        ['model_used', 'String', 'command-r7b-arabic'],
        ['response_time_ms', 'Integer', '2340'],
        ['tokens_used', 'Integer', '512'],
    ])

    doc.add_page_break()

    # ═══════════ 9. ESCALATION ═══════════
    add_h(doc, '9. Human Escalation Flow', 1)

    add_t(doc, ['Step', 'Action', 'System Behavior'], [
        ['1', 'Student clicks "Escalate to Human Advisor"', 'Opens escalation form overlay'],
        ['2', 'Selects category from 5 options', 'Academic advising, Registration, Absence, Financial, General'],
        ['3', 'Fills subject + description', 'Free-text with Arabic support'],
        ['4', 'System creates inquiry record', 'Status: pending. Attaches last 20 messages (200 chars each) as context'],
        ['5', 'Admin sees inquiry in Filament dashboard', 'Can view conversation summary, student data'],
        ['6', 'Admin responds', 'Updates advisor_response field, status \u2192 resolved'],
        ['7', 'Student retrieves response', 'Via GET /api/v1/smart-advisor/inquiries'],
    ])

    doc.add_page_break()

    # ═══════════ 10. QUALITY ═══════════
    add_h(doc, '10. Response Quality Validation', 1)

    add_p(doc, 'Every AI response is scored 0\u2013100 using 7 automated checks:', size=11)
    add_t(doc, ['Check', 'Penalty', 'What It Catches'], [
        ['Empty response', '-100', 'Blank or null AI output'],
        ['Language mismatch', '-30', 'Arabic question, English/gibberish answer'],
        ['Too short (<50 chars)', '-20', 'Incomplete or cut-off responses'],
        ['Asking for existing data', '-30', 'AI says "please provide your GPA" when GPA is already in context'],
        ['No data reference', '-15', 'Student data provided but AI doesn\'t use any of it'],
        ['Error indicators', '-10 each', 'Contains "error", "can\'t", "unable to"'],
        ['Insufficient detail', '-10', 'Short answer for detailed question (e.g., regulation query < 100 chars)'],
    ])

    add_p(doc, 'If quality score < 60: system retries with alternate model before returning to user.', size=10, bold=True, color=BLUE)

    doc.add_page_break()

    # ═══════════ 11. SECURITY ═══════════
    add_h(doc, '11. Security & Content Moderation', 1)

    add_t(doc, ['Layer', 'Implementation', 'Details'], [
        ['Authentication', 'SAML 2.0 SSO via MyQU', 'Only authenticated university users'],
        ['Authorization', 'EnsureStudentAccess middleware', 'Blocks non-students; super admins exempt for testing'],
        ['Content Moderation', 'SDAIA-compliant filter', 'Checks all user input before AI processing'],
        ['Prompt Injection Prevention', 'System prompt hardening', 'AI cannot reveal system instructions or internal tools'],
        ['Data Privacy', 'Student-scoped data only', 'Each student sees only their own data; no cross-student access'],
        ['Rate Limiting', 'throttle:api middleware', 'Prevents abuse and excessive API calls'],
        ['Caching', '5-minute TTL on student data', 'Reduces API load; configurable'],
        ['Conversation Isolation', 'user_id verification', 'Cannot access another student\'s conversations'],
        ['Contradiction Detection', 'Post-processing check', 'Flags if AI says contradictory things (e.g., "allowed" and "not allowed")'],
    ])

    doc.add_page_break()

    # ═══════════ 12. CONFIG ═══════════
    add_h(doc, '12. Technical Configuration', 1)

    add_t(doc, ['Setting', 'Default', 'Description'], [
        ['SMART_ADVISOR_MAX_TOKENS', '8192', 'Max response tokens'],
        ['temperature', '0.3', 'AI creativity (lower = more factual)'],
        ['timeout', '300s', 'Max AI response wait time'],
        ['SMART_ADVISOR_MAX_HISTORY', '30', 'Max conversation messages to include'],
        ['SMART_ADVISOR_KB_PATH', 'advisor_knowledge', 'Knowledge base folder in storage/app/'],
        ['SMART_ADVISOR_MAX_KB_CHARS', '30000', 'Max knowledge context chars'],
        ['SMART_ADVISOR_CACHE_TTL', '300', 'Student data cache (seconds)'],
        ['SMART_ADVISOR_AI_INTENT', 'true', 'AI intent detection enabled'],
        ['SMART_ADVISOR_GROQ_ENABLED', 'false', 'Groq cloud fallback'],
        ['SMART_ADVISOR_SEMANTIC_SEARCH', 'false', 'Semantic search microservice'],
        ['BLACKBOARD_API_BASE_URL', 'http://127.0.0.1:8009/api/v2', 'University API base'],
        ['BLACKBOARD_TIMEOUT', '30', 'API call timeout (seconds)'],
    ])

    doc.add_page_break()

    # ═══════════ 13. FILE MAP ═══════════
    add_h(doc, '13. File Locations & Codebase Map', 1)

    add_t(doc, ['Component', 'File', 'Lines', 'Purpose'], [
        ['Web Controller', 'app/Http/Controllers/SmartAdvisorViewController.php', '33', 'UI entry point'],
        ['API Controller', 'app/Http/Controllers/Api/SmartAdvisorController.php', '369', 'REST API endpoints'],
        ['Core Service', 'app/Services/SmartAdvisorService.php', '3,087', 'All business logic'],
        ['University API', 'app/Services/UniversityApiClient.php', '500+', 'Blackboard + SIS calls'],
        ['Blackboard API', 'app/Services/BlackboardApiClient.php', '200+', 'LMS-specific calls'],
        ['Knowledge Retrieval', 'app/Services/KnowledgeRetrievalService.php', '350+', 'RAG search engine'],
        ['AI Completion', 'app/Services/AiCompletionService.php', '250+', 'Provider abstraction'],
        ['Chat View', 'resources/views/smart-advisor/index.blade.php', '1,299+', 'Chat UI (Blade)'],
        ['Conversation Model', 'app/Models/AdvisorConversation.php', '44', 'DB model'],
        ['Message Model', 'app/Models/AdvisorMessage.php', '30', 'DB model'],
        ['Inquiry Model', 'app/Models/AdvisorInquiry.php', '48', 'Escalation model'],
        ['Routes (web)', 'routes/web.php (line 118)', '1', '/smart-advisor'],
        ['Routes (API)', 'routes/api.php (lines 111\u2013125)', '7', '/api/v1/smart-advisor/*'],
        ['Config', 'config/quai.php (lines 202\u2013245)', '44', 'All settings'],
        ['Middleware', 'app/Http/Middleware/EnsureStudentAccess.php', '30+', 'Student-only gate'],
    ])

    doc.add_page_break()

    # ═══════════ 14. LIMITATIONS ═══════════
    add_h(doc, '14. Current Limitations & Known Gaps', 1)

    add_t(doc, ['#', 'Limitation', 'Impact', 'Planned Resolution'], [
        ['1', 'No proactive monitoring (reactive only)', 'System only responds when student asks', 'QMentor V2: 24/7 agent loop'],
        ['2', 'No risk scoring or prediction', 'Cannot flag at-risk students automatically', 'QMentor V2: 66-indicator predictive engine'],
        ['3', 'No advisor dashboard', 'Advisors don\'t see student risk data', 'QMentor V2: Full Digital Twin dashboard'],
        ['4', 'No multi-channel alerts', 'Only web chat available', 'QMentor V2: SMS, email, WhatsApp, push'],
        ['5', 'LMS engagement data gaps', 'Login frequency, time-on-platform not available', 'Blackboard Data integration'],
        ['6', 'Knowledge base is manual', 'PDFs must be manually uploaded/updated', 'Auto-sync pipeline in V2'],
        ['7', 'Context window limits (Groq)', 'Cloud fallback limited to 12K chars', 'AWS Bedrock with larger context'],
        ['8', 'No mobile app', 'Web-only access', 'QMentor V2: Flutter iOS/Android'],
        ['9', 'No Teams scheduling', 'Cannot book advisor meetings from chat', 'QMentor V2: Graph API integration'],
        ['10', 'No study plan builder', 'Cannot recommend electives or build schedules', 'QMentor V2: GPO module'],
        ['11', 'Hallucination not fully eliminated', 'Prompt mitigation + contradiction detection, but gaps remain', 'Grounded generation + output validation'],
        ['12', 'No voice input', 'Text-only input', 'QMentor V2: Speech-to-text (Whisper)'],
    ])

    doc.add_page_break()

    # ═══════════ 15. ROADMAP ═══════════
    add_h(doc, '15. Roadmap: Smart Advisor V1 \u2192 QMentor V2', 1)

    blue_box(doc, 'Smart Advisor V1 is the foundation. QMentor V2 builds on top of it with proactive AI, predictive analytics, and multi-channel reach.')

    add_t(doc, ['Capability', 'V1 (Current)', 'V2 (QMentor)'], [
        ['Architecture', 'Reactive chatbot (student asks)', 'Agentic AI (proactive 24/7 monitoring)'],
        ['Risk Detection', 'None', '66 indicators, 9 categories, 0\u2013100 scoring'],
        ['Advisor Tools', 'None', 'Full Digital Twin dashboard + intervention plans'],
        ['Channels', 'Web chat only', 'Web + Mobile + SMS + Email + WhatsApp + Teams'],
        ['Alerts', 'None', 'Multi-channel with auto-escalation'],
        ['Academic Planning', 'None', 'AI Study Plan Builder + Graduation Optimizer'],
        ['Meeting Scheduling', 'None', 'Microsoft Teams integration via Graph API'],
        ['Student Data', 'Real-time via API (on-demand)', 'Continuous monitoring + Digital Twin'],
        ['Knowledge Base', 'BM25 keyword search', 'Semantic search + expanded knowledge'],
        ['AI Models', 'Self-hosted Ollama (7\u201314b)', 'AWS Bedrock (Claude/Llama) + self-hosted'],
        ['Hosting', 'On-premises / local', 'AWS Cloud-Native with GPU servers'],
        ['Compliance', 'Basic SDAIA content moderation', 'Full SDAIA ethics (37 reqs) + NCA + PDPL'],
        ['Scale', 'Single instance', '70,000+ students with auto-scaling'],
    ])

    # ── FOOTER ──
    doc.add_paragraph('')
    add_p(doc, '\u2501' * 50, size=8, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, 'Prepared by: Digital Transformation Team \u2014 Deanship of E-Learning & IT, Qassim University',
          size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    path = os.path.join(os.path.dirname(__file__), 'SmartAdvisor_V1_Technical.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    main()

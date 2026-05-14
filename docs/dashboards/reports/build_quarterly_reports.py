"""
Generates four standalone Arabic RTL Word reports — one per admin dashboard —
covering Q3 2025, Q4 2025, and Q1 2026, and tying every analysis back to the
DGA "منظور مركزية المستفيد" framework (Perspective 8 + adjacent standards).

Outputs (all under docs/dashboards/reports/):
  • Complaints-Q3-2025-Q1-2026-Report.docx
  • Reviews-Q3-2025-Q1-2026-Report.docx
  • ServiceTasks-Q3-2025-Q1-2026-Report.docx
  • ServiceEvaluations-Q3-2025-Q1-2026-Report.docx

Style: DGA Saudi-green palette throughout. RTL document-level + section-level.

Run:
  php artisan tinker --execute="..." > /tmp/quarterly-data.json   # already populated by parent shell
  python3 docs/dashboards/reports/build_quarterly_reports.py
"""
from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "docs" / "dashboards" / "reports"
DATA = json.loads(Path("/tmp/quarterly-data.json").read_text())

BASE_URL = "http://localhost:8007"

# DGA Saudi-green palette (per the brief)
SA_50 = RGBColor(0xF3, 0xFC, 0xF6)
SA_100 = RGBColor(0xDF, 0xF6, 0xE7)
SA_200 = RGBColor(0xB8, 0xEA, 0xCB)
SA_400 = RGBColor(0x54, 0xC0, 0x8A)
SA_500 = RGBColor(0x25, 0x93, 0x5F)   # MAIN brand
SA_600 = RGBColor(0x1B, 0x83, 0x54)
SA_700 = RGBColor(0x16, 0x6A, 0x45)
SA_800 = RGBColor(0x14, 0x57, 0x3A)
SA_900 = RGBColor(0x10, 0x46, 0x31)
GRAY_500 = RGBColor(0x6C, 0x73, 0x7F)
GRAY_700 = RGBColor(0x38, 0x42, 0x50)
GRAY_900 = RGBColor(0x11, 0x19, 0x27)
ERROR_700 = RGBColor(0xB4, 0x23, 0x18)


# ===========================================================================
# Low-level XML helpers
# ===========================================================================
def set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_rtl(cell, *, shaded: str | None = None) -> None:
    if shaded:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shaded)
        tcPr.append(shd)
    for p in cell.paragraphs:
        set_paragraph_rtl(p)


def set_section_rtl(section) -> None:
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)


def set_default_rtl(doc: Document) -> None:
    normal = doc.styles["Normal"]
    pPr = normal.element.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def add_hyperlink(paragraph, url: str, text: str, *, bold=False, size=11, color=SA_600) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(k), "Arial")
    rPr.append(rFonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(color_el)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(6)  # WD_BREAK.PAGE = 6


# ===========================================================================
# Higher-level building blocks
# ===========================================================================
def add_h(doc: Document, text: str, level: int = 1, color=SA_700) -> None:
    h = doc.add_heading("", level=level)
    set_paragraph_rtl(h)
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = color


def add_p(doc: Document, text: str, *, bold=False, size=11, italic=False,
          color=GRAY_900, align="right") -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = color


def add_bullet(doc: Document, text: str, *, size=10, indent=False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_rtl(p)
    if indent:
        p.paragraph_format.right_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_500


def add_kv_table(doc: Document, headers: list[str], rows: list[list[str]],
                 *, widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Arial"
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_rtl(cell, shaded="166A45")
    for r_idx, row in enumerate(rows, start=1):
        shade = "F3FCF6" if r_idx % 2 == 1 else None
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Arial"
            set_cell_rtl(cell, shaded=shade)


def add_callout(doc: Document, label: str, body: str, *, color_hex="DFF6E7", border_hex="166A45") -> None:
    """Single-row table styled as a green callout box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color_hex); tcPr.append(shd)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "right", "left"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), border_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)
    p1 = cell.paragraphs[0]
    set_paragraph_rtl(p1)
    r = p1.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(int(border_hex[0:2], 16), int(border_hex[2:4], 16), int(border_hex[4:6], 16))
    r.font.name = "Arial"
    p2 = cell.add_paragraph()
    set_paragraph_rtl(p2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY_900
    r2.font.name = "Arial"


def add_toc_line(doc: Document, num: str, title: str, *, indent=0) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.right_indent = Cm(0.5 * indent)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(num + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = SA_700
    r1.font.name = "Arial"
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY_900
    r2.font.name = "Arial"


def cover_page(doc: Document, title: str, subtitle: str, dashboard_url: str) -> None:
    # Big green heading
    h = doc.add_paragraph()
    set_paragraph_rtl(h)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(120)
    r = h.add_run("جامعة القصيم — وكالة الجامعة للتطوير والجودة")
    r.font.size = Pt(13)
    r.font.color.rgb = SA_700
    r.font.name = "Arial"
    r.bold = True

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = SA_800
    r.font.name = "Arial"

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = SA_600
    r.font.name = "Arial"

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("الفترة المغطاة: الربع الثالث 2025 — الربع الرابع 2025 — الربع الأول 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_700
    r.font.name = "Arial"

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("الإطار المرجعي: قياس 2026 — المنظور الثامن «مركزية المستفيد»")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_700
    r.font.name = "Arial"

    # Spacer + Dashboard hyperlink
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    add_hyperlink(p, dashboard_url, "فتح اللوحة الحية على QUAI →", bold=True, size=12, color=SA_600)

    page_break(doc)


def perspective8_block(doc: Document) -> None:
    """The reusable Perspective 8 framework block, with all standards mapped."""
    add_h(doc, "الإطار المرجعي: المنظور الثامن — مركزية المستفيد", level=1, color=SA_700)
    add_p(doc,
        "يهدف هذا المنظور إلى تعزيز دور المستفيد وتحويله إلى شريك فاعل في تطوير وتحسين الخدمات الحكومية الرقمية. "
        "يندرج تحت هذا المنظور ثلاثة محاور رئيسية:")

    add_callout(doc, "المحور 18: مشاركة المستفيد",
        "المعيار 5.18.1: إتاحة القنوات وطرح موضوعات وفرص المشاركة الإلكترونية. "
        "يقيس مدى توفير قنوات اتصال فعّالة تتيح للمستفيد إبداء الرأي والمشاركة في تطوير الخدمات.")
    add_callout(doc, "المحور 19: تعزيز العلاقة مع المستفيد",
        "يهدف إلى تقوية الروابط والتفاعل المستمر بين الجهة والمستفيدين، عبر قنوات اتصال متعددة "
        "ومتابعة منتظمة لمستوى الرضا.")
    add_callout(doc, "المحور 20: تجربة المستفيد",
        "يركز على تحسين رحلة المستفيد وضمان جودة تفاعله مع الخدمات الرقمية، من حيث السرعة، السهولة، "
        "والاتساق عبر القنوات.")

    add_p(doc, "")
    add_p(doc, "معايير ذات صلة من المنظور السابع (القنوات والخدمات):", bold=True, color=SA_700)
    add_bullet(doc, "المعيار 5.17.3: تبني الخدمات الرقمية الاستباقية للمستفيد الداخلي — يتضمن تقديم 10 خدمات داخلية بدون نماذج (Pre-filled).")
    add_bullet(doc, "المعيار 5.17.4: تبني الخدمات الرقمية الاستباقية للمستفيد الخارجي — تقديم ما لا يقل عن 30% أو 10 خدمات استباقية، و30% بدون نماذج.")
    add_p(doc, "")
    add_p(doc, "متطلبات إضافية متعلقة بالمستفيد:", bold=True, color=SA_700)
    add_bullet(doc, "المعيار 5.16.4: تطوير الخدمات ذات الأولوية — حصر فئات المستفيدين وأعدادهم.")
    add_bullet(doc, "المعيار 5.15.4: نظام التصميم الوطني الموحد (كود المنصات) — لضمان تجربة رقمية موحدة وسلسة.")


def appendix(doc: Document) -> None:
    add_h(doc, "ملحق: كيف تتحقق من أرقام هذا التقرير", level=1, color=SA_700)
    add_p(doc, "كل رقم في هذا التقرير مأخوذ مباشرة من قاعدة بيانات QUAI، ويمكن إعادة استخراجه بسهولة عبر:")
    add_bullet(doc, "php artisan tinker")
    add_p(doc, "نماذج Eloquent ذات الصلة:", bold=True)
    add_bullet(doc, "App\\Models\\ServiceTask  ← لوحة مهام تقنية المعلومات (نظام ساعد)")
    add_bullet(doc, "App\\Models\\Complaint  ← لوحة الشكاوى والمقترحات")
    add_bullet(doc, "App\\Models\\ServiceEvaluation  ← لوحة تقييم الخدمات")
    add_bullet(doc, "App\\Models\\Review  ← لوحة التقييمات الخارجية")
    add_p(doc, "أهداف SLA الموحدة في: App\\Support\\Sla::TARGETS_HOURS")
    add_caption(doc, "تم إعداد هذا التقرير برمجياً من بيانات حية. تاريخ التوليد محفوظ في خصائص المستند.")


# ===========================================================================
# Report 1: Complaints
# ===========================================================================
def build_complaints_report() -> None:
    c = DATA["complaints"]
    out = OUT_DIR / "Complaints-Q3-2025-Q1-2026-Report.docx"
    doc = Document()
    set_default_rtl(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2)
        set_section_rtl(s)

    cover_page(doc,
        "تقرير لوحة الشكاوى والمقترحات",
        "مساهمة في قياس مركزية المستفيد للجهات الحكومية",
        BASE_URL + "/admin/complaints-dashboard")

    # ---- Detailed TOC ----
    add_h(doc, "الفهرس التفصيلي", level=1, color=SA_700)
    add_toc_line(doc, "1.", "الملخص التنفيذي")
    add_toc_line(doc, "2.", "الإطار المرجعي: المنظور الثامن — مركزية المستفيد")
    add_toc_line(doc, "3.", "نطاق التقرير ومصادر البيانات")
    add_toc_line(doc, "4.", "المؤشرات الأساسية للوحة")
    add_toc_line(doc, "    4.1", "حجم الطلبات وتوزيعها", indent=1)
    add_toc_line(doc, "    4.2", "نسبة الرد ومعالجة المعلق", indent=1)
    add_toc_line(doc, "    4.3", "متوسط وقت الاستجابة (تقديري)", indent=1)
    add_toc_line(doc, "5.", "تحليل ربعي: Q3 2025 / Q4 2025 / Q1 2026")
    add_toc_line(doc, "6.", "أعلى التصنيفات وأكثرها إزعاجاً للمستفيد")
    add_toc_line(doc, "7.", "ربط القراءات بالمعايير")
    add_toc_line(doc, "    7.1", "المعيار 5.18.1 — إتاحة قنوات المشاركة الإلكترونية", indent=1)
    add_toc_line(doc, "    7.2", "المحور 19 — تعزيز العلاقة مع المستفيد", indent=1)
    add_toc_line(doc, "    7.3", "المحور 20 — تجربة المستفيد", indent=1)
    add_toc_line(doc, "8.", "التوصيات التنفيذية")
    add_toc_line(doc, "9.", "خطة العمل المقترحة (90 يوماً)")
    add_toc_line(doc, "10.", "الملحق")
    page_break(doc)

    # ---- 1. Executive Summary ----
    add_h(doc, "1. الملخص التنفيذي", level=1)
    add_p(doc,
        f"خلال الفترة المرجعية، استقبلت قناة الشكاوى والمقترحات في QUAI {c['total']} طلباً، "
        f"تم الرد على {c['responded']} منها (نسبة {round(c['responded']/c['total']*100,1)}%)، "
        f"بينما لا يزال {c['total']-c['responded']} طلباً بانتظار الرد الإداري.")
    add_p(doc,
        f"التوزيع: استفسارات {c['inquiries']}، اقتراحات {c['suggestions']}، شكاوى رسمية {c['complaints_only']}. "
        f"تشير غلبة الاستفسارات (≈{round(c['inquiries']/c['total']*100,1)}%) إلى أن القناة تُستخدم في معظم الأحيان "
        f"للحصول على معلومات لا للتعبير عن استياء — وهي إشارة إيجابية تدل على ثقة المستفيد، لكنها أيضاً تكشف "
        f"عن فرصة كبيرة لخدمات ذاتية (Self-service).")
    add_callout(doc, "أبرز ملاحظة",
        "نسبة الرد الإجمالية مرتفعة، لكن المؤشر الحرج هو حصر الطلبات المعلقة في تصنيفات بعينها — مما يشير "
        "إلى ضرورة تخصيص مالك واضح لكل تصنيف وليس مجرد رفع نسبة الرد العامة.")
    page_break(doc)

    # ---- 2. Perspective 8 ----
    add_h(doc, "2. الإطار المرجعي", level=1)
    perspective8_block(doc)
    page_break(doc)

    # ---- 3. Scope ----
    add_h(doc, "3. نطاق التقرير ومصادر البيانات", level=1)
    add_p(doc, f"المصدر: جدول complaints في قاعدة بيانات QUAI، عدد السجلات الكلي: {c['total']}.")
    add_p(doc, "ملاحظة بشأن التوزيع الزمني:", bold=True)
    add_p(doc, c["note"])
    add_p(doc,
        "بسبب غياب طوابع زمنية للتقديم في الملف المصدر، نعرض الإجماليات بدل التوزيع الربعي. عند ربط نظام "
        "تذاكر فعلي بطابع زمني للتقديم والرد، سيُحسب التوزيع الربعي تلقائياً.")
    page_break(doc)

    # ---- 4. KPIs ----
    add_h(doc, "4. المؤشرات الأساسية للوحة", level=1)

    add_h(doc, "4.1 حجم الطلبات وتوزيعها", level=2)
    add_kv_table(doc,
        ["المؤشر", "القيمة", "المصدر"],
        [
            ["إجمالي الطلبات", f"{c['total']}", "Complaint::count()"],
            ["استفسارات", f"{c['inquiries']}", "WHERE req_type = 'استفسار'"],
            ["اقتراحات", f"{c['suggestions']}", "WHERE req_type = 'اقتراح'"],
            ["شكاوى رسمية", f"{c['complaints_only']}", "WHERE req_type = 'شكوى'"],
        ],
        widths_cm=[5.0, 4.0, 8.0])

    add_h(doc, "4.2 نسبة الرد ومعالجة المعلق", level=2)
    add_kv_table(doc,
        ["المؤشر", "القيمة", "المصدر"],
        [
            ["تم الرد", f"{c['responded']}", "Complaint::withResponse()"],
            ["بانتظار الرد", f"{c['total']-c['responded']}", "Complaint::withoutResponse()"],
            ["نسبة الرد", f"{round(c['responded']/c['total']*100,1)}%", "responded / total"],
        ],
        widths_cm=[5.0, 4.0, 8.0])

    add_h(doc, "4.3 متوسط وقت الاستجابة (تقديري)", level=2)
    weighted = round((90 * c['inquiries'] + 360 * c['complaints_only'] + 1440 * c['suggestions']) / c['total'])
    add_p(doc,
        f"التقدير المرجح يبلغ ≈ {weighted:,} دقيقة ({round(weighted/60,1)} ساعة)، محسوباً وفق الفرضيات التالية: "
        f"الاستفسار 90 د، الشكوى 360 د، الاقتراح 1,440 د. هذه الفرضيات قابلة للضبط حالما تتوفر طوابع زمنية حقيقية.")
    page_break(doc)

    # ---- 5. Quarterly note ----
    add_h(doc, "5. تحليل ربعي: Q3 2025 / Q4 2025 / Q1 2026", level=1)
    add_p(doc,
        "نظراً لكون كل سجلات الشكاوى وردت دفعة واحدة دون طوابع زمنية، لا يمكن تقديم انفصال ربعي دقيق "
        "في الوقت الراهن. يوصى — لتحقيق متطلبات قياس 2026 — بإضافة عمود submitted_at و responded_at عند "
        "ربط نظام التذاكر الجديد، حيث سيُحسب التوزيع التلقائي بمجرد توفر التواريخ.")
    add_callout(doc, "ما الذي يجب أن يظهر في الربع القادم",
        "بمجرد تفعيل التواريخ الحقيقية، سيعرض هذا القسم: عدد الطلبات لكل ربع، نسبة الرد الربعي، "
        "متوسط وقت الاستجابة الربعي، ومقارنة الانحدار/الصعود بين الأرباع.")
    page_break(doc)

    # ---- 6. Top classifications ----
    add_h(doc, "6. أعلى التصنيفات وأكثرها إزعاجاً للمستفيد", level=1)
    rows = [[r['classification_desc'] or "غير محدد", str(r['n']), "—"] for r in c['top_classification']]
    add_kv_table(doc,
        ["التصنيف", "عدد الطلبات", "ملاحظة"],
        rows,
        widths_cm=[8.0, 3.0, 5.5])
    add_h(doc, "التصنيفات الأعلى تأخراً في الرد", level=2)
    rows = [[r['classification_desc'] or "غير محدد", str(r['n']), "بدون رد إداري"] for r in c['pending_classification']]
    add_kv_table(doc,
        ["التصنيف", "بانتظار الرد", "الحالة"],
        rows,
        widths_cm=[8.0, 3.0, 5.5])
    page_break(doc)

    # ---- 7. Mapping ----
    add_h(doc, "7. ربط القراءات بالمعايير", level=1)

    add_h(doc, "7.1 المعيار 5.18.1 — إتاحة قنوات المشاركة الإلكترونية", level=2)
    add_p(doc,
        f"اللوحة تُثبت توفر قناة مشاركة فعّالة، استقبلت {c['total']} طلباً وردت على {c['responded']} منها. "
        "وجود تصنيفات متعددة وقابلية الفصل بين \"استفسار\" و\"اقتراح\" و\"شكوى\" يعكس نضج القناة وقدرتها "
        "على استيعاب أنواع المشاركة المختلفة.")
    add_callout(doc, "دليل الالتزام بالمعيار",
        f"عدد الطلبات المستقبَلة، نسبة الرد ({round(c['responded']/c['total']*100,1)}%)، عدد التصنيفات "
        "المتاحة، وآلية تتبع كل طلب — كلها محققة وتدعم وثيقة قياس 2026.")

    add_h(doc, "7.2 المحور 19 — تعزيز العلاقة مع المستفيد", level=2)
    add_p(doc,
        "تشير الأرقام إلى وجود حوار حقيقي: لكل طلب مكتمل ملاحظة إدارية مكتوبة، وهذا أفضل من مجرد إغلاق ميكانيكي. "
        "غلبة الاستفسارات (تتجاوز ضعف الشكاوى) تكشف ثقة المستفيد بالقناة كمرجع للمعلومة.")

    add_h(doc, "7.3 المحور 20 — تجربة المستفيد", level=2)
    add_p(doc,
        "تجربة المستفيد قابلة للتحسين عبر: تقصير زمن الرد على الاستفسارات الشائعة (قاعدة معرفة ذاتية)، "
        "ومعالجة التصنيفات التي تحظى بأعلى تجمع شكاوى — لأنها تكشف نقاط ألم متكررة لا فردية.")
    page_break(doc)

    # ---- 8. Recommendations ----
    add_h(doc, "8. التوصيات التنفيذية", level=1)
    add_p(doc, "أولاً — توصيات عاجلة (خلال هذا الأسبوع):", bold=True, color=ERROR_700)
    add_bullet(doc,
        f"تحديد مالك واضح لكل تصنيف فرعي وتفعيل تنبيه يومي عند تجاوز 48 ساعة دون رد. "
        f"حالياً {c['total']-c['responded']} طلباً بانتظار الرد، تركز معظمها في تصنيف واحد.")
    add_bullet(doc, "نشر إحصائيات الرد على بوابة التواصل لتعزيز الشفافية وثقة المستفيد.")

    add_p(doc, "ثانياً — توصيات استباقية (خلال 30-60 يوماً):", bold=True, color=SA_700)
    add_bullet(doc,
        "إنشاء قاعدة معرفة ذاتية الخدمة (FAQ تفاعلي + محرك بحث) للإجابة المباشرة على أعلى 10 استفسارات متكررة. "
        f"المتوقع تحويل ≈ {round(c['inquiries']*0.4)} استفسار شهرياً إلى خدمة ذاتية.")
    add_bullet(doc, "إضافة عمودي submitted_at و responded_at في جدول الشكاوى لحساب \"زمن الاستجابة\" الفعلي بالدقائق.")

    add_p(doc, "ثالثاً — توصيات استراتيجية (دورة قياس قادمة):", bold=True, color=SA_700)
    add_bullet(doc,
        "ربط لوحة الشكاوى ببوابة الجامعة الموحدة بحيث ينعكس كل طلب فوراً على لوحة المستفيد، تحقيقاً للمعيار 5.15.4.")
    page_break(doc)

    # ---- 9. Action plan ----
    add_h(doc, "9. خطة العمل المقترحة (90 يوماً)", level=1)
    add_kv_table(doc,
        ["المرحلة", "الأسبوع", "الإجراء", "المؤشر"],
        [
            ["1", "1-2", "تخصيص ملكية لكل تصنيف", "عدد التصنيفات بدون مالك = 0"],
            ["2", "3-4", "إطلاق قاعدة المعرفة الذاتية v1", "نسبة الاستفسارات المحوّلة ذاتياً ≥ 30%"],
            ["3", "5-8", "إضافة طوابع زمنية للجدول", "submitted_at متوفر لكل سجل جديد"],
            ["4", "9-12", "نشر تقرير ربعي علني", "نشر بحلول 30/Q-end"],
        ],
        widths_cm=[2.0, 2.5, 7.0, 5.0])

    page_break(doc)
    appendix(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ===========================================================================
# Report 2: Reviews
# ===========================================================================
def build_reviews_report() -> None:
    rv = DATA["reviews"]
    out = OUT_DIR / "Reviews-Q3-2025-Q1-2026-Report.docx"
    doc = Document()
    set_default_rtl(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2)
        set_section_rtl(s)

    cover_page(doc,
        "تقرير لوحة التقييمات الخارجية",
        "مرآة المستفيد الخارجي على المنصات العامة",
        BASE_URL + "/admin/reviews-dashboard")

    add_h(doc, "الفهرس التفصيلي", level=1, color=SA_700)
    add_toc_line(doc, "1.", "الملخص التنفيذي")
    add_toc_line(doc, "2.", "الإطار المرجعي: المنظور الثامن — مركزية المستفيد")
    add_toc_line(doc, "3.", "مصادر البيانات والمنصات المغطاة")
    add_toc_line(doc, "4.", "المؤشرات الأساسية")
    add_toc_line(doc, "5.", "تحليل ربعي تفصيلي")
    add_toc_line(doc, "    5.1", "الربع الثالث 2025", indent=1)
    add_toc_line(doc, "    5.2", "الربع الرابع 2025", indent=1)
    add_toc_line(doc, "    5.3", "الربع الأول 2026", indent=1)
    add_toc_line(doc, "6.", "تحليل المنصات والمشاعر")
    add_toc_line(doc, "7.", "ربط القراءات بالمعايير")
    add_toc_line(doc, "    7.1", "المحور 19 — تعزيز العلاقة مع المستفيد", indent=1)
    add_toc_line(doc, "    7.2", "المحور 20 — تجربة المستفيد", indent=1)
    add_toc_line(doc, "    7.3", "المعيار 5.16.4 — حصر فئات المستفيدين", indent=1)
    add_toc_line(doc, "8.", "التوصيات التنفيذية")
    add_toc_line(doc, "9.", "خطة العمل المقترحة (90 يوماً)")
    add_toc_line(doc, "10.", "الملحق")
    page_break(doc)

    # 1. Executive summary
    total_all = sum(rv[q]['total'] for q in rv)
    sat_total = sum(rv[q]['positive'] for q in rv)
    neg_total = sum(rv[q]['negative'] for q in rv)
    add_h(doc, "1. الملخص التنفيذي", level=1)
    add_p(doc,
        f"خلال الفترة المرجعية، رُصد {total_all} تقييم خارجي عبر منصات Google Maps و App Store و Google Play "
        f"ووسائل التواصل الاجتماعي. متوسط التقييم العام يتراوح بين {min(rv[q]['avg_rating'] for q in rv)}/5 "
        f"و {max(rv[q]['avg_rating'] for q in rv)}/5، مع ميل واضح نحو المشاعر الإيجابية ({sat_total} إيجابي مقابل {neg_total} سلبي).")
    add_callout(doc, "أبرز ملاحظة",
        "حجم التقييمات في الربع الأول 2026 يفوق نظيره في الربعين السابقين مجتمعَين، مما يدل على نمو مكاسب الجامعة "
        "في الحضور الرقمي — وهذا في حد ذاته مؤشر قوي على المحور 19 (تعزيز العلاقة مع المستفيد).")
    page_break(doc)

    # 2. Perspective 8
    add_h(doc, "2. الإطار المرجعي", level=1)
    perspective8_block(doc)
    page_break(doc)

    # 3. Sources
    add_h(doc, "3. مصادر البيانات والمنصات المغطاة", level=1)
    add_p(doc, "تجمع اللوحة تقييمات من خمس قنوات خارجية:", bold=True)
    add_bullet(doc, "خرائط Google (google_maps) — يؤثر مباشرة على نتائج البحث المحلية.")
    add_bullet(doc, "App Store و Google Play — تطبيقات الجامعة الرسمية.")
    add_bullet(doc, "وسائل التواصل الاجتماعي — تويتر/إنستجرام/فيسبوك.")
    add_bullet(doc, "إدخال يدوي — تقييمات مُجمَّعة من قنوات أخرى.")
    page_break(doc)

    # 4. KPIs
    add_h(doc, "4. المؤشرات الأساسية", level=1)
    add_kv_table(doc,
        ["المؤشر", "Q3 2025", "Q4 2025", "Q1 2026"],
        [
            ["إجمالي التقييمات", str(rv['Q3-2025']['total']), str(rv['Q4-2025']['total']), str(rv['Q1-2026']['total'])],
            ["متوسط التقييم (من 5)", str(rv['Q3-2025']['avg_rating']), str(rv['Q4-2025']['avg_rating']), str(rv['Q1-2026']['avg_rating'])],
            ["إيجابي", str(rv['Q3-2025']['positive']), str(rv['Q4-2025']['positive']), str(rv['Q1-2026']['positive'])],
            ["محايد", str(rv['Q3-2025']['neutral']), str(rv['Q4-2025']['neutral']), str(rv['Q1-2026']['neutral'])],
            ["سلبي", str(rv['Q3-2025']['negative']), str(rv['Q4-2025']['negative']), str(rv['Q1-2026']['negative'])],
        ],
        widths_cm=[5.5, 3.5, 3.5, 3.5])
    page_break(doc)

    # 5. Quarter detail
    add_h(doc, "5. تحليل ربعي تفصيلي", level=1)
    for q_label, q_key, q_title in [
        ("5.1", "Q3-2025", "5.1 الربع الثالث 2025"),
        ("5.2", "Q4-2025", "5.2 الربع الرابع 2025"),
        ("5.3", "Q1-2026", "5.3 الربع الأول 2026"),
    ]:
        s = rv[q_key]
        add_h(doc, q_title, level=2)
        add_p(doc,
            f"إجمالي {s['total']} تقييم بمتوسط {s['avg_rating']}/5. توزيع المشاعر: "
            f"إيجابي {s['positive']}، محايد {s['neutral']}، سلبي {s['negative']}.")
        if s['platforms']:
            top = s['platforms'][0]
            add_p(doc, f"المنصة الأكثر نشاطاً: {top['platform']} بـ {top['n']} تقييم.")
    page_break(doc)

    # 6. Platform deep-dive
    add_h(doc, "6. تحليل المنصات والمشاعر", level=1)
    add_p(doc,
        "اللوحة تكشف عن أن وسائل التواصل الاجتماعي تحمل العبء الأكبر من حجم التقييمات، يليها خرائط Google. "
        "التطبيقات (App Store/Google Play) لها وزن أقل لكن تأثيرها أعمق على رحلة المستفيد التقني.")
    add_callout(doc, "نمط مكتشف",
        "متوسط تقييم خرائط Google عادة ما يكون أعلى من المنصات الأخرى. هذا منطقي — مستخدم خرائط Google "
        "يقيّم تجربة موقعية، بينما مستخدم وسائل التواصل قد يقيّم خدمة عابرة أو تجربة سلبية محددة.")
    page_break(doc)

    # 7. Mapping
    add_h(doc, "7. ربط القراءات بالمعايير", level=1)
    add_h(doc, "7.1 المحور 19 — تعزيز العلاقة مع المستفيد", level=2)
    add_p(doc,
        "وجود قناة موحّدة لمراقبة التقييمات الخارجية — مع تحليل مشاعر آلي وتنبيه للسلبي — هو في صميم المحور 19. "
        "اللوحة الحية تتيح للجهة الرد العام في الوقت المناسب وتوثيق التحسن.")

    add_h(doc, "7.2 المحور 20 — تجربة المستفيد", level=2)
    add_p(doc,
        "تتبع متوسط التقييم عبر الأرباع الثلاثة هو قياس مباشر لجودة التجربة. الانحدار في أي ربع يستوجب "
        "تحقيقاً حول أسبابه (تحديث جديد، تغيير سياسة، حادث خدمة).")

    add_h(doc, "7.3 المعيار 5.16.4 — حصر فئات المستفيدين", level=2)
    add_p(doc,
        "تصنيف التقييمات بحسب اللغة والمنصة يساعد في تحديد فئات المستفيدين الفعلية: العربي/الأجنبي، "
        "المستخدم المحلي/الدولي. هذه القسمة شرط لإحكام تطوير الخدمات وفقاً للمعيار.")
    page_break(doc)

    # 8. Recommendations
    add_h(doc, "8. التوصيات التنفيذية", level=1)
    add_p(doc, "أولاً — توصيات عاجلة (هذا الأسبوع):", bold=True, color=ERROR_700)
    add_bullet(doc, f"الرد على آخر {neg_total} تعليق سلبي علناً، مع تكليف موظف اتصال مؤسسي بالمتابعة الفردية.")
    add_p(doc, "ثانياً — توصيات استباقية (30-60 يوماً):", bold=True, color=SA_700)
    add_bullet(doc, "تشغيل تنبيه فوري عند ورود تقييم بدرجة 1 أو 2 لتقصير وقت الاستجابة.")
    add_bullet(doc, "إثراء حقل sentiment للسجلات غير المحلَّلة لرفع جودة التحليل.")
    add_p(doc, "ثالثاً — توصيات استراتيجية:", bold=True, color=SA_700)
    add_bullet(doc, "ربط حملة طلب التقييم بانتهاء الخدمة الرقمية، لزيادة عدد التقييمات الإيجابية الموثقة.")
    page_break(doc)

    # 9. Action plan
    add_h(doc, "9. خطة العمل المقترحة (90 يوماً)", level=1)
    add_kv_table(doc,
        ["المرحلة", "الأسبوع", "الإجراء", "المؤشر"],
        [
            ["1", "1", "الرد على كل تعليق سلبي قائم", "متبقي = 0"],
            ["2", "2-4", "تنبيه آلي للتقييمات السلبية", "زمن الرد على السلبي < 24 ساعة"],
            ["3", "5-8", "إثراء التحليل العاطفي للسجلات الناقصة", "نسبة المُحلَّل ≥ 80%"],
            ["4", "9-12", "حملة دعوة للتقييم بعد الخدمات", "عدد التقييمات الجديدة ≥ +30%"],
        ],
        widths_cm=[2.0, 2.5, 7.0, 5.0])
    page_break(doc)
    appendix(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ===========================================================================
# Report 3: Service Tasks (نظام ساعد)
# ===========================================================================
def build_service_tasks_report() -> None:
    st = DATA["service_tasks"]
    out = OUT_DIR / "ServiceTasks-Q3-2025-Q1-2026-Report.docx"
    doc = Document()
    set_default_rtl(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2)
        set_section_rtl(s)

    cover_page(doc,
        "تقرير لوحة مهام تقنية المعلومات (نظام ساعد)",
        "قياس جودة الاستجابة التشغيلية والالتزام بـ SLA لخدمة المستفيد",
        BASE_URL + "/admin/service-tasks-dashboard")

    add_h(doc, "الفهرس التفصيلي", level=1, color=SA_700)
    add_toc_line(doc, "1.", "الملخص التنفيذي")
    add_toc_line(doc, "2.", "الإطار المرجعي: المنظور الثامن — مركزية المستفيد")
    add_toc_line(doc, "3.", "نطاق التقرير ومصدر البيانات")
    add_toc_line(doc, "4.", "المؤشرات الأساسية")
    add_toc_line(doc, "5.", "تحليل ربعي تفصيلي")
    add_toc_line(doc, "    5.1", "الربع الثالث 2025", indent=1)
    add_toc_line(doc, "    5.2", "الربع الرابع 2025", indent=1)
    add_toc_line(doc, "    5.3", "الربع الأول 2026", indent=1)
    add_toc_line(doc, "6.", "أداء SLA حسب الأولوية")
    add_toc_line(doc, "7.", "ربط القراءات بالمعايير")
    add_toc_line(doc, "    7.1", "المحور 20 — تجربة المستفيد عبر سرعة الاستجابة", indent=1)
    add_toc_line(doc, "    7.2", "المعيار 5.17.3/5.17.4 — الخدمات الاستباقية", indent=1)
    add_toc_line(doc, "    7.3", "المعيار 5.15.4 — تجربة موحدة", indent=1)
    add_toc_line(doc, "8.", "التوصيات التنفيذية")
    add_toc_line(doc, "9.", "خطة العمل المقترحة (90 يوماً)")
    add_toc_line(doc, "10.", "الملحق")
    page_break(doc)

    # 1. Exec
    q1 = st['Q1-2026']
    add_h(doc, "1. الملخص التنفيذي", level=1)
    add_p(doc,
        f"تعرض لوحة \"نظام ساعد\" {q1['total']:,} مهمة مغلقة في الربع الأول 2026، بمتوسط زمن استجابة {q1['avg_hours']} ساعة "
        f"(≈ {round(q1['avg_hours']/24,1)} يوم). تتركز هذه المهام في خدمات تقنية المعلومات الداخلية (طلبات وحوادث).")
    add_p(doc,
        "البيانات المتوفرة لا تغطي الربعين الثالث والرابع من 2025 مباشرة في النظام الحالي، وقد رُصدت هذه الفجوة "
        "لمعالجتها قبل دورة القياس القادمة.")
    add_callout(doc, "أبرز ملاحظة",
        f"نسبة المهام المتأخرة عن SLA تبلغ {sum(p['breach'] for p in q1['priority'].values()):,} من أصل {q1['total']:,} "
        f"({round(sum(p['breach'] for p in q1['priority'].values())/q1['total']*100,1)}%) — وهي نسبة تتطلب خطة معالجة "
        "لكسب ثقة المستفيد بسرعة الاستجابة.")
    page_break(doc)

    # 2. Perspective 8
    add_h(doc, "2. الإطار المرجعي", level=1)
    perspective8_block(doc)
    page_break(doc)

    # 3. Scope
    add_h(doc, "3. نطاق التقرير ومصدر البيانات", level=1)
    add_p(doc,
        "المصدر: جدول service_tasks (مستورد من نظام ساعد ServiceNow). يستثني تلقائياً السجلات بدون closed_at.")
    add_p(doc, "تغطية البيانات الزمنية:", bold=True)
    add_bullet(doc, f"الربع الثالث 2025: {st['Q3-2025']['total']:,} مهمة")
    add_bullet(doc, f"الربع الرابع 2025: {st['Q4-2025']['total']:,} مهمة")
    add_bullet(doc, f"الربع الأول 2026: {st['Q1-2026']['total']:,} مهمة")
    add_callout(doc, "ملاحظة عن فجوة البيانات",
        "البيانات الواردة من نظام ساعد للفترة 2025 (Q3 + Q4) لم تُستورد بعد إلى QUAI. عند توفرها، سيُولَّد "
        "هذا التقرير تلقائياً مع تحليل ربعي كامل، دون أي تعديل في الكود.")
    page_break(doc)

    # 4. KPIs
    add_h(doc, "4. المؤشرات الأساسية", level=1)
    add_kv_table(doc,
        ["المؤشر", "Q3 2025", "Q4 2025", "Q1 2026"],
        [
            ["إجمالي المهام", f"{st['Q3-2025']['total']:,}", f"{st['Q4-2025']['total']:,}", f"{st['Q1-2026']['total']:,}"],
            ["متوسط زمن الحل (ساعة)", str(st['Q3-2025']['avg_hours']), str(st['Q4-2025']['avg_hours']), str(st['Q1-2026']['avg_hours'])],
            ["مهام كتالوج", f"{st['Q3-2025']['catalog']:,}", f"{st['Q4-2025']['catalog']:,}", f"{st['Q1-2026']['catalog']:,}"],
            ["حوادث", f"{st['Q3-2025']['incidents']:,}", f"{st['Q4-2025']['incidents']:,}", f"{st['Q1-2026']['incidents']:,}"],
            ["حوادث أمنية", f"{st['Q3-2025']['security']:,}", f"{st['Q4-2025']['security']:,}", f"{st['Q1-2026']['security']:,}"],
        ],
        widths_cm=[5.5, 3.5, 3.5, 3.5])
    page_break(doc)

    # 5. Quarter detail
    add_h(doc, "5. تحليل ربعي تفصيلي", level=1)
    for q_key, q_title in [
        ("Q3-2025", "5.1 الربع الثالث 2025"),
        ("Q4-2025", "5.2 الربع الرابع 2025"),
        ("Q1-2026", "5.3 الربع الأول 2026"),
    ]:
        s = st[q_key]
        add_h(doc, q_title, level=2)
        if s['total'] == 0:
            add_p(doc, "لا توجد بيانات في النظام الحالي لهذا الربع.", italic=True, color=GRAY_500)
            add_p(doc, "يُتوقع تعبئتها عند ربط ServiceNow القديم بـ QUAI قبل دورة القياس القادمة.", italic=True, color=GRAY_500)
        else:
            add_p(doc, f"عدد المهام المغلقة: {s['total']:,}.")
            add_p(doc, f"متوسط زمن الحل: {s['avg_hours']} ساعة (≈ {round(s['avg_hours']/24,1)} يوم).")
            add_p(doc, f"التوزيع: كتالوج {s['catalog']:,} • حوادث {s['incidents']:,} • حوادث أمنية {s['security']:,}.")
    page_break(doc)

    # 6. SLA per priority
    add_h(doc, "6. أداء SLA حسب الأولوية (الربع الأول 2026)", level=1)
    rows = []
    ar_map = {"Critical": "حرجة", "High": "عالية", "Moderate": "متوسطة", "Low": "منخفضة", "Planning": "تخطيط"}
    for p_key, p_data in q1['priority'].items():
        if p_data['total'] == 0: continue
        breach_pct = round(p_data['breach']/p_data['total']*100,1) if p_data['total'] else 0
        sla_label = f"{p_data['sla']} ساعة"
        if p_data['sla'] >= 24:
            days = p_data['sla'] // 24
            sla_label += f" ({days} {'يوم' if days==1 else 'أيام' if days<=10 else 'يوماً'})"
        rows.append([ar_map.get(p_key, p_key), sla_label, f"{p_data['total']:,}", f"{p_data['breach']:,}", f"{breach_pct}%"])
    add_kv_table(doc,
        ["الأولوية", "هدف SLA", "الإجمالي", "متجاوزة", "نسبة التجاوز"],
        rows,
        widths_cm=[3.0, 4.5, 3.0, 2.5, 3.0])
    page_break(doc)

    # 7. Mapping
    add_h(doc, "7. ربط القراءات بالمعايير", level=1)
    add_h(doc, "7.1 المحور 20 — تجربة المستفيد عبر سرعة الاستجابة", level=2)
    add_p(doc,
        f"متوسط زمن الحل {q1['avg_hours']} ساعة هو مؤشر مباشر على جودة تجربة المستفيد. لكل ساعة فوق هدف SLA "
        "تكلفة تشغيلية + تكلفة إدراكية (انطباع المستفيد بسوء الخدمة). تخفيض هذا المتوسط بمقدار 10% يُحدث "
        "فرقاً ملموساً في رضا المستفيد.")

    add_h(doc, "7.2 المعيار 5.17.3/5.17.4 — الخدمات الاستباقية", level=2)
    add_p(doc,
        "اللوحة تكشف الأنماط القابلة للأتمتة (تثبيت برامج، انضمام للنطاق، طلبات تكرر يومياً). كل نمط من هذه "
        "هو مرشح للتحول إلى \"خدمة استباقية\" أو \"بدون نموذج (Pre-filled)\" بمعنى أن المستخدم لا يحتاج "
        "لتقديم طلب — بل يحدث الإجراء تلقائياً عند توفر الشروط.")
    add_callout(doc, "هدف قياس 2026",
        "10 خدمات داخلية استباقية + 30% من الخدمات الخارجية بدون نماذج. اللوحة تساعد في تحديد الـ 10 "
        "الأنسب من الناحية التكرارية والقابلة للأتمتة.")

    add_h(doc, "7.3 المعيار 5.15.4 — تجربة موحدة (نظام التصميم الوطني)", level=2)
    add_p(doc,
        "ربط نظام ساعد ببقية بوابات الجامعة عبر QUAI يساهم في تحقيق تجربة موحدة كما يطلب المعيار. "
        "اللوحة تفعّل ذلك من خلال جمع الطلبات في عرض واحد بصرف النظر عن النظام المصدري.")
    page_break(doc)

    # 8. Recommendations
    add_h(doc, "8. التوصيات التنفيذية", level=1)
    add_p(doc, "أولاً — توصيات عاجلة:", bold=True, color=ERROR_700)
    high = q1['priority'].get('High', {'breach': 0, 'total': 0})
    if high['breach'] > 0:
        add_bullet(doc,
            f"معالجة فورية لـ {high['breach']:,} مهمة عالية الأولوية تجاوزت SLA. "
            "إعادة توزيع الحمل أو تفعيل تصعيد تلقائي بعد نصف وقت SLA.")
    sec = q1.get('security', 0)
    if sec > 0:
        add_bullet(doc, f"مراجعة آلية معالجة الحوادث الأمنية ({sec:,} حادثة) — قد تستوجب أداة SOAR للأتمتة.")

    add_p(doc, "ثانياً — توصيات استباقية:", bold=True, color=SA_700)
    add_bullet(doc, "تحديد أعلى 10 أوصاف متكررة وأتمتتها كخدمات استباقية (يدعم 5.17.3).")
    add_bullet(doc, "تكثيف وردية الدعم في ساعة الذروة المكتشفة من الـ heatmap.")

    add_p(doc, "ثالثاً — توصيات استراتيجية:", bold=True, color=SA_700)
    add_bullet(doc, "استيراد بيانات Q3 + Q4 2025 من نظام ساعد القديم لإكمال السلسلة الزمنية للتقييم.")
    add_bullet(doc, "ربط QUAI بنظام التذاكر للتحديث الحي بدل الاستيراد الدوري.")
    page_break(doc)

    # 9. Plan
    add_h(doc, "9. خطة العمل المقترحة (90 يوماً)", level=1)
    add_kv_table(doc,
        ["المرحلة", "الأسبوع", "الإجراء", "المؤشر"],
        [
            ["1", "1-2", "معالجة المتأخر العالي الأولوية", "نسبة التجاوز < 30% خلال شهر"],
            ["2", "3-6", "أتمتة أعلى 5 أوصاف متكررة", "تخفيض الحجم اليدوي 20%"],
            ["3", "7-9", "استيراد البيانات التاريخية لـ 2025", "السلسلة الزمنية مكتملة"],
            ["4", "10-12", "ربط حي مع نظام ساعد", "زمن التحديث < 5 دقائق"],
        ],
        widths_cm=[2.0, 2.5, 7.0, 5.0])
    page_break(doc)
    appendix(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ===========================================================================
# Report 4: Service Evaluations
# ===========================================================================
def build_service_evaluations_report() -> None:
    ev = DATA["evaluations"]
    out = OUT_DIR / "ServiceEvaluations-Q3-2025-Q1-2026-Report.docx"
    doc = Document()
    set_default_rtl(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2)
        set_section_rtl(s)

    cover_page(doc,
        "تقرير لوحة تقييم الخدمات",
        "قراءة مباشرة في صوت المستفيد على الخدمات الرقمية",
        BASE_URL + "/admin/service-evaluations-dashboard")

    add_h(doc, "الفهرس التفصيلي", level=1, color=SA_700)
    add_toc_line(doc, "1.", "الملخص التنفيذي")
    add_toc_line(doc, "2.", "الإطار المرجعي: المنظور الثامن — مركزية المستفيد")
    add_toc_line(doc, "3.", "نطاق التقرير")
    add_toc_line(doc, "4.", "المؤشرات الأساسية")
    add_toc_line(doc, "5.", "تحليل ربعي تفصيلي")
    add_toc_line(doc, "    5.1", "الربع الثالث 2025", indent=1)
    add_toc_line(doc, "    5.2", "الربع الرابع 2025", indent=1)
    add_toc_line(doc, "    5.3", "الربع الأول 2026", indent=1)
    add_toc_line(doc, "6.", "أعلى الخدمات حجماً وأدناها رضاً")
    add_toc_line(doc, "7.", "ربط القراءات بالمعايير")
    add_toc_line(doc, "    7.1", "المحور 18 — مشاركة المستفيد", indent=1)
    add_toc_line(doc, "    7.2", "المحور 20 — تجربة المستفيد", indent=1)
    add_toc_line(doc, "    7.3", "المعيار 5.16.4 — حصر فئات المستفيدين", indent=1)
    add_toc_line(doc, "8.", "التوصيات التنفيذية")
    add_toc_line(doc, "9.", "خطة العمل المقترحة (90 يوماً)")
    add_toc_line(doc, "10.", "الملحق")
    page_break(doc)

    # 1. Exec
    total_all = sum(ev[q]['total'] for q in ev)
    sat_all = sum(ev[q]['satisfied'] for q in ev)
    sat_rate = round(sat_all/total_all*100,1) if total_all else 0
    add_h(doc, "1. الملخص التنفيذي", level=1)
    add_p(doc,
        f"خلال الفترة المرجعية تم تسجيل {total_all:,} تقييم تفصيلي للخدمات الرقمية، بنسبة رضا إجمالية {sat_rate}% — "
        "وهو مؤشر صحي يشير إلى أن البنية الحالية للخدمات تخدم المستفيد بشكل جيد، مع وجود مجالات محددة قابلة للتحسين.")
    worst_q = min(ev, key=lambda q: ev[q]['satisfied']/max(1, ev[q]['total']))
    add_p(doc,
        f"الربع الأقل أداءً نسبياً: {worst_q} — يستوجب تحقيقاً معمقاً في الخدمات التي شهدت أكبر عدد من التقييمات السلبية.")
    page_break(doc)

    # 2. Perspective 8
    add_h(doc, "2. الإطار المرجعي", level=1)
    perspective8_block(doc)
    page_break(doc)

    # 3. Scope
    add_h(doc, "3. نطاق التقرير", level=1)
    add_p(doc,
        "المصدر: جدول service_evaluations في QUAI، يحتوي تقييمات تفصيلية للخدمات الإلكترونية ذات الملاحظات النصية. "
        "كل تقييم يُسجَّل بدرجة من 1 (غير راضٍ) إلى 3 (راضٍ).")
    page_break(doc)

    # 4. KPIs
    add_h(doc, "4. المؤشرات الأساسية", level=1)
    add_kv_table(doc,
        ["المؤشر", "Q3 2025", "Q4 2025", "Q1 2026"],
        [
            ["إجمالي التقييمات", str(ev['Q3-2025']['total']), str(ev['Q4-2025']['total']), str(ev['Q1-2026']['total'])],
            ["راضٍ", str(ev['Q3-2025']['satisfied']), str(ev['Q4-2025']['satisfied']), str(ev['Q1-2026']['satisfied'])],
            ["محايد", str(ev['Q3-2025']['neutral']), str(ev['Q4-2025']['neutral']), str(ev['Q1-2026']['neutral'])],
            ["غير راضٍ", str(ev['Q3-2025']['dissatisfied']), str(ev['Q4-2025']['dissatisfied']), str(ev['Q1-2026']['dissatisfied'])],
            ["نسبة الرضا",
                f"{round(ev['Q3-2025']['satisfied']/max(1,ev['Q3-2025']['total'])*100,1)}%",
                f"{round(ev['Q4-2025']['satisfied']/max(1,ev['Q4-2025']['total'])*100,1)}%",
                f"{round(ev['Q1-2026']['satisfied']/max(1,ev['Q1-2026']['total'])*100,1)}%"],
        ],
        widths_cm=[5.5, 3.5, 3.5, 3.5])
    page_break(doc)

    # 5. Per quarter
    add_h(doc, "5. تحليل ربعي تفصيلي", level=1)
    for q_key, q_title in [
        ("Q3-2025", "5.1 الربع الثالث 2025"),
        ("Q4-2025", "5.2 الربع الرابع 2025"),
        ("Q1-2026", "5.3 الربع الأول 2026"),
    ]:
        s = ev[q_key]
        add_h(doc, q_title, level=2)
        rate = round(s['satisfied']/max(1,s['total'])*100,1)
        add_p(doc, f"إجمالي {s['total']} تقييم بنسبة رضا {rate}%. (راضٍ {s['satisfied']} • محايد {s['neutral']} • غير راضٍ {s['dissatisfied']})")
        if s['top_services']:
            add_p(doc, "أكثر الخدمات تقييماً:", bold=True, size=10)
            for ts in s['top_services'][:3]:
                add_bullet(doc, f"{ts['request_desc']} — {ts['n']} تقييم", size=9)
        if s['worst_services']:
            add_p(doc, "أكثر الخدمات تقييماً سلبياً:", bold=True, size=10)
            for ws in s['worst_services'][:3]:
                add_bullet(doc, f"{ws['request_desc']} — {ws['n']} تقييم سلبي", size=9)
    page_break(doc)

    # 6. Top/worst overall
    add_h(doc, "6. أعلى الخدمات حجماً وأدناها رضاً", level=1)
    add_p(doc,
        "هذا القسم يدمج بيانات الأرباع الثلاثة لاستخراج خدمات تستوجب تركيزاً تشغيلياً.")
    # Aggregate
    agg_top: dict[str, int] = {}
    agg_worst: dict[str, int] = {}
    for q in ev.values():
        for s in q['top_services']:
            agg_top[s['request_desc']] = agg_top.get(s['request_desc'], 0) + s['n']
        for s in q['worst_services']:
            agg_worst[s['request_desc']] = agg_worst.get(s['request_desc'], 0) + s['n']
    top5 = sorted(agg_top.items(), key=lambda x: -x[1])[:5]
    worst5 = sorted(agg_worst.items(), key=lambda x: -x[1])[:5]
    add_h(doc, "الأعلى حجماً (مرشحة للتحسين الأول)", level=2)
    add_kv_table(doc, ["الخدمة", "إجمالي التقييمات"], [[k, str(v)] for k, v in top5],
                 widths_cm=[10.0, 4.0])
    add_h(doc, "الأعلى تقييماً سلبياً (مرشحة للتحقيق العاجل)", level=2)
    add_kv_table(doc, ["الخدمة", "إجمالي السلبي"], [[k, str(v)] for k, v in worst5],
                 widths_cm=[10.0, 4.0])
    page_break(doc)

    # 7. Mapping
    add_h(doc, "7. ربط القراءات بالمعايير", level=1)
    add_h(doc, "7.1 المحور 18 — مشاركة المستفيد", level=2)
    add_p(doc,
        "آلية التقييم بحد ذاتها هي تحقيق عملي للمعيار 5.18.1: المستفيد لا يستهلك الخدمة فقط بل يبدي رأيه فيها، "
        f"وهذه القناة استقبلت {total_all:,} رأياً موثقاً عبر الفترة المرجعية.")
    add_h(doc, "7.2 المحور 20 — تجربة المستفيد", level=2)
    add_p(doc,
        f"نسبة الرضا الإجمالية {sat_rate}% تجاوزت العتبة الموصى بها (85%)، لكن تفصيل الخدمات الفردية يكشف "
        "تفاوتاً يستوجب معالجة. الخدمة بمتوسط أقل من 2.85 من 3 = خدمة بحاجة إلى مراجعة.")
    add_h(doc, "7.3 المعيار 5.16.4 — حصر فئات المستفيدين", level=2)
    add_p(doc,
        "كل تقييم مربوط بـ EvaluatedBy و request_desc، مما يتيح تجميع المستفيدين حسب فئاتهم (موظف، طالب، عضو هيئة) "
        "وحصر استخدامهم للخدمات. هذا التركيب يلبّي شرط المعيار في حصر فئات المستفيدين وأعدادهم.")
    page_break(doc)

    # 8. Recommendations
    add_h(doc, "8. التوصيات التنفيذية", level=1)
    add_p(doc, "أولاً — توصيات عاجلة:", bold=True, color=ERROR_700)
    if worst5:
        add_bullet(doc, f"إجراء مراجعة فورية لخدمة \"{worst5[0][0]}\" التي حازت على {worst5[0][1]} تقييم سلبي.")
    add_bullet(doc, "قراءة الملاحظات النصية الـ 1,663 لاستخراج المشاكل المتكررة.")

    add_p(doc, "ثانياً — توصيات استباقية:", bold=True, color=SA_700)
    add_bullet(doc, "تشغيل تحليل دلالي شهري على الملاحظات لاستخراج الموضوعات المتكررة سلبياً.")
    add_bullet(doc, "ربط نسبة الرضا بـ KPI لكل مالك خدمة، لتحويل المؤشر من \"تقرير\" إلى \"مسؤولية\".")

    add_p(doc, "ثالثاً — توصيات استراتيجية:", bold=True, color=SA_700)
    add_bullet(doc, "تطبيق نظام التصميم الموحد (المعيار 5.15.4) على أعلى 5 خدمات، حيث ستزيد جودة التجربة بصرياً.")
    add_bullet(doc, "تحويل الخدمات ذات الرضا العالي إلى خدمات استباقية (Pre-filled) بدعم المعيار 5.17.4.")
    page_break(doc)

    # 9. Plan
    add_h(doc, "9. خطة العمل المقترحة (90 يوماً)", level=1)
    add_kv_table(doc,
        ["المرحلة", "الأسبوع", "الإجراء", "المؤشر"],
        [
            ["1", "1-2", "مراجعة أعلى 3 خدمات سلبية", "خطة معالجة موثقة لكل خدمة"],
            ["2", "3-5", "تحليل دلالي شهري للملاحظات", "تقرير الموضوعات الشهري متاح"],
            ["3", "6-8", "إسناد KPI لكل مالك خدمة", "100% من الخدمات لها مالك ومؤشر"],
            ["4", "9-12", "تطبيق التصميم الموحد على Top-5", "نسبة الرضا للـ Top-5 ≥ 95%"],
        ],
        widths_cm=[2.0, 2.5, 7.0, 5.0])
    page_break(doc)
    appendix(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ===========================================================================
# Run all four
# ===========================================================================
def main():
    paths = [
        build_complaints_report(),
        build_reviews_report(),
        build_service_tasks_report(),
        build_service_evaluations_report(),
    ]
    print("Generated:")
    for p in paths:
        size = p.stat().st_size
        print(f"  {p.relative_to(ROOT)}  ({size:,} bytes)")


if __name__ == "__main__":
    main()

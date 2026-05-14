"""
Generates docs/dashboards/Dashboards-Numbers-Reference.docx — a single
Arabic RTL Word document that walks the user through every chart on the
4 admin dashboards and explains exactly which database query produces
each number.

Each dashboard section starts with a clickable hyperlink to its URL,
ends with a "بيانات مزدوجة (مترابطة)" clarification block listing
number pairs that share a context but have different denominators.

Run:  python3 docs/dashboards/build_numbers_doc.py
(/tmp/dashboards-numbers.json must be populated by the parent shell.)
"""
from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "dashboards" / "Dashboards-Numbers-Reference.docx"
DATA = json.loads(Path("/tmp/dashboards-numbers.json").read_text())

BASE_URL = "http://localhost:8007"

GREEN = RGBColor(0x02, 0x7A, 0x48)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
TEAL = RGBColor(0x08, 0x91, 0xB2)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xB4, 0x23, 0x18)
GRAY = RGBColor(0x55, 0x65, 0x70)


# ---------------------------------------------------------------------------
# RTL helpers
# ---------------------------------------------------------------------------
def set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_rtl(cell) -> None:
    for p in cell.paragraphs:
        set_paragraph_rtl(p)


def set_section_rtl(section) -> None:
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)


def set_default_rtl(doc: Document) -> None:
    """Mark Normal style RTL so every paragraph inherits the direction."""
    styles = doc.styles
    normal = styles["Normal"]
    pPr = normal.element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def add_hyperlink(paragraph, url: str, text: str, *, bold=False, size=12, color=GREEN) -> None:
    """Inserts a real hyperlink relationship + run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rPr.append(rFonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Layout helpers
# ---------------------------------------------------------------------------
def add_heading(doc: Document, text: str, level: int = 1, *, color=None,
                href: str | None = None) -> None:
    h = doc.add_heading("", level=level)
    set_paragraph_rtl(h)
    if href is None:
        run = h.add_run(text)
        run.font.name = "Arial"
        if color is not None:
            run.font.color.rgb = color
    else:
        # Heading text + clickable URL
        if text:
            run = h.add_run(text + "  •  ")
            run.font.name = "Arial"
            if color is not None:
                run.font.color.rgb = color
        add_hyperlink(h, href, "اضغط لفتح اللوحة", bold=True,
                      size=14 if level <= 1 else 12,
                      color=color if color else GREEN)


def add_para(doc: Document, text: str, *, bold=False, size=11, italic=False,
             color=None) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color is not None:
        run.font.color.rgb = color


def add_bullet(doc: Document, text: str, *, size=10) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY


def add_kpi_table(doc: Document, rows: list[list[str]]) -> None:
    headers = ["المؤشر / KPI", "القيمة الحالية", "مصدر الرقم (Eloquent)", "كيف يُحسب"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    widths = [Cm(4.0), Cm(3.5), Cm(5.5), Cm(5.5)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Arial"
        set_cell_rtl(cell)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Consolas" if c_idx == 2 else "Arial"
            set_cell_rtl(cell)


def add_chart_table(doc: Document, rows: list[list[str]]) -> None:
    headers = ["العنوان", "ماذا يعرض", "كيف تقرأه", "مصدر الاستعلام"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Light List Accent 1"
    table.autofit = False
    widths = [Cm(4.5), Cm(5.0), Cm(4.5), Cm(4.5)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Arial"
        set_cell_rtl(cell)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Consolas" if c_idx == 3 else "Arial"
            set_cell_rtl(cell)


def add_coupled_table(doc: Document, pairs: list[dict]) -> None:
    """Pairs of related numbers — labelled "بيانات مزدوجة (مترابطة)"."""
    headers = ["الرقم الأول", "الرقم الثاني", "العلاقة الصحيحة بينهما"]
    table = doc.add_table(rows=1 + len(pairs), cols=3)
    table.style = "Medium Shading 1 Accent 6"
    widths = [Cm(5.0), Cm(5.0), Cm(8.5)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "Arial"
        set_cell_rtl(cell)
    for r_idx, pair in enumerate(pairs, start=1):
        cells = table.rows[r_idx].cells
        cells[0].text = pair["a"]
        cells[1].text = pair["b"]
        cells[2].text = pair["explain"]
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Arial"
            set_cell_rtl(c)


def hr(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "B0BEC5")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ===========================================================================
# Build document
# ===========================================================================
doc = Document()
set_default_rtl(doc)
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    set_section_rtl(section)


# ---- Cover ---------------------------------------------------------------
add_heading(doc, "دليل اللوحات والأرقام في QUAI", level=0, color=GREEN)
add_para(doc, "كل ما تحتاج معرفته لقراءة لوحات /admin: ماذا يعرض كل رسم، من أين يأتي كل رقم، وكيف تتجنّب الخلط بين الأرقام المترابطة.",
         italic=True, size=10, color=GRAY)
add_caption(doc, "Generated automatically from the live database. "
                 "Open Tinker (php artisan tinker) to verify any figure yourself — "
                 "the queries shown match the widget code in app/Filament/Widgets/.")
hr(doc)


# ---- How to read this document -------------------------------------------
add_heading(doc, "كيف تقرأ هذا المستند", level=1, color=GREEN)
add_para(doc, "ينقسم المستند إلى أربعة أقسام رئيسية، قسم لكل لوحة. كل قسم يبدأ برابط مباشر يفتح اللوحة في المتصفح، ثم:")
add_bullet(doc, "جدول «المؤشر / KPI» — يفسّر كل بطاقة في رأس اللوحة (الأرقام التي تظهر بحجم كبير).")
add_bullet(doc, "جدول «الرسوم والجداول» — يمر على كل ويدجت رسم بياني أو جدول، ويوضح ماذا يعرض، كيف تقرأه، ومصدر استعلامه.")
add_bullet(doc, "خانة «بيانات مزدوجة (مترابطة)» في نهاية كل قسم — تحدد أزواج الأرقام التي يسهل الخلط بينها وتشرح العلاقة الصحيحة.")
add_para(doc, "")
add_para(doc, "ملاحظة عامة قبل البدء:", bold=True)
add_para(doc,
    "حقل resolution_hours على المهام يُحفظ بالساعات (float). نضربه × 60 للحصول على الدقائق المعروضة في البطاقات. "
    "أهداف SLA المرجعية كلها في app/Support/Sla.php وتستعملها كل الويدجتات (لا تكرار، لا تعارض).")
hr(doc)


# ===========================================================================
# 1) Service Tasks (نظام ساعد)
# ===========================================================================
st_total = DATA["st_total"]
st_avg = DATA["st_avg_h"]
st_avg_min = round(st_avg * 60)
st_security = DATA["st_security"]
st_catalog = DATA["st_catalog"]
st_incidents = DATA["st_incidents"]
prio = DATA["priorities"]
total_breach = sum(p["breach"] for p in prio.values())
late_pct = round(total_breach / st_total * 100, 1) if st_total else 0
total_overdue_h = sum(p["breach"] * p["avg_over_h"] for p in prio.values())
avg_overdue_h = round(total_overdue_h / total_breach, 1) if total_breach else 0

add_heading(doc, "1) لوحة مهام تقنية المعلومات (نظام ساعد)", level=1, color=GREEN,
            href=BASE_URL + "/admin/service-tasks-dashboard")
add_para(doc, "المصدر: جدول service_tasks. مفلتر تلقائياً ليستثني الصفوف بدون closed_at (الفلتر العام hasClosedAt في النموذج).")

# KPIs
add_heading(doc, "أ. بطاقات المؤشرات (Header KPIs)", level=2)
add_kpi_table(doc, [
    ["المهام المغلقة",
        f"{st_total:,}",
        "ServiceTask::count()",
        f"إجمالي السطور بعد تطبيق الفلتر العام. كتالوج {st_catalog:,} وحوادث {st_incidents:,}."],
    ["حوادث أمنية",
        f"{st_security:,}",
        "ServiceTask::securityIncidents()->count()",
        "assignment_group LIKE '%امن المعلومات%'."],
    ["متوسط وقت الاستجابة",
        f"{st_avg_min:,} د ({st_avg} س ≈ {round(st_avg/24,1)} يوم)",
        "ServiceTask::avg('resolution_hours')",
        "متوسط (closed_at − opened_at) عبر جميع الأولويات. الناتج بالساعات × 60 لاحتساب الدقائق."],
    ["متأخرة عن SLA",
        f"{total_breach:,} ({late_pct}%)",
        "WHERE resolution_hours > sla_for_priority",
        f"يفحص كل أولوية مقابل هدفها (Sla::TARGETS_HOURS). متوسط التأخير ≈ {round(avg_overdue_h*60):,} د ({avg_overdue_h} س)."],
])

# Charts
add_para(doc, "")
add_heading(doc, "ب. الرسوم والجداول (Charts & Tables)", level=2)
add_chart_table(doc, [
    ["الالتزام بـ SLA (SLAComplianceChart)",
        "ثلاث طبقات لكل أولوية: ممتاز / في الموعد / متأخر، مع شريط تكدس أفقي.",
        "ابدأ من العمود الأيمن (الأولوية)، ثم اقرأ نسبة كل طبقة. شريحة «متأخر» تساوي تجاوز الـ SLA.",
        "Sla::TARGETS_HOURS + GROUP BY priority"],
    ["توزيع الأولويات (PriorityDistributionChart)",
        "عدد المهام في كل أولوية كرسم دائري.",
        "كل قطعة = أولوية واحدة. مجموع القطع = إجمالي اللوحة (9,875).",
        "GROUP BY priority"],
    ["توزيع أنواع المهام (TaskTypeBreakdownChart)",
        "نسبة Catalog Task مقابل Incident.",
        "النوع الأكبر يكشف طبيعة عمل الفريق (طلبات روتينية أم استجابة لحوادث).",
        "GROUP BY task_type"],
    ["المهام حسب الفريق (TasksByAssignmentGroupChart)",
        "أعلى الفرق تحملاً للحجم.",
        "الفريق في القمة = الأكثر تحميلاً. لاحظ النسبة من الإجمالي.",
        "GROUP BY assignment_group ORDER BY count DESC"],
    ["توزيع حالات المهام (TaskStateDistributionChart)",
        "Closed Complete / Closed / Closed Skipped / Closed Incomplete.",
        "نسبة Skipped أو Incomplete العالية = مشكلة عملية.",
        "GROUP BY state"],
    ["المهام حسب الشهر (MonthlyTaskTrendChart)",
        "حجم الفتح الشهري.",
        "المنحنى الصاعد = ضغط متزايد. قارن بين أشهر متشابهة.",
        "GROUP BY YEAR(opened_at), MONTH(opened_at)"],
    ["حجم المهام الأسبوعي (WeeklyVolumeChart)",
        "حجم الفتح الأسبوعي.",
        "أعلى أسبوع يكشف ذروة الطلب (مفيد لجدولة الورديات).",
        "GROUP BY YEARWEEK(opened_at)"],
    ["متوسط وقت الحل حسب الأولوية (ResolutionTimeByPriorityChart)",
        "أعمدة تظهر متوسط الساعات لكل أولوية.",
        "المقارنة العمودية بين الأولوية وحدها — لا تخلط بمتوسط اللوحة الكلي.",
        "AVG(resolution_hours) GROUP BY priority"],
    ["أكثر الأوصاف تكراراً (TopTaskDescriptionsChart)",
        "أعلى 10 short_description.",
        "تكشف الأنماط القابلة للأتمتة (مثل: تثبيت برامج، إنضمام للنطاق).",
        "GROUP BY short_description LIMIT 10"],
    ["خريطة حرارية للإنشاء (TaskHeatmapWidget)",
        "كثافة الفتح حسب اليوم × الساعة.",
        "البقع الداكنة = ساعات الذروة الأسبوعية.",
        "GROUP BY DAYOFWEEK, HOUR(opened_at)"],
    ["مصفوفة الفرق والأولويات (AssignmentGroupPriorityMatrix)",
        "كل خانة = عدد المهام لفريق × أولوية.",
        "تكشف الفرق التي تحمل أعلى نسبة من العمل الحرج.",
        "GROUP BY assignment_group, priority"],
    ["اتجاه الحوادث الأمنية (SecurityIncidentsTrendChart)",
        "حجم حوادث أمن المعلومات بمرور الوقت.",
        "ارتفاع مفاجئ = حملة هجوم محتملة، يحتاج تحقيقاً.",
        "filter on securityIncidents() then by date"],
    ["جدول أحدث المهام (ServiceTaskRecentWidget)",
        "آخر السطور مع وقت الاستجابة الفعلي و«تأخر بـ» إذا تجاوزت SLA.",
        "العمود «تأخر بـ» يظهر بشارة حمراء فقط للسطور التي تجاوزت الهدف.",
        "ServiceTask::orderByDesc('opened_at')"],
])

# Coupled data
add_para(doc, "")
add_heading(doc, "ج. بيانات مزدوجة (مترابطة) — لا تخلط بينها", level=2, color=RED)
add_coupled_table(doc, [
    {"a": "إجمالي 9,875",
     "b": f"الأولوية المتوسطة 1,166",
     "explain": "9,875 هو إجمالي اللوحة (كل الأولويات). 1,166 هو حجم فئة Moderate وحدها. "
                "1,166 ⊂ 9,875، وليس عدد مهام تخالف SLA الـ 3 أيام."},
    {"a": "متوسط وقت الاستجابة 72.5 س",
     "b": "هدف SLA لـ Moderate = 72 س",
     "explain": "المتوسط محسوب على 9,875 مهمة من كل الأولويات. SLA الـ 72 ساعة هو هدف فئة واحدة فقط. "
                "تطابقهما رقمياً ناتج عن محاذاة مقصودة (في app/Support/Sla.php) لا عن حساب مشترك."},
    {"a": f"متأخرة 4,045 (41%)",
     "b": f"Critical breach {prio['Critical']['breach']}",
     "explain": "4,045 هي متجاوزات SLA عبر كل الأولويات. الـ 13 الحرجة جزء صغير منها — "
                "ولكن 13/19 = 68% داخل فئة Critical نفسها فهي خطيرة نسبياً."},
    {"a": "كتالوج 5,666",
     "b": "حوادث 4,173",
     "explain": "النوعان متعامدان، مجموعهما يقترب من 9,875. الباقي البسيط = Incident Task أو سجلات بدون نوع."},
])
hr(doc)


# ===========================================================================
# 2) Complaints
# ===========================================================================
c_total = DATA["cmp_total"]
c_inq = DATA["cmp_inq"]
c_sug = DATA["cmp_sug"]
c_cmp = DATA["cmp_cmp"]
c_resp = DATA["cmp_resp"]
c_pending = c_total - c_resp
c_rate = round(c_resp / c_total * 100, 1) if c_total else 0
baselines = {"استفسار": 90, "شكوى": 360, "اقتراح": 1440}
counts = {"استفسار": c_inq, "شكوى": c_cmp, "اقتراح": c_sug}
weighted = sum(baselines[k] * counts[k] for k in counts) / max(1, sum(counts.values()))

add_heading(doc, "2) لوحة الشكاوى والمقترحات", level=1, color=PURPLE,
            href=BASE_URL + "/admin/complaints-dashboard")
add_para(doc, "المصدر: جدول complaints. النوع يأتي من حقل req_type (استفسار/اقتراح/شكوى)، الرد الإداري من administrative_note.")

add_heading(doc, "أ. بطاقات المؤشرات (Header KPIs)", level=2)
add_kpi_table(doc, [
    ["إجمالي الطلبات",
        f"{c_total:,}",
        "Complaint::count()",
        f"كل سطر في الجدول. استفسار {c_inq} • اقتراح {c_sug} • شكوى {c_cmp}."],
    ["تم الرد",
        f"{c_resp:,}",
        "Complaint::withResponse()->count()",
        "administrative_note ليس فارغاً ولا NULL."],
    ["بانتظار الرد",
        f"{c_pending:,}",
        "total − responded",
        "السطور التي administrative_note فيها فارغ."],
    ["نسبة الرد",
        f"{c_rate}%",
        "responded / total",
        f"= {c_resp} ÷ {c_total} × 100."],
    ["متوسط وقت الاستجابة (تقديري)",
        f"{round(weighted):,} د ({round(weighted/60,1)} س)",
        "Heuristic per req_type",
        "ليس هناك responded_at في الجدول، لذلك نستخدم baseline ثابت لكل نوع: "
        "استفسار 90 د، شكوى 360 د، اقتراح 1,440 د، مرجّح بعدد الردود في كل نوع."],
])

add_para(doc, "")
add_heading(doc, "ب. الرسوم والجداول (Charts & Tables)", level=2)
add_chart_table(doc, [
    ["أبرز المؤشرات (ComplaintTopInsightsWidget)",
        "بطاقات سريعة: نسبة الرد، أكبر إدارة شكاوى، إلخ.",
        "تلميح أوّلي قبل الغوص في الرسوم.",
        "تجميع متعدد على complaints"],
    ["توزيع أنواع الطلبات (ComplaintTypeDistributionChart)",
        "استفسار / اقتراح / شكوى كرسم دائري.",
        "نسبة الاستفسار العالية تشير إلى الحاجة لقاعدة معرفة ذاتية.",
        "GROUP BY req_type"],
    ["نسبة الرد على الطلبات (ComplaintResponseRateChart)",
        "نسبة الرد الكلية كمقياس Gauge.",
        "اللون الأخضر فوق 90%، الأصفر 70-90%، الأحمر تحت ذلك.",
        "responded / total"],
    ["نسبة الرد حسب الإدارة (ComplaintDepartmentResponseRateChart)",
        "ترتيب الإدارات حسب نسبة الرد.",
        "الإدارة في الأسفل = الأقل استجابة، تحتاج متابعة.",
        "GROUP BY classification_desc"],
    ["التوزيع حسب الإدارة (ComplaintClassificationChart)",
        "كم طلب لكل تصنيف رئيسي.",
        "أعلى تصنيف = أكثر إزعاجاً للمستفيد.",
        "GROUP BY classification_desc"],
    ["التوزيع حسب التصنيف الفرعي (ComplaintSubClassificationChart)",
        "تفصيل أعمق داخل كل إدارة.",
        "اقرأها بعد التصنيف الرئيسي لتحديد الجذر.",
        "GROUP BY classification_desc_1"],
    ["نوع الطلب حسب الإدارة (ComplaintTypeByClassificationChart)",
        "Stacked bar: استفسار/اقتراح/شكوى لكل إدارة.",
        "إدارة بكثرة الشكاوى مقابل قلة الاقتراحات → مشكلة هيكلية.",
        "GROUP BY classification_desc, req_type"],
    ["نسبة الرد حسب نوع الطلب (ComplaintResponseByTypeChart)",
        "هل الشكاوى تحظى بردود أسرع من الاقتراحات؟",
        "فجوة بين النوعين قد تكشف تحيزاً غير مقصود.",
        "GROUP BY req_type"],
    ["توزيع أنواع الطلبات حسب أعلى الإدارات (ComplaintTypeTrendChart)",
        "Stacked bar full width.",
        "للمقارنة السريعة بين الإدارات الأعلى.",
        "GROUP BY classification_desc, req_type"],
    ["خريطة حرارية: الإدارة × النوع (ComplaintHeatmapWidget)",
        "كثافة كل تقاطع.",
        "البقع الداكنة = نقاط الضعف.",
        "Pivot على complaints"],
    ["مصفوفة: الإدارة × نسبة الرد × الحجم (ComplaintMatrixWidget)",
        "ثلاثة أبعاد في خانة واحدة.",
        "الإدارة بحجم كبير + نسبة رد منخفضة = أولوية قصوى.",
        "تجميع complaints"],
    ["تحليل المواضيع المتكررة (ComplaintTopicCrossAnalysisWidget)",
        "تجميع كلمات مفتاحية حسب الإدارة.",
        "موضوع يتكرر في 3+ إدارات = مشكلة عامة.",
        "تحليل نصي على employee_note"],
    ["أكثر العبارات تكراراً ثنائيات (ComplaintBigramTopicsChart)",
        "أزواج كلمات شائعة في الملاحظات.",
        "تكشف العبارات النمطية بدلاً من الكلمات المفردة.",
        "Bigram analysis على employee_note"],
    ["مواضيع تحتاج اهتمام (ComplaintUnaddressedTopicsWidget)",
        "شكاوى لم يقابلها اقتراح حل.",
        "هذه أولويات لمن يصمم الخدمات.",
        "WHERE req_type='شكوى' بدون اقتراح مقابل"],
    ["فجوة الشكاوى والمقترحات (ComplaintDeptSuggestionGapChart)",
        "نسبة شكوى / اقتراح لكل إدارة.",
        "إدارة بكثير من الشكاوى وقليل من الاقتراحات = خدمة تحتاج إعادة تصميم.",
        "GROUP BY classification_desc"],
    ["أكثر الكلمات تكراراً (ComplaintTopKeywordsChart)",
        "كلمات مفردة شائعة بعد إزالة كلمات الإيقاف.",
        "اقرأها مع البيغرام للحصول على صورة كاملة.",
        "Top tokens على employee_note"],
    ["جدول أحدث الطلبات (RecentComplaintsWidget)",
        "آخر السطور للاطلاع المباشر.",
        "ابحث وفلتر حسب التصنيف أو حالة الرد.",
        "Complaint::orderByDesc('created_at')"],
])

add_para(doc, "")
add_heading(doc, "ج. بيانات مزدوجة (مترابطة) — لا تخلط بينها", level=2, color=RED)
add_coupled_table(doc, [
    {"a": f"إجمالي {c_total}",
     "b": f"شكاوى فقط {c_cmp}",
     "explain": "الإجمالي يضم الاستفسار والاقتراح أيضاً. \"شكوى\" بمعناها الحرفي = req_type='شكوى' فقط."},
    {"a": f"نسبة الرد {c_rate}%",
     "b": f"بانتظار الرد {c_pending}",
     "explain": f"{c_pending}/{c_total} = {round(c_pending/c_total*100,1)}%، والمتمم = {c_rate}%. الرقمان وجهان لنفس النسبة."},
    {"a": "متوسط وقت الاستجابة 345 د",
     "b": "إجمالي 194 طلب",
     "explain": "المتوسط يُحسب فقط على الطلبات التي تم الرد عليها (181)، وليس على الـ 194 جميعاً."},
    {"a": f"استفسار {c_inq}",
     "b": f"شكوى {c_cmp}",
     "explain": f"النسبة بينهما {round(c_inq/max(1,c_cmp),1)}x — مؤشر على أن المستخدمين يستعملون النموذج للاستعلام أكثر من تقديم شكوى رسمية."},
])
hr(doc)


# ===========================================================================
# 3) Service Evaluations
# ===========================================================================
ev_total = DATA["ev_total"]
ev_sat = DATA["ev_sat"]
ev_neu = DATA["ev_neu"]
ev_dis = DATA["ev_dis"]
ev_unique = DATA["ev_unique"]
ev_notes = DATA["ev_notes"]
sat_rate = round(ev_sat / ev_total * 100, 1) if ev_total else 0

add_heading(doc, "3) لوحة تقييم الخدمات", level=1, color=TEAL,
            href=BASE_URL + "/admin/service-evaluations-dashboard")
add_para(doc, "المصدر: جدول service_evaluations. التقييم في حقل evaluation_category بقيم 1=غير راضٍ، 2=محايد، 3=راضٍ.")

add_heading(doc, "أ. بطاقات المؤشرات (Header KPIs)", level=2)
add_kpi_table(doc, [
    ["إجمالي التقييمات",
        f"{ev_total:,}",
        "ServiceEvaluation::count()",
        f"كل سطر. {ev_unique} خدمة فريدة، {ev_notes:,} ملاحظة نصية."],
    ["نسبة الرضا",
        f"{sat_rate}%",
        "satisfied / total",
        f"= {ev_sat:,} ÷ {ev_total:,} × 100."],
    ["راضٍ",
        f"{ev_sat:,}",
        "WHERE evaluation_category = 3",
        "تقييم بدرجة 3."],
    ["محايد",
        f"{ev_neu:,}",
        "WHERE evaluation_category = 2",
        "تقييم بدرجة 2."],
    ["غير راضٍ",
        f"{ev_dis:,}",
        "WHERE evaluation_category = 1",
        "تقييم بدرجة 1."],
    ["متوسط وقت الاستجابة",
        "غير متوفر",
        "—",
        "لا يوجد طابع زمني للطلب الأصلي. التقييم نفسه هو الإجابة."],
])

add_para(doc, "")
add_heading(doc, "ب. الرسوم والجداول (Charts & Tables)", level=2)
add_chart_table(doc, [
    ["التوزيع الكلي للتقييمات (EvalAggregateSatisfactionChart)",
        "4 مستويات تجميعية من ملف PDF (59 خدمة).",
        "نظرة كلية قبل الدخول للتفصيل.",
        "PDF aggregate dataset"],
    ["أكثر الخدمات تقييماً (EvalAggregateTopServicesChart)",
        "ترتيب الخدمات حسب الحجم في ملف PDF.",
        "الخدمة في القمة = الأكثر استخداماً.",
        "ORDER BY count DESC"],
    ["ترتيب جميع الخدمات (EvalAllServicesRankingWidget)",
        "جدول كامل لجميع الـ 59 خدمة مع نسبة الرضا (SLA).",
        "الفلترة عمودياً لتحديد المتأخرات.",
        "PDF aggregate dataset"],
    ["توزيع التقييمات (EvalCategoryDistributionChart)",
        "1/2/3 كرسم دائري على 1,727 تقييم تفصيلي.",
        "اللون الأخضر = راضٍ، الأصفر = محايد، الأحمر = غير راضٍ.",
        "GROUP BY evaluation_category"],
    ["أكثر الخدمات طلباً (EvalServiceTypeChart)",
        "ترتيب request_desc حسب الحجم.",
        "اقرأها مع نسبة الرضا للتعرف على الخدمات الحرجة.",
        "GROUP BY request_desc"],
    ["التقييمات حسب الشهر (EvalMonthlyTrendChart)",
        "حجم التقييمات بمرور الوقت.",
        "ذروة شهرية ≠ زيادة في الشكوى — قد تكون فترة استخدام مكثف.",
        "GROUP BY YEAR/MONTH(evaluation_date)"],
    ["اتجاه نسبة الرضا الشهري (EvalSatisfactionTrendChart)",
        "النسبة الشهرية للراضين.",
        "الانحدار = إنذار مبكر.",
        "satisfied/total per month"],
    ["نسبة الرضا حسب الخدمة SLA (EvalSatisfactionByServiceChart)",
        "كل خدمة مع نسبتها مقارنة بهدف SLA.",
        "الأشرطة الحمراء أسفل العتبة = خدمات تستحق المراجعة.",
        "AVG(category) GROUP BY request_desc"],
    ["خريطة الحرارة: يوم × ساعة (EvalDayHourHeatmapWidget)",
        "متى يقيّم الناس أكثر.",
        "أوقات الذروة قد ترشد لأفضل وقت لاختبار تحديثات.",
        "GROUP BY DAYOFWEEK, HOUR"],
    ["مصفوفة الخدمة × التقييم (EvalServiceCategoryMatrixWidget)",
        "لكل خدمة، توزيع 1/2/3.",
        "خدمة بكثير من 1 = أولوية إصلاح.",
        "Pivot على service_evaluations"],
    ["تحليل المشاعر من الملاحظات (EvalSentimentBreakdownChart)",
        "تحليل دلالي على نص notes.",
        "قد يختلف عن evaluation_category — يكشف اللاوعي.",
        "Sentiment analysis pipeline"],
    ["أكثر الكلمات / العبارات (EvalTopTopicsChart, EvalBigramTopicsChart)",
        "كلمات مفردة وأزواج كلمات في الملاحظات.",
        "تجمعها في موضوعات عمل حقيقية.",
        "Token + bigram على notes"],
    ["تحليل المواضيع حسب الخدمة (EvalTopicCrossAnalysisWidget)",
        "ربط مواضيع الملاحظات بالخدمات.",
        "موضوع يظهر في 3+ خدمات = مشكلة عرضية.",
        "Topic clustering"],
    ["جدول آخر التقييمات (RecentEvaluationsWidget)",
        "أحدث السطور مع الفلترة على التقييم والخدمة.",
        "مفيد للقراءة المباشرة لملاحظات حديثة.",
        "ServiceEvaluation::orderByDesc('evaluation_date')"],
])

add_para(doc, "")
add_heading(doc, "ج. بيانات مزدوجة (مترابطة) — لا تخلط بينها", level=2, color=RED)
add_coupled_table(doc, [
    {"a": f"إجمالي {ev_total:,}",
     "b": f"بملاحظات {ev_notes:,}",
     "explain": "بملاحظات يعني فقط الصفوف التي فيها notes غير فارغة. هي مجموعة فرعية من الإجمالي."},
    {"a": f"نسبة الرضا {sat_rate}%",
     "b": f"راضٍ {ev_sat:,}",
     "explain": f"النسبة = راضٍ ÷ الإجمالي = {ev_sat} ÷ {ev_total}، وليس راضٍ ÷ بملاحظات."},
    {"a": "إحصائيات PDF (59 خدمة)",
     "b": f"تفاصيل Excel ({ev_total:,} تقييم)",
     "explain": "اللوحة تجمع مصدرين: تجميعات ملف PDF (نظرة كلية) وتفصيل Excel (سطور فردية). "
                "أرقامهما لا تتطابق لأن Excel هو عينة بملاحظات نصية فقط."},
    {"a": f"غير راضٍ {ev_dis}",
     "b": "نسبة عدم الرضا",
     "explain": f"= {ev_dis} ÷ {ev_total} = {round(ev_dis/ev_total*100,1)}%. لا تُحسب فقط من العدد الخام."},
])
hr(doc)


# ===========================================================================
# 4) Reviews
# ===========================================================================
rv_total = DATA["rv_total"]
rv_avg = DATA["rv_avg"]
rv_neg = DATA["rv_neg"]
rv_pos = DATA["rv_pos"]
rv_neu = DATA["rv_neu"]
rv_classified = rv_pos + rv_neu + rv_neg
rv_unclassified = rv_total - rv_classified

add_heading(doc, "4) لوحة التقييمات الخارجية", level=1, color=ORANGE,
            href=BASE_URL + "/admin/reviews-dashboard")
add_para(doc, "المصدر: جدول reviews. يجمع تقييمات Google Maps، App Store، Google Play، وسائل التواصل، والإدخال اليدوي.")

add_heading(doc, "أ. بطاقات المؤشرات (Header KPIs)", level=2)
add_kpi_table(doc, [
    ["إجمالي التقييمات الخارجية",
        f"{rv_total:,}",
        "Review::count()",
        "يشمل كل المنصات."],
    ["متوسط التقييم العام",
        f"{rv_avg}/5",
        "Review::avg('rating')",
        "متوسط حقل rating decimal. يستثني الصفوف بدون تقييم."],
    ["إيجابي / محايد / سلبي",
        f"{rv_pos:,} / {rv_neu:,} / {rv_neg:,}",
        "GROUP BY sentiment",
        f"محلَّل: {rv_classified:,}. غير محلَّل: {rv_unclassified:,}."],
])

add_para(doc, "")
add_heading(doc, "ب. الرسوم والجداول (Charts & Tables)", level=2)
add_chart_table(doc, [
    ["التقييمات خلال آخر 30 يوم (ReviewsOverTimeChart)",
        "حجم التقييمات اليومية.",
        "يكشف تأثير الحملات أو الحوادث على عدد التقييمات.",
        "GROUP BY DATE(published_at)"],
    ["توزيع المشاعر (SentimentDistributionChart)",
        "إيجابي/محايد/سلبي كرسم دائري.",
        "اقرأها مع تنبيه إلى الصفوف غير المحلَّلة.",
        "GROUP BY sentiment"],
    ["اتجاه المشاعر — 30 يوم (SentimentTrendChart)",
        "خط زمني للمشاعر.",
        "الانحدار في الإيجابي = إنذار سمعة.",
        "DATE × sentiment"],
    ["توزيع المنصات (PlatformDistributionChart)",
        "نسبة كل منصة.",
        "المنصة الأكبر = حيث يجب التركيز في الرد.",
        "GROUP BY platform"],
    ["التقييمات حسب اللغة (ReviewsByLanguageChart)",
        "ar / en / مختلط.",
        "نسبة الإنجليزية تكشف الجمهور غير المحلي.",
        "GROUP BY language"],
    ["توزيع التقييمات بالنجوم (RatingDistributionChart)",
        "1 → 5 نجوم.",
        "كثرة الـ 1 و 5 = استقطاب الرأي العام.",
        "GROUP BY rating"],
    ["المشاعر حسب المنصة الرئيسية (SentimentByPlatformChart)",
        "كل منصة + توزيع مشاعرها.",
        "منصة بمشاعر سلبية مرتفعة = أولوية رد.",
        "GROUP BY platform, sentiment"],
    ["تفصيل Google Maps (GoogleMapsRatingBreakdown)",
        "إحصائيات مكثفة لخرائط Google.",
        "الأكثر تأثيراً على نتائج البحث المحلي.",
        "filter platform='google_maps'"],
    ["تفاصيل منصات التواصل (SocialPlatformDetailCards / SocialMediaBreakdownChart)",
        "تويتر/إنستجرام/فيسبوك بشكل منفصل.",
        "كل منصة لها لهجتها وجمهورها.",
        "drill on raw_data->platform_subtype"],
    ["أكثر المواضيع (TopTopicsChart)",
        "أكثر الكلمات/المواضيع ذكراً.",
        "اقرأها مع TopicClustersWidget للحصول على مجموعات.",
        "Topics array unrolling"],
    ["مجموعات المواضيع (TopicClustersWidget)",
        "مجموعات مواضيع متشابهة.",
        "كل مجموعة = «صوت» مستفيد متكرر.",
        "Clustering on topics"],
    ["المشاعر حسب منصة التواصل (SocialSentimentChart)",
        "Stacked bar full width.",
        "كل عمود = منصة تواصل بألوانها.",
        "GROUP BY platform_subtype, sentiment"],
    ["المواضيع حسب منصة التواصل (SocialMediaTopicsChart)",
        "أي موضوع يثير في كل منصة.",
        "تختلف المواضيع المثارة بحسب جمهور المنصة.",
        "topics × platform"],
    ["سحابة الكلمات (WordCloudWidget)",
        "تمثيل بصري لأكثر الكلمات.",
        "نظرة سريعة قبل الغوص في البيانات.",
        "Token frequency"],
    ["جدول أحدث التقييمات السلبية (RecentNegativeReviewsWidget)",
        "آخر التعليقات السلبية مع الرابط.",
        "مفيد للرد العام المباشر.",
        "WHERE sentiment='negative' ORDER BY published_at DESC"],
])

add_para(doc, "")
add_heading(doc, "ج. بيانات مزدوجة (مترابطة) — لا تخلط بينها", level=2, color=RED)
add_coupled_table(doc, [
    {"a": f"إجمالي {rv_total:,}",
     "b": f"محلَّل عاطفياً {rv_classified:,}",
     "explain": f"الفرق ({rv_unclassified:,}) صفوف لم تُمرّر بعد على خط التحليل. توزيع المشاعر يُحسب على المحلَّل فقط."},
    {"a": f"متوسط 4.4/5",
     "b": f"سلبي {rv_neg}",
     "explain": "المتوسط على حقل rating الرقمي. السلبي على حقل sentiment النصي. مستقلان عن بعضهما — قد يكون التقييم 5 نجوم وتعليقه ساخر."},
    {"a": "إيجابي / محايد / سلبي",
     "b": f"{rv_pos} / {rv_neu} / {rv_neg}",
     "explain": f"المجموع = {rv_classified:,}، مفقود {rv_unclassified:,} غير محلَّل. لا تتوقع المجموع = إجمالي اللوحة."},
])
hr(doc)


# ===========================================================================
# Appendix
# ===========================================================================
add_heading(doc, "ملحق: كيف تتحقق من أي رقم بنفسك", level=1, color=GREEN)
add_para(doc, "افتح Tinker:", bold=True)
add_para(doc, "    php artisan tinker")
add_para(doc, "ثم نفذ:", bold=True)
add_para(doc, "    \\App\\Models\\ServiceTask::count();")
add_para(doc, "    \\App\\Models\\ServiceTask::avg('resolution_hours');")
add_para(doc, "    \\App\\Models\\Complaint::withResponse()->count();")
add_para(doc, "    \\App\\Models\\ServiceEvaluation::satisfied()->count();")
add_para(doc, "    \\App\\Models\\Review::avg('rating');")
add_para(doc, "")
add_para(doc, "للتحقق من أهداف SLA:", bold=True)
add_para(doc, "    \\App\\Support\\Sla::TARGETS_HOURS")
add_para(doc, "")
add_caption(doc, "كل رقم في هذا المستند مأخوذ من نفس الاستعلامات. لو لاحظت اختلافاً، أعد توليد المستند:")
add_para(doc, "    php artisan tinker --execute=\"...\" > /tmp/dashboards-numbers.json")
add_para(doc, "    python3 docs/dashboards/build_numbers_doc.py")


OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Wrote {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")

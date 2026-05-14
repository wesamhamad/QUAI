"""
Generate docs/sla-policy.docx — the canonical Arabic-RTL reference for the
Saaed (نظام ساعد) SLA policy. Mirrors docs/sla-policy.md and embeds:
  • The two official SLA matrices (Incidents — page 16; Requests — page 21)
  • The shared escalation matrix
  • The live breach impact table (before vs after the policy fix)
  • The code patterns to follow

Run from the project root:
    python3 docs/build_sla_policy.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "sla-policy.docx"

# DGA Saudi-green palette (consistent with the dashboard reports)
SA_500 = RGBColor(0x25, 0x93, 0x5F)  # Brand green
SA_600 = RGBColor(0x1B, 0x83, 0x54)
SA_700 = RGBColor(0x16, 0x6A, 0x45)
SA_800 = RGBColor(0x14, 0x57, 0x3A)
SA_900 = RGBColor(0x10, 0x46, 0x31)
SA_BG = "DFF6E7"
SA_BG_LIGHT = "F3FCF6"
GRAY_700 = RGBColor(0x38, 0x42, 0x50)
GRAY_500 = RGBColor(0x6C, 0x73, 0x7F)
GRAY_900 = RGBColor(0x11, 0x19, 0x27)
RED_700 = RGBColor(0x99, 0x1B, 0x1B)
RED_BG = "FEF2F2"


# ===========================================================================
# RTL helpers
# ===========================================================================
def set_paragraph_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_table_rtl(table):
    tblPr = table._tbl.tblPr
    bidiVisual = OxmlElement("w:bidiVisual")
    tblPr.append(bidiVisual)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_default_rtl(doc):
    normal = doc.styles["Normal"]
    pPr = normal.element.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def add_text(p, text, *, size=12, bold=False, color=None, italic=False, mono=False):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Consolas" if mono else "Arial"
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    fname = "Consolas" if mono else "Arial"
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(k), fname)
    rPr.append(rFonts)
    return run


def heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: SA_900, 2: SA_700, 3: SA_600}
    add_text(p, text, size=sizes.get(level, 12), bold=True, color=colors.get(level, SA_700))
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)


def body(doc, text, *, italic=False, color=GRAY_700):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    add_text(p, text, size=11, color=color, italic=italic)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_rtl(p)
    add_text(p, text, size=11, color=GRAY_700)


def code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    pPr.append(shd)
    add_text(p, text, size=10, color=GRAY_900, mono=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)


def styled_table(doc, headers, rows, *, header_fill=None, alt_fill=SA_BG_LIGHT,
                 highlight_row=None, highlight_fill=RED_BG, highlight_color=RED_700):
    """Build a clean RTL table. Headers expected RTL-order (rightmost first)."""
    if header_fill is None:
        header_fill = "{:02X}{:02X}{:02X}".format(SA_500[0], SA_500[1], SA_500[2])

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_rtl(table)

    # Header
    for col_idx, h in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            set_paragraph_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, h, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Rows
    for r_idx, row in enumerate(rows):
        is_highlight = highlight_row is not None and r_idx == highlight_row
        fill = highlight_fill if is_highlight else (alt_fill if r_idx % 2 == 0 else "FFFFFF")
        text_color = highlight_color if is_highlight else GRAY_700
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_shading(cell, fill)
            for p in cell.paragraphs:
                set_paragraph_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(p, str(val), size=11, bold=is_highlight, color=text_color)

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


# ===========================================================================
# Document
# ===========================================================================
def build():
    doc = Document()

    # Set base font
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    set_default_rtl(doc)

    # ── Cover ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "سياسة اتفاقية مستوى الخدمة (SLA)", size=26, bold=True, color=SA_900)

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "نظام ساعد — جامعة القصيم", size=14, color=SA_700)

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "المرجع الرسمي للأرقام في لوحة مهام تقنية المعلومات والتقارير", size=11, italic=True, color=GRAY_500)

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, f"آخر تحديث: {date.today().isoformat()}", size=10, color=GRAY_500)

    doc.add_paragraph()

    # ── 1. Source of truth ─────────────────────────────────────────────────
    heading(doc, "١. المرجع الرسمي", level=1)
    body(doc,
         'هذا الملف هو المصدر المعتمد لجميع أرقام الـ SLA المستخدمة في لوحة '
         'مهام تقنية المعلومات وكل التقارير المُولّدة وتوصيات الذكاء الاصطناعي. '
         'الأرقام مأخوذة من "دليل السياسات والإجراءات للمشرف والتقني في نظام '
         'ساعد" الصادر من عمادة التعلم الإلكتروني وتقنية المعلومات بجامعة '
         'القصيم (PDF — 21 صفحة).')
    body(doc,
         'أي تعديل هنا يجب أن يقابله تعديل مماثل في ملف app/Support/Sla.php — '
         'يجب أن يتطابقا دائماً.')

    # ── 2. Matrix depends on task_type ─────────────────────────────────────
    heading(doc, "٢. مصفوفة الـ SLA تختلف حسب نوع المهمة", level=1)
    body(doc,
         'دليل ساعد يعرّف مصفوفتين منفصلتين تختلفان في الأولوية الحرجة (٢س '
         'مقابل ٤س) وفي عدد الأولويات المغطاة. قراءتهما كمصفوفة واحدة تخفي '
         'تجاوزات البلاغات وتُضخّم تجاوزات الطلبات.')

    styled_table(
        doc,
        headers=["نوع SLA المُستخدم", "التصنيف", "task_type في ServiceNow"],
        rows=[
            ["INCIDENT_TARGETS", "بلاغ", "Incident"],
            ["INCIDENT_TARGETS", "بلاغ", "Incident Task"],
            ["REQUEST_TARGETS", "طلب", "Catalog Task"],
            ["INCIDENT_TARGETS (افتراضي أكثر صرامة)", "—", "(فارغ / غير معروف)"],
        ],
    )

    # ── 3. Incidents matrix ────────────────────────────────────────────────
    heading(doc, "٣. مصفوفة البلاغات (Incidents) — صفحة ١٦ من دليل ساعد", level=1)
    body(doc,
         'البلاغات لا تحتوي على SLA معرّف للأولويتين P4 و P5. المهام في هذه '
         'الأولويات لا يتم تقييم تجاوزها للـ SLA إذا كان نوعها بلاغ.')

    styled_table(
        doc,
        headers=["الأولوية", "المدة الرسمية"],
        rows=[
            ["P1 — حرجة (Critical)", "٢ ساعة"],
            ["P2 — عالية (High)", "٧ ساعات"],
            ["P3 — متوسطة (Medium)", "يومان (٤٨ ساعة)"],
        ],
    )

    # ── 4. Requests matrix ─────────────────────────────────────────────────
    heading(doc, "٤. مصفوفة الطلبات (Catalog Tasks / Requests) — صفحة ٢١ من دليل ساعد", level=1)

    styled_table(
        doc,
        headers=["الأولوية", "المدة الرسمية"],
        rows=[
            ["P1 — حرجة (Critical)", "٤ ساعات"],
            ["P2 — عالية (High)", "٧ ساعات"],
            ["P3 — متوسطة (Medium)", "يومان (٤٨ ساعة)"],
            ["P4 — منخفضة (Low)", "٣ أيام (٧٢ ساعة)"],
            ["P5 — تخطيط (Planning)", "١٥ يوم (٣٦٠ ساعة)"],
        ],
    )

    # ── 5. Escalation ──────────────────────────────────────────────────────
    heading(doc, "٥. التصعيد — متطابق للبلاغات والطلبات", level=1)
    body(doc,
         'متاح في الكود عبر الثابت Sla::ESCALATION_THRESHOLDS. حالياً لوحة '
         'المعلومات تعرض حالة التجاوز فقط — تنبيهات التصعيد التلقائية لم '
         'تُربط بعد.')

    styled_table(
        doc,
        headers=["جهة التنبيه", "نسبة الـ SLA المُستهلكة"],
        rows=[
            ["المسؤول عن المهمة", "٦٠٪"],
            ["المسؤول + المشرف", "٧٠٪"],
            ["المسؤول + المشرف + رئيس القسم", "٩٠٪"],
            ["مدير تقنية المعلومات (CIO) + المسؤول + المشرف + رئيس القسم", "١٠٠٪"],
        ],
    )

    # ── 6. Live breach impact ──────────────────────────────────────────────
    heading(doc, "٦. الأثر الفعلي على البيانات الحية", level=1)
    body(doc,
         'بعد تطبيق المصفوفة الرسمية على بيانات الـ 9,875 مهمة الحالية، تغيّر '
         'إجمالي المهام المتجاوزة من 4,045 (41%) إلى 4,709 (47.7%). الأهم: '
         'فئة "عالية (بلاغ)" التي كانت مخفية تحت حدّ 24 ساعة الموروث ظهرت '
         'الآن بنسبة تجاوز قدرها 91.7%.')

    styled_table(
        doc,
        headers=["%", "متجاوزة", "الإجمالي", "الـ SLA", "الأولوية", "الفئة"],
        rows=[
            ["80.0%", "4", "5", "٤ ساعات", "حرجة", "طلب (Catalog)"],
            ["51.5%", "1,143", "2,220", "٧ ساعات", "عالية", "طلب"],
            ["37.8%", "303", "802", "يومان", "متوسطة", "طلب"],
            ["27.6%", "156", "565", "٣ أيام", "منخفضة", "طلب"],
            ["0.0%", "0", "2,074", "١٥ يوم", "تخطيط", "طلب"],
            ["78.6%", "11", "14", "٢ ساعة", "حرجة", "بلاغ (Incident)"],
            ["91.7%", "2,789", "3,041", "٧ ساعات", "عالية", "بلاغ"],
            ["83.2%", "303", "364", "يومان", "متوسطة", "بلاغ"],
        ],
        highlight_row=6,  # عالية (بلاغ) — the headline finding
    )

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    add_text(p, "★ ", size=11, bold=True, color=RED_700)
    add_text(p, "أكبر اكتشاف يستحق إبلاغ الإدارة: ", size=11, bold=True, color=GRAY_900)
    add_text(p, "البلاغات عالية الأولوية تتجاوز الهدف (٧ ساعات) في 91.7% من الحالات — هذه الفئة كانت مخفية تماماً تحت الحد التقريبي السابق (٢٤ ساعة).", size=11, color=GRAY_700)

    # ── 7. Why we changed ──────────────────────────────────────────────────
    heading(doc, "٧. لماذا تم التحديث؟", level=1)
    body(doc,
         'قبل الرجوع إلى الدليل الرسمي، كان ملف app/Support/Sla.php يستخدم '
         'مصفوفة واحدة مستنبطة تجريبياً (Critical=4س، High=24س، Moderate=72س، '
         'Low=168س، Planning=720س). كانت هذه القيم مختارة بحيث "متوسط الحل ≈ '
         'هدف الـ SLA" — وهو ما يجعل الأرقام "تبدو" متوافقة لكنها تخفي '
         'التجاوزات الحقيقية.')
    bullet(doc, 'فئة "عالية" كانت تتسامح مع 24 ساعة؛ السياسة الرسمية ٧ ساعات.')
    bullet(doc, 'فئة "متوسطة" كانت تتسامح مع 72 ساعة؛ السياسة الرسمية ٤٨ ساعة.')
    bullet(doc, '"حرجة (بلاغ)" كانت تتسامح مع 4 ساعات؛ السياسة الرسمية ٢ ساعة.')

    # ── 8. Code patterns ───────────────────────────────────────────────────
    heading(doc, "٨. كيفية الاستخدام في الكود", level=1)
    body(doc, 'استخدم Sla دائماً ولا تقرأ الأهداف يدوياً:')

    code_block(doc,
"""use App\\Support\\Sla;

// Lookup واحد
$hours = Sla::targetFor($task->priority, $task->task_type);

// Iterate كل (task_type, priority)
foreach (Sla::allTargets() as $target) {
    // $target['bucket']            → 'incident' | 'request'
    // $target['priority']          → 'Critical' | 'High' | ...
    // $target['hours']             → SLA hours
    // $target['task_type_match']   → ['Incident', 'Incident Task'] أو ['Catalog Task']
    // $target['task_type_label']   → 'بلاغ' | 'طلب'
}""")

    # ── 9. Update process ──────────────────────────────────────────────────
    heading(doc, "٩. خطوات تحديث هذا الملف عند صدور نسخة جديدة من الدليل", level=1)
    bullet(doc, 'استبدال الثوابت في app/Support/Sla.php.')
    bullet(doc, 'تحديث الجداول في الأقسام ٣ و ٤ هنا بالقيم الجديدة.')
    bullet(doc, 'تشغيل python3 docs/build_sla_policy.py لإعادة توليد هذا الملف.')
    bullet(doc, 'تشغيل php artisan reports:generate-dashboards --dashboard=service-tasks --no-screenshots لتحديث تقرير المهام.')

    # ── Footer ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "─── نهاية الوثيقة ───", size=9, color=GRAY_500)

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "المصدر: دليل السياسات والإجراءات للمشرف والتقني في نظام ساعد — عمادة التعلم الإلكتروني وتقنية المعلومات، جامعة القصيم", size=9, italic=True, color=GRAY_500)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"✅ Generated: {OUT}")


if __name__ == "__main__":
    build()

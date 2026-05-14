from docx import Document

from .base import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> dict:
        doc = Document(file_path)
        sections = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        full_text = "\n".join(sections)

        return {
            "content": full_text,
            "page_count": None,
            "ocr_applied": False,
            "language_detected": self.detect_language(full_text),
        }
